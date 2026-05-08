"""Build APA-formatted docx variants for the HEXACO Workplace Harassment Microsim paper.

Generates four output artifacts:
  - manuscript_preprint.docx        : single-file APA preprint (Times New Roman 12pt, double-spaced)
  - manuscript_journal.docx         : RSOS Stage 2 RR journal format (slightly different metadata block)
  - split/01_title_declarations.docx
  - split/02_body.docx               (Abstract + Introduction + Methods + Results + Discussion)
  - split/03_tables.docx             (Tables 1-6)
  - split/04_figures.docx            (Figure captions; figure files referenced separately)

Usage:
    python build_docx.py

Reads:
  - 01_intro.md, 02_methods.md, 03_results.md, 04_discussion.md, 05_refs.md, 06_tables_figures.md

All content sourced from markdown files; no inline duplication.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ============================================================================
# Paths and constants
# ============================================================================

PAPER_DIR = Path(__file__).resolve().parent
SPLIT_DIR = PAPER_DIR / "split"
SPLIT_DIR.mkdir(exist_ok=True)

OUT_PREPRINT = PAPER_DIR / "manuscript_preprint.docx"
OUT_JOURNAL = PAPER_DIR / "manuscript_journal.docx"
OUT_SPLIT_01 = SPLIT_DIR / "01_title_declarations.docx"
OUT_SPLIT_02 = SPLIT_DIR / "02_body.docx"
OUT_SPLIT_03 = SPLIT_DIR / "03_tables.docx"
OUT_SPLIT_04 = SPLIT_DIR / "04_figures.docx"

MD_INTRO = PAPER_DIR / "01_intro.md"
MD_METHODS = PAPER_DIR / "02_methods.md"
MD_RESULTS = PAPER_DIR / "03_results.md"
MD_DISCUSSION = PAPER_DIR / "04_discussion.md"
MD_REFS = PAPER_DIR / "05_refs.md"
MD_TABLES_FIGURES = PAPER_DIR / "06_tables_figures.md"

TITLE = (
    "Person-Level versus System-Level Anti-Harassment Interventions: "
    "A HEXACO 7-Typology Counterfactual Microsimulation "
    "in Japanese Workplaces"
)
RUNNING_HEAD = "HEXACO COUNTERFACTUAL MICROSIMULATION"  # APA 7 ≤ 50 chars (37 chars)
AUTHOR = "Eisuke Tokiwa"
ORCID = "0009-0009-7124-6669"
AFFILIATION = "SUNBLAZE Co., Ltd., Tokyo, Japan"
EMAIL = "eisuke.tokiwa@sunblaze.jp"
OSF_DOI = "10.17605/OSF.IO/3Y54U"
OSF_URL = "https://osf.io/3y54u"
GITHUB_URL = "https://github.com/etoki/paper"


# ============================================================================
# Style helpers (APA 7th edition; Times New Roman 12pt, double-spaced, 1" margins)
# ============================================================================


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool | None = None,
             italic: bool | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_double_space(p):
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             align: int | None = None, indent_first: bool = False, size: int = 12,
             style: str | None = None) -> "Paragraph":
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_first:
        p.paragraph_format.first_line_indent = Inches(0.5)
    set_double_space(p)
    if text:
        # If text contains markdown inline markers (**bold**, *italic*, `code`),
        # split into multi-run with proper formatting; otherwise single run.
        if re.search(r"\*\*|`|(?<!\w)\*(?!\w)", text) or re.search(r"\*[^*]+\*", text):
            render_inline(p, text, base_bold=bold, base_italic=italic, size=size)
        else:
            run = p.add_run(text)
            set_font(run, size=size, bold=bold or None, italic=italic or None)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    """Heading style (Level 1 = bold centered; Level 2/3 = bold left, no italic).

    Note: strict APA 7 specifies Level 3 as bold italic; this implementation
    follows the author's stylistic preference (no italic at any heading level)
    for compatibility with journals that diverge from strict APA 7 heading
    italics. Level 1: centered bold; Level 2: left bold; Level 3+: left bold,
    no indentation, no terminal period.
    """
    p = doc.add_paragraph()
    set_double_space(p)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=12, bold=True)
    else:  # level 2, 3, 4 — all flush-left bold, no italic
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, size=12, bold=True)
    return p


def configure_styles(doc: Document):
    """Set default Normal style to APA 7th: Times New Roman 12pt, double-spaced."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_running_head(doc: Document, running_head: str = RUNNING_HEAD):
    """APA 7 professional-paper running head: short title (≤ 50 chars,
    typically uppercase) flush-left; page number flush-right; appears on
    every page including title page.
    """
    assert len(running_head) <= 50, (
        f"Running head exceeds APA 7 limit (50 chars): '{running_head}' = {len(running_head)} chars"
    )
    for section in doc.sections:
        header = section.header
        # Clear any default
        for p in list(header.paragraphs):
            p.clear()
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Use a tab stop at right margin for page number
        p.paragraph_format.tab_stops.add_tab_stop(
            section.page_width - section.left_margin - section.right_margin,
            WD_TAB_ALIGNMENT.RIGHT,
        )
        run_left = p.add_run(running_head)
        set_font(run_left, size=12, bold=False)
        # Tab + page number field
        run_tab = p.add_run("\t")
        set_font(run_tab, size=12)
        # PAGE field code
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        page_run = p.add_run()
        page_run._r.append(fld_begin)
        page_run._r.append(instr)
        page_run._r.append(fld_sep)
        page_run._r.append(fld_end)
        set_font(page_run, size=12)


# ============================================================================
# Markdown → docx conversion (lightweight, paper-specific)
# ============================================================================

# Lines we want to skip (markdown front-matter, our internal notes)
SKIP_PATTERNS = [
    re.compile(r"^\s*---\s*$"),  # horizontal rules
]


def parse_md_blocks(md_path: Path) -> list[tuple[str, str]]:
    """Yield (block_type, text) tuples from a markdown file.

    Block types: h1, h2, h3, h4, para, table, list_item, blockquote.
    Tables are returned as a single multi-line string for downstream parsing.
    """
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    blocks: list[tuple[str, str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if any(p.match(line) for p in SKIP_PATTERNS):
            i += 1
            continue
        if line.startswith("#### "):
            blocks.append(("h4", line[5:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
        elif line.startswith("|") and "|" in line[1:]:
            # Markdown table: collect contiguous lines starting with |
            tbl_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith("|"):
                tbl_lines.append(lines[j])
                j += 1
            blocks.append(("table", "\n".join(tbl_lines)))
            i = j
        elif line.startswith("> "):
            blocks.append(("blockquote", line[2:].strip()))
            i += 1
        elif line.startswith("- ") or line.startswith("* "):
            # List item; group consecutive
            list_lines = [line[2:]]
            j = i + 1
            while j < len(lines) and (lines[j].startswith("- ") or lines[j].startswith("* ")):
                list_lines.append(lines[j][2:])
                j += 1
            for li in list_lines:
                blocks.append(("list_item", li.strip()))
            i = j
        elif line.strip() == "":
            i += 1
        else:
            # Plain paragraph; collect until blank line / structural marker
            para_lines = [line]
            j = i + 1
            while j < len(lines):
                nxt = lines[j]
                if nxt.strip() == "" or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("- ") or nxt.startswith("> "):
                    break
                para_lines.append(nxt)
                j += 1
            blocks.append(("para", " ".join(line.strip() for line in para_lines)))
            i = j
    return blocks


def render_inline(run_or_paragraph, text: str, *, base_bold: bool = False,
                  base_italic: bool = False, size: int = 12):
    """Render `text` as one or more runs, parsing inline markdown markers.

    Supported markers:
        **bold**       → bold run
        *italic*       → italic run
        `code`         → plain run (backticks stripped)
        plain text     → plain run

    Multiple markers can appear in one paragraph; each becomes its own run with
    appropriate font properties.

    Accepts either a Run (for single-run / table-cell legacy callers) or a
    Paragraph (for multi-run inline rendering).
    """
    # Detect Paragraph vs Run by checking for `add_run` attribute
    if hasattr(run_or_paragraph, "add_run"):
        p = run_or_paragraph
        # Tokenize: capture markdown spans, leave plain text between
        pattern = r"(\*\*[^*\n]+?\*\*|(?<![*])\*(?!\s|\*)[^*\n]+?(?<![\s*])\*(?!\*)|`[^`\n]+?`)"
        parts = re.split(pattern, text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                set_font(r, size=size, bold=True, italic=base_italic or None)
            elif part.startswith("*") and part.endswith("*"):
                r = p.add_run(part[1:-1])
                set_font(r, size=size, bold=base_bold or None, italic=True)
            elif part.startswith("`") and part.endswith("`"):
                r = p.add_run(part[1:-1])
                set_font(r, size=size, bold=base_bold or None, italic=base_italic or None)
            else:
                r = p.add_run(part)
                set_font(r, size=size, bold=base_bold or None, italic=base_italic or None)
        return
    # Legacy Run-based path: strip markdown markers, set whole-run formatting
    run = run_or_paragraph
    text = text.strip()
    if text.startswith("**") and text.endswith("**") and text.count("**") == 2:
        run.text = text[2:-2]
        run.bold = True
    elif text.startswith("*") and text.endswith("*") and text.count("*") == 2:
        run.text = text[1:-1]
        run.italic = True
    else:
        text = re.sub(r"\*\*([^\*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^\*]+)\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        run.text = text


def render_table(doc: Document, table_md: str):
    """Render a markdown pipe table as a native docx table with APA borders."""
    rows = [r for r in table_md.split("\n") if r.strip()]
    # Strip outer pipes and split each row
    parsed_rows: list[list[str]] = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed_rows.append(cells)
    # Detect separator row (---|---|---) and skip it
    parsed_rows = [r for r in parsed_rows if not all(re.match(r"^[-:\s]+$", c) for c in r)]
    if not parsed_rows:
        return
    n_cols = len(parsed_rows[0])
    table = doc.add_table(rows=len(parsed_rows), cols=n_cols)
    for i, cells in enumerate(parsed_rows):
        for j, cell in enumerate(cells[:n_cols]):
            tcell = table.cell(i, j)
            # Replace existing paragraph text
            p = tcell.paragraphs[0]
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            run = p.add_run()
            render_inline(run, cell)
            set_font(run, size=11, bold=(i == 0))
    add_para(doc, "")  # spacer after table


def render_blocks(doc: Document, blocks: list[tuple[str, str]], *, title_seen: bool = True):
    """Map markdown heading levels to APA 7 paper levels.

    Mapping (per author's request: Level 1 = centered, no italic anywhere):
      markdown #     (h1)  → paper Level 1 (centered, bold)
      markdown ##    (h2)  → paper Level 1 (centered, bold)   ← top-level sections
      markdown ###   (h3)  → paper Level 2 (left, bold)        ← subsections
      markdown ####  (h4)  → paper Level 3 (left, bold)        ← sub-subsections
    """
    for kind, text in blocks:
        if kind == "h1":
            add_heading(doc, text, level=1)
            title_seen = True
        elif kind == "h2":
            add_heading(doc, text, level=1)  # treat as Level 1 = centered
        elif kind == "h3":
            add_heading(doc, text, level=2)
        elif kind == "h4":
            add_heading(doc, text, level=3)
        elif kind == "para":
            add_para(doc, text, indent_first=True)
        elif kind == "list_item":
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run()
            render_inline(run, text)
            set_font(run, size=12)
            set_double_space(p)
        elif kind == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run()
            render_inline(run, text)
            set_font(run, size=12, italic=True)
            set_double_space(p)
        elif kind == "table":
            render_table(doc, text)


# ============================================================================
# Title page builder (APA + journal variants)
# ============================================================================


def build_title_page(doc: Document, *, journal_variant: bool = False):
    """APA 7 title page modeled on metaanalysis/paper/split/01_title_declarations.docx.

    Layout:
      - Title (centered, bold, ~ middle of page)
      - Author name (centered, plain — NOT italic)
      - Affiliation (centered, plain — NOT italic)
      - Author Note (centered, bold) header
      - Author Note content: ORCID line, COI line, correspondence line
      - Declarations (left-aligned, bold) header
      - Sub-sections: Conflict of Interest, Funding, Ethics Approval,
        Informed Consent, Pre-registration, Availability of data and material,
        Authors' contributions, Preprint Statement, Acknowledgments
    """
    # === Title block (centered) ===
    for _ in range(2):
        add_para(doc, "")  # spacer
    add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, AFFILIATION, align=WD_ALIGN_PARAGRAPH.CENTER)  # plain (not italic)
    add_para(doc, "")
    add_para(doc, "")

    # === Author Note (centered header, left-aligned content) ===
    add_para(doc, "Author Note", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"{AUTHOR}  https://orcid.org/{ORCID}")
    add_para(doc,
             "The author discloses a potential competing interest as detailed in the "
             "Declarations below.")
    add_para(doc,
             f"Correspondence concerning this article should be addressed to {AUTHOR}, "
             f"{AFFILIATION}. Email: {EMAIL}")
    add_para(doc, "")
    add_para(doc, "")

    # === Declarations (per metaanalysis template) ===
    add_para(doc, "Declarations", bold=True)
    add_para(doc, "")

    add_para(doc, "Competing Interests", bold=True)
    add_para(doc,
             "The author discloses the following potential competing interest: the "
             "author is the representative of SUNBLAZE Co., Ltd. (Tokyo, Japan), which "
             "provides HEXACO-JP, a proprietary Japanese-language HEXACO-based "
             "personality assessment service. Three points qualify the disclosure. "
             "(i) The present study does not use HEXACO-JP. The HEXACO instrument used "
             "here is the HEXACO-PI-R Japanese 60-item adaptation published by "
             "Wakabayashi (2014), a separately validated and openly cited instrument "
             "that predates and is independent of HEXACO-JP. (ii) The present "
             "manuscript does not reference, evaluate, market, recommend, or otherwise "
             "promote HEXACO-JP or any other SUNBLAZE product or service, and SUNBLAZE "
             "Co., Ltd. did not commission, sponsor, or direct the research. (iii) The "
             "headline empirical findings (a null personality-intervention contrast, a "
             "positive structural-intervention contrast, and the reframing of HEXACO "
             "7-typology as a stratification variable rather than an intervention or "
             "screening target) do not differentially favor commercial HEXACO-based "
             "assessment products.")
    add_para(doc, "")

    add_para(doc, "Funding", bold=True)
    add_para(doc,
             "No funding was received for this research. SUNBLAZE Co., Ltd. did not "
             "provide any funding, equipment, or compensation, and no external grant "
             "funding from public, commercial, or not-for-profit agencies was received. "
             "The author conducted the research as an independent academic project at "
             "no monetary cost, using publicly available analytical software and "
             "previously collected data.")
    add_para(doc, "")

    add_para(doc, "Ethics Approval", bold=True)
    add_para(doc,
             "The N = 354 individual-level harassment data re-analyzed in this study "
             "were originally collected under an IRB-approved protocol described in "
             "Tokiwa (2025, Research Square preprint, DOI 10.21203/rs.3.rs-7756124/v1). "
             "The present study is a secondary analysis of these de-identified data "
             "and does not introduce any new contact with participants; consistent with "
             "standard practice for secondary analysis of de-identified data covered by "
             "the original IRB approval, no additional ethics review was sought.")
    add_para(doc, "")

    add_para(doc, "Informed Consent", bold=True)
    add_para(doc,
             "All N = 354 participants in the original survey provided informed consent "
             "for research use of their de-identified responses (per Tokiwa, 2025). The "
             "present secondary analysis introduces no new participant contact.")
    add_para(doc, "")

    add_para(doc, "Pre-registration", bold=True)
    add_para(doc,
             f"This Stage 2 Registered Report implements, without deviation, the analysis "
             f"plan pre-registered as v2.0 at the Open Science Framework "
             f"(DOI {OSF_DOI}; {OSF_URL}; registered 2026-04-30). The Section 6.5 "
             f"Level 1 Methods Clarifications Log accompanying the v2.0 registration "
             f"documents 16 minor specification clarifications applied prior to Stage 0 "
             f"implementation.")
    add_para(doc, "")

    add_para(doc, "Availability of data and material", bold=True)
    add_para(doc,
             f"Public-tier supplementary artifacts (Stage 0–8 HDF5, Figures 1–6 in "
             f"PNG/PDF/SVG, canonical numerical record, SHA-256 reference hashes) are "
             f"openly available at the v2.0 OSF working project (osf.io/3hxz6, "
             f"v2.0/v2.0_supplementary.tar.gz). The N = 354 individual-level dataset is "
             f"governed by the IRB-approved data-sharing protocol and hosted in a Private "
             f"OSF component with a documented Request-Access mechanism (parent project "
             f"Wiki at osf.io/3hxz6/wiki/home/ lists access criteria, IRB requirements, "
             f"and data-use terms). All analysis code is publicly archived under MIT "
             f"license at {GITHUB_URL} (directory: simulation/).")
    add_para(doc, "")

    add_para(doc, "Authors' contributions", bold=True)
    add_para(doc,
             "The sole author (Eisuke Tokiwa) was responsible for conceptualization, "
             "methodology, software, formal analysis, investigation, data curation, "
             "writing — original draft, and writing — review & editing.")
    add_para(doc, "")

    add_para(doc, "Preprint Statement", bold=True)
    add_para(doc,
             f"A preprint of the present manuscript is publicly available at "
             f"SocArXiv (DOI 10.31235/osf.io/p2d8w_v1; "
             f"https://osf.io/preprints/socarxiv/p2d8w_v1). The underlying N = 354 "
             f"dataset and its descriptive analyses are available as a separate "
             f"preprint at Research Square (DOI 10.21203/rs.3.rs-7756124/v1; "
             f"Tokiwa, 2025). The seven HEXACO cluster centroids used as fixed "
             f"parameters in this study are from the companion paper Tokiwa (2026, "
             f"IEEE Access, DOI 10.1109/ACCESS.2026.3651324).")
    add_para(doc, "")

    add_para(doc, "Acknowledgments", bold=True)
    add_para(doc,
             "The author thanks the anonymous external methodologist (mode B; mathematical "
             "biology background) whose review memo led to the Section 6.5 Level 1 Methods "
             "Clarifications Log, and the N = 354 survey respondents whose anonymized "
             "data made this analysis possible.")

    if journal_variant:
        add_para(doc, "")
        add_para(doc,
                 "Note (journal submission): This manuscript is being submitted as a Stage 2 "
                 f"Registered Report. Conditional acceptance based on Stage 1 review of the "
                 f"v2.0 protocol document ({OSF_URL}) is acknowledged in the accompanying "
                 f"cover letter.")
    doc.add_page_break()


def build_declarations_page(doc: Document):
    """No-op: Declarations are now built inline in build_title_page() per the
    metaanalysis/paper/ template, which combines title page + author note +
    declarations on the same opening page-block."""
    return


# ============================================================================
# Document builders
# ============================================================================


def build_preprint(out_path: Path, *, journal_variant: bool = False):
    """Single-file APA preprint (or journal variant)."""
    doc = Document()
    configure_styles(doc)
    add_running_head(doc)
    build_title_page(doc, journal_variant=journal_variant)
    build_declarations_page(doc)

    # Body: Abstract + Introduction + Methods + Results + Discussion
    for md_path in [MD_INTRO, MD_METHODS, MD_RESULTS, MD_DISCUSSION]:
        blocks = parse_md_blocks(md_path)
        # Drop the entire markdown front-matter: everything before the first
        # "## " (H2) heading is treated as metadata and stripped (this includes
        # the file's "# 01. ..." H1 wrapper, the bold-prefixed metadata lines
        # like "**Author**", "**Pre-registration**", any parenthetical reader
        # notes, and the "---" front-matter separator). Body rendering begins
        # at the first H2 (e.g., "## Abstract").
        filtered = []
        body_started = False
        for kind, text in blocks:
            if not body_started:
                if kind == "h2":
                    body_started = True
                else:
                    continue
            filtered.append((kind, text))
        render_blocks(doc, filtered)
        doc.add_page_break()

    # References
    add_heading(doc, "References", level=1)
    refs_blocks = parse_md_blocks(MD_REFS)
    # Skip the file's own "# 05. References ..." header
    refs_filtered = [(k, t) for k, t in refs_blocks if not (k == "h1" and "References" in t)]
    for kind, text in refs_filtered:
        if kind == "para":
            # Hanging indent for references
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            set_double_space(p)
            run = p.add_run()
            render_inline(run, text)
            set_font(run, size=12)
        elif kind == "h2" or kind == "h3":
            add_heading(doc, text, level=1 if kind == "h2" else 2)
        elif kind == "list_item":
            # Treat as reference item with hanging indent
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            set_double_space(p)
            run = p.add_run()
            render_inline(run, text)
            set_font(run, size=12)
    doc.add_page_break()

    # Tables
    add_heading(doc, "Tables", level=1)
    tf_blocks = parse_md_blocks(MD_TABLES_FIGURES)
    in_tables = False
    in_figures = False
    for kind, text in tf_blocks:
        if kind == "h2" and text.strip().lower().startswith("tables"):
            in_tables = True
            in_figures = False
            continue
        if kind == "h2" and text.strip().lower().startswith("figures"):
            in_tables = False
            in_figures = True
            add_heading(doc, "Figures", level=1)
            continue
        if kind == "h2" and "generation notes" in text.lower():
            in_tables = False
            in_figures = False
            continue
        if in_tables or in_figures:
            if kind == "h3":
                add_heading(doc, text, level=2)
            elif kind == "table":
                render_table(doc, text)
            elif kind == "para":
                add_para(doc, text, italic=text.lower().startswith("note") or text.lower().startswith("caption"))
            elif kind == "list_item":
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run()
                render_inline(run, text)
                set_font(run, size=12)
                set_double_space(p)

    doc.save(str(out_path))
    print(f"  Wrote {out_path}")


def build_split_01(out_path: Path):
    doc = Document()
    configure_styles(doc)
    add_running_head(doc)
    build_title_page(doc)
    build_declarations_page(doc)
    doc.save(str(out_path))
    print(f"  Wrote {out_path}")


def build_split_02(out_path: Path):
    doc = Document()
    configure_styles(doc)
    add_running_head(doc)
    for md_path in [MD_INTRO, MD_METHODS, MD_RESULTS, MD_DISCUSSION]:
        blocks = parse_md_blocks(md_path)
        filtered = []
        skip_metadata = True
        for kind, text in blocks:
            if skip_metadata and (kind == "h1" or (kind == "para" and ("OSF DOI" in text or "Working title" in text or text.startswith("**Author**") or text.startswith("**Pre-registration**") or text.startswith("**Reporting standard**")))):
                continue
            skip_metadata = False
            filtered.append((kind, text))
        render_blocks(doc, filtered)
        doc.add_page_break()
    # References live in the body for split-02
    add_heading(doc, "References", level=1)
    refs_blocks = parse_md_blocks(MD_REFS)
    refs_filtered = [(k, t) for k, t in refs_blocks if not (k == "h1" and "References" in t)]
    for kind, text in refs_filtered:
        if kind == "para":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            set_double_space(p)
            run = p.add_run()
            render_inline(run, text)
            set_font(run, size=12)
        elif kind == "h2" or kind == "h3":
            add_heading(doc, text, level=1 if kind == "h2" else 2)
    doc.save(str(out_path))
    print(f"  Wrote {out_path}")


def build_split_03(out_path: Path):
    """Tables only."""
    doc = Document()
    configure_styles(doc)
    add_running_head(doc)
    add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Tables", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    tf_blocks = parse_md_blocks(MD_TABLES_FIGURES)
    in_tables = False
    for kind, text in tf_blocks:
        if kind == "h2" and text.strip().lower().startswith("tables"):
            in_tables = True
            continue
        if kind == "h2" and (text.strip().lower().startswith("figures") or "generation" in text.lower()):
            in_tables = False
            continue
        if in_tables:
            if kind == "h3":
                add_heading(doc, text, level=2)
            elif kind == "table":
                render_table(doc, text)
            elif kind == "para":
                add_para(doc, text, italic=text.lower().startswith("note"))
    doc.save(str(out_path))
    print(f"  Wrote {out_path}")


def build_split_04(out_path: Path):
    """Figure captions only (figure files referenced separately)."""
    doc = Document()
    configure_styles(doc)
    add_running_head(doc)
    add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Figures (captions)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    tf_blocks = parse_md_blocks(MD_TABLES_FIGURES)
    in_figures = False
    for kind, text in tf_blocks:
        if kind == "h2" and text.strip().lower().startswith("figures"):
            in_figures = True
            continue
        if kind == "h2" and "generation" in text.lower():
            in_figures = False
            continue
        if in_figures:
            if kind == "h3":
                add_heading(doc, text, level=2)
            elif kind == "para":
                add_para(doc, text, italic=text.lower().startswith("caption"))
    doc.save(str(out_path))
    print(f"  Wrote {out_path}")


def main():
    print("[build_docx] HEXACO Workplace Harassment Microsim manuscript builder")
    print(f"  Source markdown: {PAPER_DIR}/0[1-6]_*.md")
    print(f"  Output preprint: {OUT_PREPRINT}")
    print(f"  Output journal: {OUT_JOURNAL}")
    print(f"  Split outputs: {SPLIT_DIR}/")
    print()
    build_preprint(OUT_PREPRINT, journal_variant=False)
    build_preprint(OUT_JOURNAL, journal_variant=True)
    build_split_01(OUT_SPLIT_01)
    build_split_02(OUT_SPLIT_02)
    build_split_03(OUT_SPLIT_03)
    build_split_04(OUT_SPLIT_04)
    print()
    print("[build_docx] Done.")


if __name__ == "__main__":
    main()
