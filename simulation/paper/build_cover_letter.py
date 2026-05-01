"""Build the cover letter docx for the HEXACO Workplace Harassment Microsim paper.

Reads cover_letter.md and generates cover_letter.docx with APA-ish business-letter
formatting (Times New Roman 12pt, single-spaced, 1-inch margins, block paragraphs).

Note: Journal name is left as the placeholder 'XXX X' for the investigator to substitute
at submission time.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PAPER_DIR = Path(__file__).resolve().parent

# Build matrix: each cover letter source markdown → output docx.
# The "template" (cover_letter.md) keeps the placeholder *XXXX* journal name and
# generic suggested reviewers; the four journal-specific variants are tailored
# to PAID / RSOS / JOHP / PCI RR with the title that matches each journal's fit
# (T2.2 / T2.1 / T4.1 / T1.1 respectively per title_candidates.md).
COVER_LETTER_BUILDS = [
    ("cover_letter.md",        "cover_letter.docx"),
    ("cover_letter_paid.md",   "cover_letter_paid.docx"),
    ("cover_letter_rsos.md",   "cover_letter_rsos.docx"),
    ("cover_letter_johp.md",   "cover_letter_johp.docx"),
    ("cover_letter_pcirr.md",  "cover_letter_pcirr.docx"),
]


def set_font(run, name="Times New Roman", size=12, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def configure(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def strip_md(text: str) -> str:
    """Strip simple markdown markers for plain text rendering."""
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)  # links
    text = re.sub(r"\*\*([^\*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^\*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def add_para(doc, text, *, bold=False, italic=False, align=None, single_space=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.15 if single_space else 2.0
    pf.space_after = Pt(6)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold or None, italic=italic or None)
    return p


def build_one(src_filename: str, out_filename: str):
    """Build a single cover_letter*.docx from its markdown source."""
    src = PAPER_DIR / src_filename
    out = PAPER_DIR / out_filename
    if not src.exists():
        print(f"  Skipping (source missing): {src_filename}")
        return

    doc = Document()
    configure(doc)

    md_text = src.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    # Skip the first H1 ("# Cover Letter") and the front-matter block
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# Cover Letter"):
            i += 1
            continue
        if line.startswith("**Submission Type:") or line.startswith("**Manuscript:") \
                or line.startswith("**Author:") or line.startswith("**Pre-registration:"):
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            break
        i += 1

    # Now render the body
    in_list = False
    while i < len(lines):
        line = lines[i]

        # H2 sections
        if line.startswith("## "):
            heading = strip_md(line[3:])
            add_para(doc, heading, bold=True)
            i += 1
            continue

        # Numbered or bullet list
        if line.strip().startswith("- ") or re.match(r"^\d+\.\s", line.strip()):
            content = re.sub(r"^[-*]\s+", "", line.strip())
            content = re.sub(r"^\d+\.\s+", "", content)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.15
            pf.left_indent = Inches(0.25)
            pf.space_after = Pt(3)
            run = p.add_run("• " + strip_md(content))
            set_font(run)
            i += 1
            continue

        # Italicized journal placeholder line "*XXX X*" already handled in body
        # Date marker [Date]
        if line.strip() == "[Date]":
            add_para(doc, "[Date]")
            i += 1
            continue

        # Editor block
        if line.strip() == "The Editors":
            add_para(doc, "The Editors")
            i += 1
            continue

        # Plain paragraph (multi-line until blank)
        if line.strip() == "":
            i += 1
            continue

        # Collect contiguous non-empty lines as paragraph
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() != "" \
                and not lines[j].startswith("## ") \
                and not lines[j].strip().startswith("- ") \
                and not re.match(r"^\d+\.\s", lines[j].strip()):
            para_lines.append(lines[j])
            j += 1
        para_text = " ".join(strip_md(l) for l in para_lines)
        if para_text:
            # Italicize journal name placeholder line
            if para_text == "XXXX":
                add_para(doc, para_text, italic=True)
            else:
                add_para(doc, para_text)
        i = j

    doc.save(str(out))
    print(f"  Wrote {out}")


def main():
    """Build all cover letter variants (template + 4 journal-specific)."""
    for src_filename, out_filename in COVER_LETTER_BUILDS:
        build_one(src_filename, out_filename)


if __name__ == "__main__":
    print("[build_cover_letter] Building cover letter docx variants ...")
    main()
    print("[build_cover_letter] Done.")
