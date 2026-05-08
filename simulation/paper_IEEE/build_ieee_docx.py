"""Build an IEEE Access–format docx for the HEXACO Workplace Harassment
Microsimulation paper.

Reuses the existing APA markdown sources (paper/01_intro.md, ..., paper/05_refs.md)
as a single source of truth. Performs three transformations to produce the
IEEE Access manuscript:

1. Layout: 8.00" × 10.875" page, 0.89/0.72/0.51 margins, single-column title
   block + 2-column body (per the clustering/paper_IEEE template).
2. Heading hierarchy: section headings → 9pt bold ALL CAPS;
   subsection headings → 10pt bold mixed case;
   body text → 10pt Times New Roman justified.
3. Citations: replace APA-style (Author, Year) with IEEE-style numbered [N];
   reformat the references list in IEEE author-initials style with order of
   first appearance.

Outputs:
    paper_IEEE/manuscript_ieee.docx

Usage:
    cd simulation
    uv run python paper_IEEE/build_ieee_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips, RGBColor

# ============================================================================
# Paths and constants
# ============================================================================

PAPER_DIR = Path(__file__).resolve().parent
SIM_PAPER_DIR = PAPER_DIR.parent / "paper"
OUT_DOCX = PAPER_DIR / "manuscript_ieee.docx"

MD_INTRO = SIM_PAPER_DIR / "01_intro.md"
MD_METHODS = SIM_PAPER_DIR / "02_methods.md"
MD_RESULTS = SIM_PAPER_DIR / "03_results.md"
MD_DISCUSSION = SIM_PAPER_DIR / "04_discussion.md"
MD_REFS = SIM_PAPER_DIR / "05_refs.md"
MD_TABLES_FIGURES = SIM_PAPER_DIR / "06_tables_figures.md"

# Manuscript metadata
TITLE = (
    "Person-Level versus System-Level Anti-Harassment Interventions: "
    "A HEXACO 7-Typology Counterfactual Microsimulation "
    "in Japanese Workplaces"
)
AUTHOR_LASTNAME = "TOKIWA"
AUTHOR_FULLNAME = "Eisuke Tokiwa"
AUTHOR_INITIAL = "E. Tokiwa"
AFFILIATION = (
    "SUNBLAZE Co., Ltd., Nishi-Shinjuku 3-3-13, Shinjuku-ku, "
    "Tokyo 160-0023, Japan"
)
EMAIL = "eisuke.tokiwa@sunblaze.jp"
ORCID = "0009-0009-7124-6669"
FUNDING = (
    "No funding was received for this research. SUNBLAZE Co., Ltd. did not "
    "provide any funding, equipment, or compensation, and no external grant "
    "funding from public, commercial, or not-for-profit agencies was received. "
    "The author conducted the research as an independent academic project at "
    "no monetary cost."
)

INDEX_TERMS = (
    "Bootstrap, counterfactual analysis, HEXACO model, intersection-union "
    "test, microsimulation, target trial emulation, workplace harassment."
)

# ============================================================================
# IEEE Access page setup constants (twips: 1 inch = 1440 twips)
# ============================================================================

PAGE_WIDTH_TWIPS = 11520    # 8.00"
PAGE_HEIGHT_TWIPS = 15660   # 10.875"
TOP_MARGIN_TWIPS = 1280     # 0.89"
BOTTOM_MARGIN_TWIPS = 1040  # 0.72"
LEFT_MARGIN_TWIPS = 740     # 0.51"
RIGHT_MARGIN_TWIPS = 740    # 0.51"
HEADER_TWIPS = 360
FOOTER_TWIPS = 640

COL_WIDTH_TWIPS = 4820      # 3.35" per column
COL_GAP_TWIPS = 400         # 0.28" gap

# ============================================================================
# References: IEEE-format mapping (numbered by order of first appearance)
# ============================================================================

# APA citation key → IEEE [N] number; IEEE-style formatted reference text.
# Order of N is determined by first appearance in body markdown.
# (Reference text adapted from 05_refs.md per IEEE Access bibliography style:
#  Initials. Lastname, "Title," Journal, vol. X, no. Y, pp. X–Y, Year, doi: ...)

IEEE_REFS: dict[str, str] = {
    # Author-Year key (matches APA in-text) → IEEE-style reference text
    "Hershcovis & Barling, 2010":
        "M. S. Hershcovis and J. Barling, \"Towards a multi-foci approach to "
        "workplace aggression: A meta-analytic review of outcomes from "
        "different perpetrators,\" J. Organ. Behav., vol. 31, no. 1, pp. "
        "24–44, 2010, doi: 10.1002/job.621.",
    "Salin, 2003":
        "D. Salin, \"Ways of explaining workplace bullying: A review of "
        "enabling, motivating and precipitating structures and processes in "
        "the work environment,\" Hum. Relat., vol. 56, no. 10, pp. "
        "1213–1232, 2003, doi: 10.1177/00187267035610003.",
    "MHLW, 2017, 2021, 2024":
        "Ministry of Health, Labour and Welfare (MHLW), \"Workplace "
        "harassment surveys: H28 / R2 / R5 results [Original report in "
        "Japanese],\" 2017, 2021, 2024. [Online]. Available: "
        "https://www.mhlw.go.jp/",
    "Mikkelsen et al., 2020":
        "E. G. Mikkelsen, Å. M. Hansen, R. Persson, M. F. Byrgesen, and "
        "A. Hogh, \"Individual consequences of being exposed to workplace "
        "bullying,\" in Bullying and Harassment in the Workplace: Theory, "
        "Research and Practice, 3rd ed., S. Einarsen, H. Hoel, D. Zapf, "
        "and C. L. Cooper, Eds. Boca Raton, FL: CRC Press, 2020, pp. "
        "163–208, doi: 10.1201/9780429462528-6.",
    "Salin, 2021":
        "D. Salin, \"Workplace bullying and gender: An overview of "
        "empirical findings,\" in Dignity and Inclusion at Work: Handbooks "
        "of Workplace Bullying, Emotional Abuse and Harassment, vol. 3, "
        "P. D'Cruz et al., Eds. Singapore: Springer, 2021, pp. 331–361, "
        "doi: 10.1007/978-981-13-0218-3_12.",
    "Nielsen et al., 2017":
        "M. B. Nielsen, L. Glasø, and S. Einarsen, \"Exposure to workplace "
        "harassment and the Five Factor Model of personality: A "
        "meta-analysis,\" Pers. Individ. Differ., vol. 104, pp. 195–206, "
        "2017, doi: 10.1016/j.paid.2016.08.015.",
    "Pilch & Turska, 2015":
        "I. Pilch and E. Turska, \"Relationships between Machiavellianism, "
        "organizational culture, and workplace bullying: Emotional abuse "
        "from the target's and the perpetrator's perspective,\" J. Bus. "
        "Ethics, vol. 128, no. 1, pp. 83–93, 2015, doi: "
        "10.1007/s10551-014-2081-3.",
    "Ashton & Lee, 2007":
        "M. C. Ashton and K. Lee, \"Empirical, theoretical, and practical "
        "advantages of the HEXACO model of personality structure,\" Pers. "
        "Soc. Psychol. Rev., vol. 11, no. 2, pp. 150–166, 2007, doi: "
        "10.1177/1088868306294907.",
    "Glasø et al., 2007":
        "L. Glasø, S. B. Matthiesen, M. B. Nielsen, and S. Einarsen, "
        "\"Do targets of workplace bullying portray a general victim "
        "personality profile?,\" Scand. J. Psychol., vol. 48, no. 4, pp. "
        "313–319, 2007, doi: 10.1111/j.1467-9450.2007.00554.x.",
    "Lee et al., 2013":
        "K. Lee, M. C. Ashton, J. Wiltshire, J. S. Bourdage, B. A. Visser, "
        "and A. Gallucci, \"Sex, power, and money: Prediction from the "
        "Dark Triad and Honesty-Humility,\" Eur. J. Pers., vol. 27, no. "
        "2, pp. 169–184, 2013, doi: 10.1002/per.1860.",
    "Linton & Power, 2013":
        "D. K. Linton and J. L. Power, \"The personality traits of "
        "workplace bullies are often shared by their victims: Is there a "
        "dark side to victims?,\" Pers. Individ. Differ., vol. 54, no. 6, "
        "pp. 738–743, 2013, doi: 10.1016/j.paid.2012.11.026.",
    "Power et al., 2013":
        "J. L. Power, et al., \"Acceptability of workplace bullying: A "
        "comparative study on six continents,\" J. Bus. Res., vol. 66, "
        "no. 3, pp. 374–380, 2013, doi: 10.1016/j.jbusres.2011.08.018.",
    "Tokiwa, 2026":
        "E. Tokiwa, \"Cultural influences on personality types: A cluster "
        "analysis of HEXACO traits in Japan,\" IEEE Access, 2026, doi: "
        "10.1109/ACCESS.2026.3651324.",
    "Tokiwa, 2025":
        "E. Tokiwa, \"Predicting power and gender harassment from HEXACO "
        "and Dark Traits,\" Research Square, Preprint, version 1, 2025, "
        "doi: 10.21203/rs.3.rs-7756124/v1.",
    "Spielauer, 2011":
        "M. Spielauer, \"What is social science microsimulation?,\" Soc. "
        "Sci. Comput. Rev., vol. 29, no. 1, pp. 9–20, 2011, doi: "
        "10.1177/0894439310370085.",
    "Rutter et al., 2011":
        "C. M. Rutter, A. M. Zaslavsky, and E. J. Feuer, \"Dynamic "
        "microsimulation models for health outcomes: A review,\" Med. "
        "Decis. Making, vol. 31, no. 1, pp. 10–18, 2011, doi: "
        "10.1177/0272989X10369005.",
    "Hernán & Robins, 2016":
        "M. A. Hernán and J. M. Robins, \"Using big data to emulate a "
        "target trial when a randomized trial is not available,\" Am. J. "
        "Epidemiol., vol. 183, no. 8, pp. 758–764, 2016, doi: "
        "10.1093/aje/kwv254.",
    "Hernán & Robins, 2020":
        "M. A. Hernán and J. M. Robins, Causal Inference: What If. Boca "
        "Raton, FL: Chapman & Hall/CRC, 2020.",
    "Hudgens & Halloran, 2008":
        "M. G. Hudgens and M. E. Halloran, \"Toward causal inference with "
        "interference,\" J. Amer. Statist. Assoc., vol. 103, no. 482, pp. "
        "832–842, 2008, doi: 10.1198/016214508000000292.",
    "Pearl, 2009":
        "J. Pearl, Causality: Models, Reasoning, and Inference, 2nd ed. "
        "Cambridge, U.K.: Cambridge Univ. Press, 2009.",
    "Hudson & Fraley, 2015":
        "N. W. Hudson and R. C. Fraley, \"Volitional personality trait "
        "change: Can people choose to change their personality traits?,\" "
        "J. Pers. Soc. Psychol., vol. 109, no. 3, pp. 490–507, 2015, "
        "doi: 10.1037/pspp0000021.",
    "Roberts et al., 2017":
        "B. W. Roberts, J. Luo, D. A. Briley, P. I. Chow, R. Su, and "
        "P. L. Hill, \"A systematic review of personality trait change "
        "through intervention,\" Psychol. Bull., vol. 143, no. 2, pp. "
        "117–141, 2017, doi: 10.1037/bul0000088.",
    "Escartín, 2016":
        "J. Escartín, \"Insights into workplace bullying: Psychosocial "
        "drivers and effective interventions,\" Psychol. Res. Behav. "
        "Manage., vol. 9, pp. 157–169, 2016, doi: 10.2147/PRBM.S91211.",
    "Hodgins et al., 2014":
        "M. Hodgins, S. MacCurtain, and P. Mannix-McNamara, \"Workplace "
        "bullying and incivility: A systematic review of interventions,\" "
        "Int. J. Workplace Health Manage., vol. 7, no. 1, pp. 54–72, "
        "2014, doi: 10.1108/IJWHM-08-2013-0030.",
    "Statistics Bureau, 2023":
        "Statistics Bureau, Ministry of Internal Affairs and Communications, "
        "\"Labour Force Survey (Basic Tabulation): 2022 Annual Average "
        "Results Summary,\" Tokyo, Japan, 2023. [Online]. Available: "
        "https://www.e-stat.go.jp/",
    "Wakabayashi, 2014":
        "A. Wakabayashi, \"A sixth personality domain that is independent "
        "of the Big Five domains: The psychometric properties of the "
        "HEXACO Personality Inventory in a Japanese sample,\" Jpn. "
        "Psychol. Res., vol. 56, no. 3, pp. 211–223, 2014, doi: "
        "10.1111/jpr.12045.",
    "Ashton & Lee, 2009":
        "M. C. Ashton and K. Lee, \"The HEXACO-60: A short measure of "
        "the major dimensions of personality,\" J. Pers. Assess., vol. "
        "91, no. 4, pp. 340–345, 2009, doi: 10.1080/00223890902935878.",
    "Tou et al., 2017":
        "K. Tou, A. Tsuda, M. Nii, T. Yamahiro, and M. Irie, \"Development "
        "of new power harassment questionnaire in workplace,\" in Proc. "
        "81st Annu. Conv. Japan. Psychol. Assoc., 2017.",
    "Kobayashi & Tanaka, 2010":
        "A. Kobayashi and K. Tanaka, \"Development of the Gender Harassment "
        "Scale,\" Jpn. J. Ind. Organ. Psychol., vol. 24, no. 1, pp. 15–27, "
        "2010, doi: 10.32222/jaiop.24.1_15.",
    "Efron, 1987":
        "B. Efron, \"Better bootstrap confidence intervals,\" J. Amer. "
        "Statist. Assoc., vol. 82, no. 397, pp. 171–185, 1987, doi: "
        "10.1080/01621459.1987.10478410.",
    "Carlin & Louis, 2000":
        "B. P. Carlin and T. A. Louis, Bayes and Empirical Bayes Methods "
        "for Data Analysis, 2nd ed. Boca Raton, FL: Chapman & Hall/CRC, "
        "2000.",
    "Page, 1963":
        "E. B. Page, \"Ordered hypotheses for multiple treatments: A "
        "significance test for linear ranks,\" J. Amer. Statist. Assoc., "
        "vol. 58, no. 301, pp. 216–230, 1963, doi: "
        "10.1080/01621459.1963.10500843.",
    "Podsakoff et al., 2003":
        "P. M. Podsakoff, S. B. MacKenzie, J.-Y. Lee, and N. P. Podsakoff, "
        "\"Common method biases in behavioral research: A critical review "
        "of the literature and recommended remedies,\" J. Appl. Psychol., "
        "vol. 88, no. 5, pp. 879–903, 2003, doi: "
        "10.1037/0021-9010.88.5.879.",
    "Lindell & Whitney, 2001":
        "M. K. Lindell and D. J. Whitney, \"Accounting for common method "
        "variance in cross-sectional research designs,\" J. Appl. Psychol., "
        "vol. 86, no. 1, pp. 114–121, 2001, doi: 10.1037/0021-9010.86.1.114.",
    "Berger & Hsu, 1996":
        "R. L. Berger and J. C. Hsu, \"Bioequivalence trials, "
        "intersection-union tests and equivalence confidence sets,\" "
        "Statist. Sci., vol. 11, no. 4, pp. 283–319, 1996, doi: "
        "10.1214/ss/1032280304.",
    "Schelling, 1971":
        "T. C. Schelling, \"Dynamic models of segregation,\" J. Math. "
        "Sociol., vol. 1, no. 2, pp. 143–186, 1971, doi: "
        "10.1080/0022250X.1971.9989794.",
    "Park et al., 2023":
        "J. S. Park, J. C. O'Brien, C. J. Cai, M. R. Morris, P. Liang, "
        "and M. S. Bernstein, \"Generative agents: Interactive simulacra "
        "of human behavior,\" in Proc. 36th Annu. ACM Symp. User Interface "
        "Softw. Technol. (UIST), 2023, doi: 10.1145/3586183.3606763.",
}


# ============================================================================
# Citation extraction + IEEE numbering
# ============================================================================

def extract_apa_citations(text: str) -> list[str]:
    """Find all APA-style citations like (Author, Year), (A & B, Year),
    (A et al., Year), or narrative-form 'Author (Year)' / 'Author et al. (Year)'.
    Returns a list in order of appearance.
    """
    cits: list[str] = []
    # Parenthetical: (Author, Year), (A & B, Year), (A et al., Year),
    # (A, B, & C, Year), or multi-cite (...; ...)
    paren_pat = re.compile(r"\(([^()]*?\b\d{4}[a-z]?\b[^()]*?)\)")
    # Narrative: Author (Year), A & B (Year), A et al. (Year)
    name_chars = r"A-Za-zÀ-ÿ\-'"
    narr_pat = re.compile(
        rf"((?:[A-Z][{name_chars}]+(?:\s*&\s*[A-Z][{name_chars}]+|\s+et\s+al\.|\s*,\s*[A-Z][{name_chars}]+\s*&\s*[A-Z][{name_chars}]+)?))\s*\((\d{{4}}[a-z]?)\b[^)]*?\)"
    )

    for m in paren_pat.finditer(text):
        body = m.group(1)
        for piece in re.split(r"\s*;\s*", body):
            piece = piece.strip()
            if re.search(r"\b\d{4}\b", piece):
                cits.append(piece)
    for m in narr_pat.finditer(text):
        author = m.group(1).strip()
        year = m.group(2).strip()
        cits.append(f"{author}, {year}")
    return cits


def normalize_citation_key(cit: str) -> str | None:
    """Return the IEEE_REFS dict key matching a raw citation string, or None
    if no mapping is found."""
    cit = cit.strip()
    cit = re.sub(r"\s+", " ", cit)
    cit = cit.replace(" and ", " & ")  # normalize narrative form
    # Direct match
    if cit in IEEE_REFS:
        return cit
    # Try simplified: drop trailing parts after the year
    simplified = re.sub(r"(\d{4}[a-z]?)[,\s]*.*$", r"\1", cit).strip().rstrip(",")
    if simplified in IEEE_REFS:
        return simplified
    # Author + year only matching
    m = re.match(r"(.+?),\s*(\d{4}[a-z]?)\b", cit)
    if m:
        author_part = m.group(1).strip()
        year = m.group(2)
        candidate = f"{author_part}, {year}"
        if candidate in IEEE_REFS:
            return candidate
    # Try author & year substring search across IEEE_REFS keys
    yr_match = re.search(r"\b(\d{4})\b", cit)
    if yr_match:
        year = yr_match.group(1)
        for key in IEEE_REFS:
            # Match if key starts with same first surname AND has same year
            cit_words = cit.split()
            if not cit_words:
                continue
            first_surname = re.split(r"[ ,]", cit)[0]
            if key.startswith(first_surname) and year in key:
                return key
    return None


def build_citation_index(body_text: str) -> tuple[dict[str, int], list[str]]:
    """Walk the body, assign [N] in order of first appearance, return:
        - mapping: {ieee_ref_key: N}
        - ordered list of (ieee_ref_key) by N
    """
    mapping: dict[str, int] = {}
    order: list[str] = []
    for cit in extract_apa_citations(body_text):
        key = normalize_citation_key(cit)
        if key and key not in mapping:
            mapping[key] = len(order) + 1
            order.append(key)
    return mapping, order


def replace_citations_with_ieee(text: str, mapping: dict[str, int]) -> str:
    """Replace APA citations in `text` with IEEE numbered [N] form."""
    # Parenthetical (single or multi-citation)
    def paren_repl(match: re.Match) -> str:
        body = match.group(1)
        pieces = re.split(r"\s*;\s*", body)
        nums: list[int] = []
        leftover_pieces: list[str] = []
        for piece in pieces:
            piece = piece.strip()
            if not piece:
                continue
            key = normalize_citation_key(piece)
            if key and key in mapping:
                nums.append(mapping[key])
            else:
                leftover_pieces.append(piece)
        if nums and not leftover_pieces:
            nums = sorted(set(nums))
            # IEEE Editorial Style Manual: separate multiple citations with
            # comma between brackets, e.g., [1], [2], [3] — NOT [1, 2, 3].
            return ", ".join(f"[{n}]" for n in nums)
        if nums and leftover_pieces:
            nums = sorted(set(nums))
            cite_str = ", ".join(f"[{n}]" for n in nums)
            return f"{cite_str} (" + "; ".join(leftover_pieces) + ")"
        # No matched citations; return original
        return match.group(0)

    paren_pat = re.compile(r"\(([^()]*?\b\d{4}[a-z]?\b[^()]*?)\)")
    text = paren_pat.sub(paren_repl, text)

    # Narrative: Author (Year) → Author [N]
    def narr_repl(match: re.Match) -> str:
        author = match.group(1).strip()
        year = match.group(2).strip()
        cit = f"{author}, {year}"
        key = normalize_citation_key(cit)
        if key and key in mapping:
            return f"{author} [{mapping[key]}]"
        return match.group(0)

    name_chars = r"A-Za-zÀ-ÿ\-'"
    narr_pat = re.compile(
        rf"([A-Z][{name_chars}]+(?:\s*&\s*[A-Z][{name_chars}]+|\s+et\s+al\.)?)\s*\((\d{{4}}[a-z]?)\b[^)]*?\)"
    )
    text = narr_pat.sub(narr_repl, text)
    return text


# ============================================================================
# DOCX rendering helpers (IEEE Access format)
# ============================================================================

def set_font(run, name="Times New Roman", size=10, bold=None, italic=None,
             smallcaps=False):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if smallcaps:
        run.font.small_caps = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_para_format(p, *, align=None, single_space=True, space_before=0,
                    space_after=0, first_line_indent=None):
    pf = p.paragraph_format
    pf.line_spacing = 1.0 if single_space else 1.15
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if align is not None:
        p.alignment = align


def configure_page(doc):
    sec = doc.sections[0]
    sec.page_width = Twips(PAGE_WIDTH_TWIPS)
    sec.page_height = Twips(PAGE_HEIGHT_TWIPS)
    sec.top_margin = Twips(TOP_MARGIN_TWIPS)
    sec.bottom_margin = Twips(BOTTOM_MARGIN_TWIPS)
    sec.left_margin = Twips(LEFT_MARGIN_TWIPS)
    sec.right_margin = Twips(RIGHT_MARGIN_TWIPS)
    sec.header_distance = Twips(HEADER_TWIPS)
    sec.footer_distance = Twips(FOOTER_TWIPS)
    # Set default Normal style font
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_section_columns(section, num_cols: int):
    """Set the number of columns for a section by editing sectPr."""
    sectPr = section._sectPr
    # Remove any existing <w:cols>
    for cols in sectPr.findall(qn("w:cols")):
        sectPr.remove(cols)
    cols = OxmlElement("w:cols")
    if num_cols == 1:
        cols.set(qn("w:num"), "1")
    else:
        cols.set(qn("w:num"), str(num_cols))
        cols.set(qn("w:equalWidth"), "0")
        cols.set(qn("w:space"), str(COL_GAP_TWIPS))
        for i in range(num_cols):
            col = OxmlElement("w:col")
            col.set(qn("w:w"), str(COL_WIDTH_TWIPS))
            col.set(qn("w:space"), str(COL_GAP_TWIPS) if i < num_cols - 1 else "0")
            cols.append(col)
    sectPr.append(cols)


def add_continuous_section_break(doc, num_cols: int):
    """Add a continuous section break and configure the new section's column
    count. Returns the new (current-last) section object."""
    # Insert a paragraph that contains the section break properties
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    sectType = OxmlElement("w:type")
    sectType.set(qn("w:val"), "continuous")
    sectPr.append(sectType)
    # Page size & margins inherit from previous section by default; explicitly set
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), str(PAGE_WIDTH_TWIPS))
    pgSz.set(qn("w:h"), str(PAGE_HEIGHT_TWIPS))
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), str(TOP_MARGIN_TWIPS))
    pgMar.set(qn("w:bottom"), str(BOTTOM_MARGIN_TWIPS))
    pgMar.set(qn("w:left"), str(LEFT_MARGIN_TWIPS))
    pgMar.set(qn("w:right"), str(RIGHT_MARGIN_TWIPS))
    sectPr.append(pgMar)
    # Add cols
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    if num_cols > 1:
        cols.set(qn("w:equalWidth"), "0")
        cols.set(qn("w:space"), str(COL_GAP_TWIPS))
        for i in range(num_cols):
            col = OxmlElement("w:col")
            col.set(qn("w:w"), str(COL_WIDTH_TWIPS))
            col.set(qn("w:space"), str(COL_GAP_TWIPS) if i < num_cols - 1 else "0")
            cols.append(col)
    sectPr.append(cols)
    pPr.append(sectPr)


# ============================================================================
# Markdown body parser → IEEE rendering
# ============================================================================

def strip_markdown_inline(text: str) -> str:
    """Strip basic markdown inline markers (already handled elsewhere in
    main APA build, but required here for IEEE rendering)."""
    text = re.sub(r"\*\*([^*\n]+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`\n]+?)`", r"\1", text)
    return text


def add_inline_runs(p, text: str, *, base_size=10, base_bold=False,
                   base_italic=False):
    """Add runs to paragraph p, parsing **bold** / *italic* / `code` markers
    so the docx output reflects the inline emphasis without literal markers."""
    pattern = (
        r"(\*\*[^*\n]+?\*\*|"
        r"(?<![*])\*(?!\s|\*)[^*\n]+?(?<![\s*])\*(?!\*)|"
        r"`[^`\n]+?`)"
    )
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, size=base_size, bold=True, italic=base_italic or None)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            set_font(r, size=base_size, bold=base_bold or None, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_font(r, size=base_size, bold=base_bold or None,
                     italic=base_italic or None)
        else:
            r = p.add_run(part)
            set_font(r, size=base_size, bold=base_bold or None,
                     italic=base_italic or None)


def parse_md_blocks(md_text: str) -> list[tuple[str, str]]:
    """Yield (kind, text) blocks. kinds: h1, h2, h3, h4, para, list_item, table_md."""
    blocks: list[tuple[str, str]] = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# ") and not line.startswith("## "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("#### "):
            blocks.append(("h4", line[5:].strip()))
            i += 1
        elif line.strip().startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            # Table
            tbl_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                tbl_lines.append(lines[j])
                j += 1
            blocks.append(("table_md", "\n".join(tbl_lines)))
            i = j
        elif line.strip().startswith("- ") or re.match(r"^\d+\.\s", line.strip()):
            blocks.append(("list_item", re.sub(r"^[\d-]+\.?\s+", "",
                                                line.strip(), count=1)))
            i += 1
        elif line.strip() == "" or line.strip() == "---":
            i += 1
        else:
            # Paragraph
            para_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip() != "" and \
                    not lines[j].startswith("#") and \
                    not lines[j].strip().startswith("- ") and \
                    not re.match(r"^\d+\.\s", lines[j].strip()) and \
                    not lines[j].strip().startswith("|"):
                para_lines.append(lines[j])
                j += 1
            text = " ".join(l.strip() for l in para_lines).strip()
            if text:
                blocks.append(("para", text))
            i = j
    return blocks


def to_roman(n: int) -> str:
    """Integer → Roman numeral (1..30)."""
    pairs = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for v, s in pairs:
        while n >= v:
            out += s
            n -= v
    return out


def to_letter(n: int) -> str:
    """Integer → uppercase letter (1..26)."""
    return chr(ord("A") + n - 1) if 1 <= n <= 26 else f"AA{n}"


def add_ieee_section_heading(doc, text: str, *, roman: int | None = None,
                             centered: bool = True):
    """9pt bold ALL CAPS section heading. If `roman` is given, prefixes the
    heading with Roman numerals "I. ", "II. ", etc., and centers the heading
    per IEEE Editorial Style Manual. Unnumbered top-level sections (REFERENCES,
    ACKNOWLEDGMENT, BIOGRAPHIES) pass roman=None.
    """
    p = doc.add_paragraph()
    align = WD_ALIGN_PARAGRAPH.CENTER if centered else None
    set_para_format(p, align=align, space_before=8, space_after=4)
    label = f"{to_roman(roman)}. " if roman is not None else ""
    r = p.add_run(label + text.upper())
    set_font(r, size=9, bold=True)


def add_ieee_subsection_heading(doc, text: str, *, letter: int | None = None):
    """10pt bold mixed-case subsection heading.
    Prefixed with a letter "A. ", "B. ", ... per IEEE Style Manual.
    """
    p = doc.add_paragraph()
    set_para_format(p, space_before=4, space_after=2)
    label = f"{to_letter(letter)}. " if letter is not None else ""
    r = p.add_run(label + text)
    set_font(r, size=10, bold=True)


def add_ieee_subsubsection_heading(doc, text: str, *, number: int | None = None):
    """10pt bold italic run-in style sub-subsection heading.
    Prefixed with "1) ", "2) ", ... and ends with colon (per IEEE Style Manual).
    """
    p = doc.add_paragraph()
    set_para_format(p, space_before=2, space_after=2)
    label = f"{number}) " if number is not None else ""
    r = p.add_run(label + text + ":")
    set_font(r, size=10, bold=True, italic=True)


def add_ieee_body_para(doc, text: str, *, indent=True):
    """10pt justified body paragraph with first-line indent."""
    p = doc.add_paragraph()
    set_para_format(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Twips(180) if indent else None,
    )
    add_inline_runs(p, text, base_size=10)


def add_ieee_list_item(doc, text: str):
    """10pt body bullet list item."""
    p = doc.add_paragraph(style="List Bullet")
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_inline_runs(p, text, base_size=10)


def add_ieee_table(doc, table_md: str):
    """Render markdown pipe table as IEEE-style table (small font)."""
    rows = [r for r in table_md.split("\n") if r.strip()]
    parsed_rows: list[list[str]] = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed_rows.append(cells)
    parsed_rows = [r for r in parsed_rows
                   if not all(re.match(r"^[-:\s]+$", c) for c in r)]
    if not parsed_rows:
        return
    n_cols = len(parsed_rows[0])
    table = doc.add_table(rows=len(parsed_rows), cols=n_cols)
    for i, cells in enumerate(parsed_rows):
        for j, cell in enumerate(cells[:n_cols]):
            tcell = table.cell(i, j)
            p = tcell.paragraphs[0]
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            add_inline_runs(p, strip_markdown_inline(cell), base_size=8,
                           base_bold=(i == 0))


# ============================================================================
# Main builder
# ============================================================================

# Skip patterns: section headings we don't want in IEEE body or duplicates
SKIP_HEADINGS = {
    # File-level h1 wrappers in our markdown
    "01. Abstract + Introduction", "02. Methods", "03. Results",
    "04. Discussion + Limitations", "05. References (APA 7th edition)",
    "06. Tables and Figures",
    # Front-matter / metadata
    "Abstract",  # rendered inline as ABSTRACT below
}

# Map our section / subsection text to IEEE-friendly equivalents
HEADING_MAPPING = {
    "Introduction": "INTRODUCTION",
    "Methods": "METHODS",
    "Results": "RESULTS",
    "Discussion": "DISCUSSION",
    "Conclusion": "CONCLUSION",
    "References": "REFERENCES",
}


def build():
    doc = Document()
    configure_page(doc)
    # NOTE on python-docx section model:
    # In Word's sectPr layout, a section break paragraph carries the sectPr
    # for the section ENDING at that break, while the document's trailing
    # sectPr applies to the FINAL section. So our pattern is:
    #   1. Render the title block (page-wide / 1-column).
    #   2. Insert a continuous section break whose sectPr declares cols=1
    #      (binds the title block as 1-column).
    #   3. Render the body (multi-column).
    #   4. After all content, set the final (trailing) sectPr to cols=2.

    # === Title block (single column) ===
    # Manuscript ID block (italic 8pt) — review/version metadata placeholder
    p = doc.add_paragraph()
    set_para_format(p, space_before=0, space_after=0)
    r = p.add_run(
        "Received XX Month, 20XX; revised XX Month, 20XX; accepted XX Month, "
        "20XX. Date of publication XX Month, 20XX; date of current version "
        "XX Month, 20XX."
    )
    set_font(r, size=8, italic=True)

    p = doc.add_paragraph()
    set_para_format(p, space_before=0, space_after=8)
    r = p.add_run("Digital Object Identifier 10.1109/ACCESS.20XX.DOI")
    set_font(r, size=8, italic=True)

    # Title
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    r = p.add_run(TITLE)
    set_font(r, size=22, bold=True)

    # Author + superscript
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    r = p.add_run(AUTHOR_INITIAL.upper())
    set_font(r, size=10, bold=True)
    r2 = p.add_run("1")
    set_font(r2, size=7, bold=True)
    r2.font.superscript = True

    # Affiliation
    p = doc.add_paragraph()
    set_para_format(p, space_after=2)
    r = p.add_run("1")
    set_font(r, size=7)
    r.font.superscript = True
    r2 = p.add_run(AFFILIATION)
    set_font(r2, size=7)

    # Corresponding author
    p = doc.add_paragraph()
    set_para_format(p, space_after=2)
    r = p.add_run(f"Corresponding author: {AUTHOR_INITIAL} (e-mail: {EMAIL}).")
    set_font(r, size=8)

    # Funding
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    r = p.add_run(FUNDING)
    set_font(r, size=8)

    # First-page IEEE Access copyright + CC-BY notice
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    r = p.add_run(
        "© 2026 The Authors. This work is licensed under a Creative Commons "
        "Attribution 4.0 License. For more information, see "
        "https://creativecommons.org/licenses/by/4.0/. "
        "VOLUME XX, 2026"
    )
    set_font(r, size=7)

    # === Read body markdown ===
    intro_md = MD_INTRO.read_text(encoding="utf-8")
    methods_md = MD_METHODS.read_text(encoding="utf-8")
    results_md = MD_RESULTS.read_text(encoding="utf-8")
    discussion_md = MD_DISCUSSION.read_text(encoding="utf-8")
    body_md = "\n\n".join([intro_md, methods_md, results_md, discussion_md])

    # Build citation index from full body
    cit_mapping, cit_order = build_citation_index(body_md)

    # === Abstract (single column, inline bold "ABSTRACT") ===
    # Extract abstract text from intro_md
    abs_match = re.search(r"## Abstract\s*\n\n(.+?)\n\n", intro_md, re.DOTALL)
    abstract_text = abs_match.group(1).strip() if abs_match else ""
    abstract_text = strip_markdown_inline(abstract_text)
    abstract_text = replace_citations_with_ieee(abstract_text, cit_mapping)

    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    r = p.add_run("ABSTRACT")
    set_font(r, size=10, bold=True)
    r2 = p.add_run(" " + abstract_text)
    set_font(r2, size=10)

    # === Index Terms (inline bold) ===
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
    r = p.add_run("INDEX TERMS")
    set_font(r, size=10, bold=True)
    r2 = p.add_run(" " + INDEX_TERMS)
    set_font(r2, size=10)

    # === Section break: title block ends here as 1-column ===
    add_continuous_section_break(doc, num_cols=1)

    # === Body section structure ===
    # IEEE Access uses Roman-numeral major sections (I., II., III., ...).
    # Each markdown file becomes one major section; the ## headings inside
    # each file become A., B., C., ... subsections; ### headings become
    # 1), 2), ... sub-subsections.
    # Conclusion is extracted from discussion_md and rendered as its own
    # major section (V. CONCLUSION).
    SECTION_PLAN = [
        ("INTRODUCTION", intro_md),
        ("METHODS", methods_md),
        ("RESULTS", results_md),
        ("DISCUSSION", discussion_md),
    ]

    def _strip_metadata(md_text: str) -> str:
        md_text = re.sub(r"^# 0\d\..*?\n", "", md_text)
        md_text = re.sub(
            r"^\*\*(?:Working title|Author|Pre-registration|Reporting standard).*?\n",
            "", md_text, flags=re.MULTILINE)
        md_text = re.sub(r"## Abstract\s*\n\n.*?\n\n\*\*Keywords.*?\n",
                         "", md_text, flags=re.DOTALL)
        md_text = re.sub(r"^\*\*Keywords.*?\n", "", md_text, flags=re.MULTILINE)
        md_text = re.sub(r"^---\s*$", "", md_text, flags=re.MULTILINE)
        return md_text

    def _extract_conclusion(discussion_md: str) -> tuple[str, str]:
        """Split discussion_md into (discussion_without_conclusion, conclusion_md).
        The Conclusion subsection becomes its own major section in IEEE format.
        """
        # Match `## Conclusion\n\n...` until the end of file
        m = re.search(r"\n(##\s+Conclusion\b.*?)$", discussion_md,
                      re.DOTALL | re.IGNORECASE)
        if m:
            return discussion_md[:m.start()].rstrip(), m.group(1).lstrip()
        return discussion_md, ""

    discussion_body, conclusion_md = _extract_conclusion(discussion_md)
    SECTION_PLAN[3] = ("DISCUSSION", discussion_body)
    SECTION_PLAN.append(("CONCLUSION", conclusion_md))

    # === Render body sections ===
    for roman_idx, (section_label, md_text) in enumerate(SECTION_PLAN, start=1):
        if not md_text.strip():
            continue
        md_text = _strip_metadata(md_text)
        md_text = replace_citations_with_ieee(md_text, cit_mapping)

        # Emit Roman-numeral major section heading
        add_ieee_section_heading(doc, section_label, roman=roman_idx)

        # Walk blocks, applying letter / number prefix to ## / ### / ####
        letter_counter = 0   # A. B. C. ...
        number_counter = 0   # 1) 2) 3) ...
        last_h2_was_conclusion_root = False

        blocks = parse_md_blocks(md_text)
        for kind, txt in blocks:
            txt_clean = strip_markdown_inline(txt) if kind != "table_md" else txt
            if kind == "h1":
                continue  # file-title h1 already stripped
            if kind == "h2":
                if txt_clean in SKIP_HEADINGS:
                    continue
                # If this is the section's bare CONCLUSION (== section_label),
                # don't repeat as subsection
                if txt_clean.upper() == section_label:
                    last_h2_was_conclusion_root = True
                    continue
                letter_counter += 1
                number_counter = 0
                add_ieee_subsection_heading(doc, txt_clean, letter=letter_counter)
            elif kind == "h3":
                number_counter += 1
                add_ieee_subsubsection_heading(doc, txt_clean, number=number_counter)
            elif kind == "h4":
                add_ieee_subsubsection_heading(doc, txt_clean)
            elif kind == "para":
                add_ieee_body_para(doc, txt_clean)
            elif kind == "list_item":
                add_ieee_list_item(doc, txt_clean)
            elif kind == "table_md":
                add_ieee_table(doc, txt)

    # === ACKNOWLEDGMENT (unnumbered, before References) ===
    add_ieee_section_heading(doc, "ACKNOWLEDGMENT", roman=None)
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(
        "The author thanks the anonymous external methodologist (mode B; "
        "mathematical biology background) whose review memo led to the Section "
        "6.5 Level 1 Methods Clarifications Log, and the N = 354 survey "
        "respondents whose anonymized data made this analysis possible. "
    )
    set_font(r, size=10)
    r = p.add_run(
        "Competing Interests: The author discloses the following potential "
        "competing interest: the author is the founder and representative of "
        "SUNBLAZE Co., Ltd. (Tokyo, Japan), which provides HEXACO-JP, a "
        "proprietary Japanese-language HEXACO-based personality assessment "
        "service. The present study does not use HEXACO-JP; the HEXACO "
        "instrument used here is the HEXACO-PI-R Japanese 60-item adaptation "
        "published by Wakabayashi (2014), a separately validated and openly "
        "cited instrument that predates and is independent of HEXACO-JP. The "
        "manuscript does not reference, evaluate, market, recommend, or "
        "otherwise promote HEXACO-JP or any other SUNBLAZE product or "
        "service, and the headline empirical findings (a null personality-"
        "intervention contrast and a positive structural-intervention "
        "contrast) do not differentially favor commercial HEXACO-based "
        "assessment products. "
    )
    set_font(r, size=10, italic=True)
    r = p.add_run(
        "Ethics statement: The N = 354 individual-level harassment data "
        "re-analyzed in this study were originally collected under an "
        "IRB-approved protocol described in [29]; the present secondary "
        "analysis of de-identified records does not require additional "
        "ethics review. "
    )
    set_font(r, size=10, italic=True)
    r = p.add_run(
        "Data availability: Public-tier supplementary artifacts (Stage 0–8 "
        "HDF5, Figures 1–6 in PNG/PDF/SVG, canonical numerical record, "
        "SHA-256 reference hashes) are openly available at the v2.0 OSF "
        "working project (osf.io/3hxz6, v2.0/v2.0_supplementary.tar.gz). "
        "The N = 354 individual-level dataset is governed by the "
        "IRB-approved data-sharing protocol and hosted in a Private OSF "
        "component with a documented Request-Access mechanism "
        "(osf.io/3hxz6/wiki/home/). "
    )
    set_font(r, size=10, italic=True)
    r = p.add_run(
        "AI use disclosure: Generative AI (Anthropic Claude) was used as a "
        "drafting and code-review assistant; all final claims, derivations, "
        "and statistical conclusions are the author's responsibility."
    )
    set_font(r, size=10, italic=True)

    # === References (IEEE-numbered) ===
    add_ieee_section_heading(doc, "REFERENCES", roman=None)
    for idx, key in enumerate(cit_order, start=1):
        ref_text = IEEE_REFS[key]
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
        r = p.add_run(f"[{idx}] ")
        set_font(r, size=9, bold=True)
        r2 = p.add_run(ref_text)
        set_font(r2, size=8)

    # === Author Biography (IEEE Access required, 100-150 words/author) ===
    add_ieee_section_heading(doc, "BIOGRAPHIES", roman=None)

    # Photo placeholder note (post-acceptance, IEEE provides typesetting frame)
    p = doc.add_paragraph()
    set_para_format(p, space_after=4)
    r = p.add_run("[Photo placeholder — 100×120 px grayscale or color, supplied separately at production stage.]")
    set_font(r, size=8, italic=True)

    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    r = p.add_run("EISUKE TOKIWA ")
    set_font(r, size=8, bold=True, smallcaps=True)
    r2 = p.add_run(
        f"(Member, IEEE) is the founder of SUNBLAZE Co., Ltd., Tokyo, Japan. "
        f"He received his graphic-design training in Japan and subsequently "
        f"redirected his work toward computational psychology and "
        f"workplace-organization research. His current research interests "
        f"span the HEXACO personality model, workplace harassment, target-"
        f"trial emulation in observational data, microsimulation pipelines "
        f"with empirical-Bayes shrinkage, and pre-registered counterfactual "
        f"causal analysis. Prior work includes a HEXACO 7-typology "
        f"clustering of N = 13,668 Japanese respondents (IEEE Access, 2026). "
        f"He is committed to fully reproducible, hash-verified analytical "
        f"pipelines (MIT-licensed, OSF-archived). ORCID: {ORCID}. Contact: "
        f"{EMAIL}."
    )
    set_font(r2, size=8)

    # === Final trailing section gets 2-column layout (applies to the body) ===
    set_section_columns(doc.sections[-1], 2)

    doc.save(str(OUT_DOCX))
    print(f"  Wrote {OUT_DOCX}")
    print(f"  Citations indexed: {len(cit_order)}")
    return cit_mapping, cit_order


def main():
    print("[build_ieee_docx] HEXACO Workplace Harassment Microsim — IEEE Access build")
    print(f"  Source markdown: {SIM_PAPER_DIR}/0[1-5]_*.md")
    print(f"  Output: {OUT_DOCX}")
    build()
    print("[build_ieee_docx] Done.")


if __name__ == "__main__":
    main()
