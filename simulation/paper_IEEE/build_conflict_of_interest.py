"""Build the standalone Conflict of Interest docx for IEEE Access submission.

Reads conflict_of_interest_ieee.md and produces conflict_of_interest_ieee.docx
with a clean, business-letter-style layout (Times New Roman 12pt, single-
spaced, 1-inch margins, plain headings) suitable for direct upload to the IEEE
Access submission portal as the required "single conflict of interest
document".

Usage:
    cd simulation
    uv run --extra paper python paper_IEEE/build_conflict_of_interest.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PAPER_IEEE_DIR = Path(__file__).resolve().parent
SRC = PAPER_IEEE_DIR / "conflict_of_interest_ieee.md"
OUT = PAPER_IEEE_DIR / "conflict_of_interest_ieee.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12,
            bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def configure(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def strip_md(text: str) -> str:
    """Strip basic markdown markers for plain text rendering."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)         # bold
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)  # italic
    text = re.sub(r"`([^`]+)`", r"\1", text)               # code
    return text.strip()


def add_para(doc: Document, text: str, *, bold: bool = False,
            italic: bool = False, align=None,
            indent: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    if indent is not None:
        pf.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold or None, italic=italic or None)


def render_inline_emph(p, text: str) -> None:
    """Add runs to paragraph p, parsing **bold** / *italic* markers so that
    the docx output reflects inline emphasis without literal markers."""
    pattern = (
        r"(\*\*[^*\n]+?\*\*|"
        r"(?<!\*)\*(?!\s|\*)[^*\n]+?(?<!\s)\*(?!\*))"
    )
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)


def build() -> None:
    if not SRC.exists():
        print(f"  Source missing: {SRC}")
        return

    doc = Document()
    configure(doc)

    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip blank lines (paragraph breaks already inserted by add_para).
        if not stripped:
            i += 1
            continue

        # Top-level title (# heading).
        if stripped.startswith("# "):
            add_para(doc, strip_md(stripped[2:]), bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # Section heading (## heading).
        if stripped.startswith("## "):
            add_para(doc, strip_md(stripped[3:]), bold=True)
            i += 1
            continue

        # Horizontal rule — render as a small spacer.
        if stripped == "---":
            add_para(doc, "")
            i += 1
            continue

        # Numbered list item (e.g., "1. **Bold heading.** Body text...").
        m_list = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m_list:
            content = m_list.group(2)
            # Collect continuation lines until blank.
            j = i + 1
            while j < n and lines[j].strip() and not (
                    lines[j].strip().startswith("#")
                    or re.match(r"^\d+\.\s+", lines[j].strip())
                    or lines[j].strip() == "---"):
                content += " " + lines[j].strip()
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.25)
            number_run = p.add_run(f"{m_list.group(1)}. ")
            set_font(number_run, bold=True)
            render_inline_emph(p, content)
            i = j
            continue

        # Plain paragraph: collect contiguous non-empty lines.
        para_lines = [stripped]
        j = i + 1
        while j < n and lines[j].strip() and not (
                lines[j].strip().startswith("#")
                or re.match(r"^\d+\.\s+", lines[j].strip())
                or lines[j].strip() == "---"):
            para_lines.append(lines[j].strip())
            j += 1
        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        render_inline_emph(p, para_text)
        i = j

    doc.save(str(OUT))
    print(f"  Wrote {OUT}")


def main() -> None:
    print("[build_conflict_of_interest] IEEE Access standalone COI document")
    print(f"  Source markdown: {SRC}")
    print(f"  Output: {OUT}")
    build()
    print("[build_conflict_of_interest] Done.")


if __name__ == "__main__":
    main()
