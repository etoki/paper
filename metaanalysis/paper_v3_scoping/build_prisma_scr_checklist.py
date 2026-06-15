"""
build_prisma_scr_checklist.py — Generate the PRISMA-ScR (Tricco et al.,
2018) 22-item checklist as a supplementary document for the
paper_v3_scoping reframing of the manuscript.

Scoping reviews are reported against PRISMA-ScR rather than PRISMA 2020.
The 22 items below follow Tricco et al. (2018), Annals of Internal
Medicine, 169(7), 467–473, with the "Location where item is reported"
column pointing to the section/subsection of
manuscript_journal_v3_scoping.docx.

Both a non-anonymous and an anonymous version are produced:
    prisma_scr_checklist.docx        — for Heliyon, Frontiers, MDPI, BMC
    prisma_scr_checklist_anon.docx   — for HSSC (double-anonymous review)
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HERE = Path(__file__).resolve().parent
OUT_NON_ANON = HERE / "prisma_scr_checklist.docx"
OUT_ANON = HERE / "prisma_scr_checklist_anon.docx"


def _set_run_font(run, font_name="Calibri", size_pt=10, bold=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_paragraph(doc, text, *, bold=False, size=10, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


def _osf_phrase(anonymous: bool, *, deposit_name: str) -> str:
    if anonymous:
        return (f"deposited as the {deposit_name} component of the time-"
                f"stamped open-science project [DOI provided post-"
                f"acceptance for double-anonymous review]")
    return ""


# PRISMA-ScR (Tricco et al., 2018) — 22 items
# Each tuple: (section, item_number, item_text, location_reported_non_anon)
# The location text deliberately mirrors the wording in the manuscript so
# the reviewer can cross-check rapidly.

def _build_items(anonymous: bool):
    osf_run = ("the search log is deposited on the associated OSF "
               "project (02_search component)" if not anonymous else
               "the search log is deposited as an anonymised "
               "supplementary file with this submission")
    return [
        # ---- TITLE ----
        ("TITLE", "", "", ""),
        ("Title", "1",
         "Identify the report as a scoping review.",
         "Title page — title explicitly identifies the report as "
         "\"A Scoping Review with Exploratory Quantitative Synthesis.\""),

        # ---- ABSTRACT ----
        ("ABSTRACT", "", "", ""),
        ("Structured summary", "2",
         "Provide a structured summary that includes (as applicable): "
         "background, objectives, eligibility criteria, sources of "
         "evidence, charting methods, results, and conclusions that "
         "relate to the review questions and objectives.",
         "Abstract — single-paragraph summary covers all PRISMA-ScR "
         "abstract elements; explicitly names PRISMA-ScR (Tricco et al., "
         "2018) as the reporting framework, names the executed sources "
         "of evidence and the unexecuted pre-registered subscription-"
         "gated databases, and identifies the conclusions as a charted "
         "map plus a research agenda rather than as established effect "
         "estimates."),

        # ---- INTRODUCTION ----
        ("INTRODUCTION", "", "", ""),
        ("Rationale", "3",
         "Describe the rationale for the review in the context of what "
         "is already known. Explain why the review questions/objectives "
         "lend themselves to a scoping review approach.",
         "Introduction → \"Personality and Academic Achievement in Face-"
         "to-Face Contexts\" and \"The Shift to Online Learning "
         "Environments\" subsections summarise eight prior face-to-face "
         "meta-analyses (Poropat 2009 through Chen et al. 2025) and "
         "establish the absence of an online-modality moderator as the "
         "stated gap. Munn et al. (2018) and the scoping-review "
         "reframing rationale are cited explicitly in the Information "
         "Sources subsection."),
        ("Objectives", "4",
         "Provide an explicit statement of the questions and objectives "
         "being addressed with reference to their key elements (e.g., "
         "population or participants, concepts, and context) or other "
         "relevant key elements used to conceptualize the review "
         "questions and/or objectives.",
         "Introduction → \"The Present Review\" subsection — three "
         "Mapping Questions (MQ1, MQ2, MQ3) and five directional "
         "Mapping Priors (MP1 through MP5) are stated explicitly. The "
         "PCC (Population–Concept–Context) frame replaces PICOS for "
         "this scoping review and is described in the Eligibility "
         "Criteria subsection."),

        # ---- METHODS ----
        ("METHODS", "", "", ""),
        ("Protocol and registration", "5",
         "Indicate whether a review protocol exists; state if and where "
         "it can be accessed (e.g., a Web address); and if available, "
         "provide registration information, including the registration "
         "number.",
         "Methods → \"Protocol and Registration\" subsection — protocol "
         f"time-stamped 23 April 2026 on OSF Registries; {osf_run}. "
         "Any deviations from the pre-registered protocol are "
         "transparently disclosed in the Methods → Information Sources "
         "subsection and in the Limitations section."),
        ("Eligibility criteria", "6",
         "Specify characteristics of the sources of evidence used as "
         "eligibility criteria (e.g., years considered, language, and "
         "publication status), and provide a rationale.",
         "Methods → \"Eligibility Criteria\" subsection — full "
         "Population–Concept–Context eligibility rules, including "
         "population (learners in online, blended, or MOOC "
         "environments), concept (Big Five or HEXACO personality "
         "instruments and academic achievement, satisfaction, "
         "engagement, or perception outcomes), context (post-secondary "
         "or K-12 online learning environments), and publication-type, "
         "language (English-only), and design restrictions."),
        ("Information sources", "7",
         "Describe all information sources (e.g., databases with dates "
         "of coverage and contact with authors to identify additional "
         "sources), as well as the date the most recent search was "
         "executed.",
         "Methods → \"Information Sources and Search Strategy\" "
         "subsection — pre-registered six databases (PubMed/MEDLINE, "
         "PsycINFO, ERIC, Web of Science, Scopus, ProQuest "
         "Dissertations) are named. Actually-executed sources: "
         "(a) WebSearch (Google Scholar-equivalent), 23 April 2026; "
         "(b) ERIC web search, 23 April 2026; (c) targeted retrieval "
         "from open-access repositories (PubMed Central, Frontiers, "
         "MDPI, Open Praxis); (d) forward and backward citation "
         "snowballing from included primary studies and from the "
         "eight benchmark face-to-face meta-analyses. PsycINFO, "
         "Scopus, Web of Science, and ProQuest Dissertations were "
         "not searched (institutional access unavailable; individual "
         "subscription not offered by Clarivate, Elsevier, or APA); "
         "this gap is the basis for reframing as a scoping review."),
        ("Search", "8",
         "Present the full electronic search strategy for at least 1 "
         "database, including any limits used, such that it could be "
         "repeated.",
         "Methods → \"Information Sources and Search Strategy\" "
         "subsection — three-concept Boolean strategy (personality × "
         "online learning × academic outcome) is reproduced in full. "
         "The complete per-query syntax, execution dates, and hit "
         f"counts are deposited; {osf_run}."),
        ("Selection of sources of evidence", "9",
         "State the process for selecting sources of evidence (i.e., "
         "screening and eligibility) included in the scoping review.",
         "Methods → \"Study Selection\" subsection — single reviewer "
         "under a pre-specified intra-rater reliability protocol with "
         "10% (title/abstract) and 20% (full-text) re-screening after "
         "a ≥7-day wash-out; target Cohen's κ ≥ 0.80 met at each "
         "stage. Single-reviewer workflow is discussed transparently "
         "as a limitation."),
        ("Data charting process", "10",
         "Describe the methods of charting data from the included "
         "sources of evidence (e.g., calibrated forms or forms that "
         "have been tested by the team before their use, and whether "
         "data charting was done independently or in duplicate) and "
         "any processes for obtaining and confirming data from "
         "investigators.",
         "Methods → \"Data Extraction\" subsection — single-reviewer "
         "extraction using a pre-piloted CSV-based form; intra-rater "
         "reliability sub-sample of 20% re-extracted after ≥7-day "
         "wash-out with target κ ≥ 0.80 met. Categorical fields and "
         "continuous fields are listed in the manuscript and on the "
         f"OSF project (04_extraction component). Where reported "
         "statistics permitted conversion to Pearson r, the Peterson "
         "and Brown (2005) β-to-r conversion was applied with a "
         "transparent two-predictor rule."),
        ("Data items", "11",
         "List and define all variables for which data were sought and "
         "any assumptions and simplifications made.",
         "Methods → \"Data Extraction\" subsection and Table 1 list all "
         "extracted variables: study identification, country, journal, "
         "sample size, age, gender composition, education level, "
         "sampling method, design, modality, modality subtype, "
         "personality instrument and reliability, outcome instrument, "
         "effect-size statistics for each Big Five trait, era coding, "
         "and JBI risk-of-bias score. Assumptions and simplifications "
         "for HEXACO-to-Big-Five mapping and β-to-r conversion are "
         "stated explicitly."),
        ("Critical appraisal of individual sources of evidence", "12",
         "If done, provide a rationale for conducting a critical "
         "appraisal of included sources of evidence; describe the "
         "methods used and how this information was used in any data "
         "synthesis (if appropriate).",
         "Methods → \"Risk-of-Bias Assessment\" subsection — Joanna "
         "Briggs Institute (JBI) checklist for analytical cross-"
         "sectional studies was applied with single-reviewer rating "
         "and 20% re-rating intra-rater check. Risk-of-bias scores "
         "are reported per-study in Table 1 and used as an exclusion "
         "criterion in a pre-specified sensitivity analysis (Table "
         "4). Reframing as a scoping review does not eliminate the "
         "appraisal; it re-positions the appraisal as a descriptive "
         "feature of the mapped evidence."),
        ("Synthesis of results", "13",
         "Describe the methods of handling and summarizing the data "
         "that were charted.",
         "Methods → \"Quantitative Synthesis\" subsection — random-"
         "effects meta-analysis with REML τ² estimation and Hartung-"
         "Knapp-Sidik-Jonkman confidence-interval adjustment, on the "
         "Fisher z scale, for the exploratory quantitative pool "
         "(k = 10). The synthesis is explicitly described as "
         "exploratory rather than confirmatory, and prediction "
         "intervals (not just confidence intervals) are reported for "
         "each trait."),

        # ---- RESULTS ----
        ("RESULTS", "", "", ""),
        ("Selection of sources of evidence", "14",
         "Give numbers of sources of evidence screened, assessed for "
         "eligibility, and included in the review, with reasons for "
         "exclusions at each stage, ideally using a flow diagram.",
         "Results → \"Study Selection\" subsection and Figure 1 (a "
         "PRISMA-ScR-style flow diagram) — numbers of records "
         "identified, screened, retrieved, assessed for eligibility, "
         "and included are presented. Full-text exclusion reasons are "
         "reported."),
        ("Characteristics of sources of evidence", "15",
         "For each included source of evidence, present characteristics "
         "for which data were charted and provide the citations.",
         "Results → \"Characteristics of Retained Studies\" subsection "
         "and Table 1 — each of the 25 scoping-retained studies is "
         "characterised by country, modality, sample, instrument, "
         "outcome, era, region, JBI risk-of-bias score, and inclusion "
         "status (exploratory-pool vs. scoping-only)."),
        ("Critical appraisal within sources of evidence", "16",
         "If done, present data on critical appraisal of included "
         "sources of evidence (see item 12).",
         "Results → \"Risk-of-Bias Distribution\" paragraph and Table "
         "1 (per-study risk-of-bias score column). Sensitivity "
         "analysis excluding RoB-low studies is reported in Table 4."),
        ("Results of individual sources of evidence", "17",
         "For each included source of evidence, present the relevant "
         "data that were charted that relate to the review questions "
         "and objectives.",
         "Results → \"Per-Study Effect Sizes\" subsection — per-trait "
         "Pearson correlations (direct or β-converted) are reported "
         "for each of the 10 studies contributing to the exploratory "
         "quantitative pool, with sample size, outcome type, and "
         "instrument explicitly noted."),
        ("Synthesis of results", "18",
         "Summarize and/or present the charting results as they relate "
         "to the review questions and objectives.",
         "Results → \"Exploratory Pooled Effect Sizes\" and \"Subgroup "
         "Contrasts\" subsections, Tables 2 and 3 — pooled trait-by-"
         "trait estimates, 95% CIs, 95% prediction intervals, I², and "
         "Q-statistics; pre-registered subgroup contrasts (Region, "
         "Outcome Type, Era) are reported descriptively rather than "
         "as confirmatory tests."),

        # ---- DISCUSSION ----
        ("DISCUSSION", "", "", ""),
        ("Summary of evidence", "19",
         "Summarize the main results (including an overview of "
         "concepts, themes, and types of evidence available), link to "
         "the review questions and objectives, and consider the "
         "relevance to key groups.",
         "Discussion → \"Summary of Findings,\" \"Mapping-Prior "
         "Evaluation,\" \"Theoretical Implications,\" \"Tentative "
         "Practical Implications,\" and \"Distinguishing Robust from "
         "Fragile Findings\" subsections — the main mapped pattern "
         "(Conscientiousness as the only relatively robust trait; "
         "Agreeableness/Extraversion subgroup contrasts as exploratory; "
         "Openness/Neuroticism essentially null in the mappable pool) "
         "is summarised against MP1-MP5 and against the eight benchmark "
         "face-to-face meta-analyses."),
        ("Limitations", "20",
         "Discuss the limitations of the scoping review process.",
         "Discussion → \"Limitations\" subsection — seven explicit "
         "limitations: (1) small exploratory pool (k = 10, 6 direct + "
         "4 β-converted); (2) executed search did not cover the four "
         "pre-registered subscription-gated databases and is described "
         "as the basis for the scoping-review reframing; (3) single-"
         "reviewer workflow; (4) English-only restriction with "
         "particular consequence for Asian-subgroup interpretation; "
         "(5) outcome heterogeneity; (6) author's own prior primary "
         "study inclusion handled by pre-specified sensitivity "
         "analysis; (7) Big-Five-dimensional rather than facet-level "
         "granularity."),
        ("Conclusions", "21",
         "Provide a general interpretation of the results with respect "
         "to the review questions and objectives, as well as potential "
         "implications and/or next steps.",
         "Conclusion — the synthesis is positioned as a charted map "
         "and a research agenda rather than as an established "
         "quantitative effect-size estimate; next steps include "
         "subscription-gated-database replication search, linguistically "
         "inclusive corpus expansion, facet-level primary studies, and "
         "within-learner modality contrasts."),

        # ---- FUNDING ----
        ("FUNDING", "", "", ""),
        ("Funding", "22",
         "Describe sources of funding for the included sources of "
         "evidence, as well as sources of funding for the scoping "
         "review. Describe the role of the funders of the scoping "
         "review.",
         "Declarations → Funding subsection — no external funding "
         "was received for the scoping review; sources of funding "
         "for individual included primary studies are recorded in "
         "the data-extraction CSV where reported."),
    ]


def build_checklist(anonymous: bool, out_path: Path):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)
    s.page_width, s.page_height = s.page_height, s.page_width  # landscape

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("PRISMA-ScR Checklist (Tricco et al., 2018)")
    _set_run_font(run, size_pt=14, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    run = sub.add_run(
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Scoping Review with Exploratory "
        "Quantitative Synthesis"
    )
    _set_run_font(run, size_pt=11, bold=False)

    if anonymous:
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.paragraph_format.space_after = Pt(10)
        run = info.add_run(
            "[Author identifying information removed for double-anonymous "
            "peer review]"
        )
        _set_run_font(run, size_pt=10, bold=False)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run(
        "This checklist follows the PRISMA Extension for Scoping Reviews "
        "(PRISMA-ScR) reporting guidelines: Tricco, A. C., Lillie, E., "
        "Zarin, W., et al. (2018). PRISMA Extension for Scoping Reviews "
        "(PRISMA-ScR): Checklist and explanation. Annals of Internal "
        "Medicine, 169(7), 467–473. https://doi.org/10.7326/M18-0850. "
        "Items are numbered 1–22. The \"Location reported\" column points "
        "to the section/subsection of the main manuscript where each "
        "item is reported; page numbers are deliberately avoided because "
        "they shift between submission and journal-formatted PDFs."
    )
    _set_run_font(run, size_pt=10, bold=False)

    items = _build_items(anonymous)

    # Build the table with 4 columns: Section / # / Item / Location reported
    headers = ["Section", "#", "Checklist item", "Location reported"]
    table = doc.add_table(rows=len(items) + 1, cols=len(headers))
    table.autofit = False
    widths = [Cm(3.2), Cm(1.2), Cm(8.5), Cm(11.5)]

    # Header row
    for ci, htext in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.width = widths[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(htext)
        _set_run_font(run, size_pt=10, bold=True)
        _set_cell_shading(cell, "D9D9D9")
        _set_cell_borders(cell)

    for ri, (section, num, item, loc) in enumerate(items, start=1):
        is_section_header = (num == "" and item == "" and loc == "")
        cells = table.rows[ri].cells
        cells[0].text = ""
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
        for cc, value in enumerate([section, num, item, loc]):
            cell = cells[cc]
            cell.width = widths[cc]
            p = cell.paragraphs[0]
            run = p.add_run(value)
            _set_run_font(run, size_pt=9, bold=is_section_header)
            if is_section_header:
                _set_cell_shading(cell, "F2F2F2")
            _set_cell_borders(cell)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main():
    build_checklist(anonymous=False, out_path=OUT_NON_ANON)
    build_checklist(anonymous=True, out_path=OUT_ANON)


if __name__ == "__main__":
    main()
