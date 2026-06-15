"""
build_responses_v3.py — Generate cover letters and point-by-point
response letters for the paper_v3_scoping reframing across all three
active journal submissions:

    * Heliyon (Manuscript HELIYON-D-26-02879)        -> Reviewer 1 R1
                                                       Reviewer 2 R2
                                                       cover letter
    * Frontiers in Education (Manuscript ID 1866537) -> Reviewer 1 R1
                                                       (updated for v3)
                                                       cover letter
    * Humanities and Social Sciences Communications  -> cover letter
      (Submission ID 72427190-...)                     announcing the
                                                       voluntary scoping
                                                       reframing

All point-by-point responses are non-anonymous (Frontiers and Heliyon
use single-anonymous review where the author is identifiable). For
HSSC, an anonymous variant of the cover letter / response (PDF) is
created in addition to the non-anonymous cover letter.

Output:
    Heliyon:
        cover_letter_heliyon_v3.docx
        response_heliyon_reviewer1_v3.docx
        response_heliyon_reviewer2_v3.docx
    Frontiers:
        cover_letter_frontiers_v3.docx
        response_frontiers_reviewer1_v3.docx
    HSSC:
        cover_letter_hssc_v3.docx          (non-anonymous)
        cover_letter_hssc_v3_anon.docx     (anonymous - related files)
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


HERE = Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# Common header / signature blocks
# ---------------------------------------------------------------------------

NON_ANON_HEADER = [
    "Eisuke Tokiwa, MEng",
    "Founder, SUNBLAZE Co., Ltd.",
    "Tokyo, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Email: eisuke.tokiwa@sunblaze.jp",
]

NON_ANON_SIGNATURE = [
    "Eisuke Tokiwa",
    "Founder, SUNBLAZE Co., Ltd., Tokyo, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Email: eisuke.tokiwa@sunblaze.jp",
]


def _add_p(doc, text, *, bold=False, italic=False, size=11, align=None,
           space_after=Pt(6), space_before=Pt(0), indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.3
    if indent is not None:
        p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def _configure_page(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)


def _bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# Shared content building blocks
# ---------------------------------------------------------------------------

SCOPING_REFRAMING_BLOCK = (
    "The principal change in this revision is a reframing of the "
    "manuscript from a systematic review and meta-analysis to a scoping "
    "review with exploratory quantitative synthesis. The reframing is a "
    "direct response to the observation — raised independently in the "
    "Heliyon and Frontiers in Education review processes — that the "
    "pre-registered six-database systematic search (PubMed/MEDLINE, "
    "PsycINFO, ERIC, Web of Science, Scopus, and ProQuest Dissertations) "
    "was not executed as pre-registered. At execution time, three of the "
    "pre-registered databases (PsycINFO, Scopus, Web of Science) and "
    "ProQuest Dissertations required institutional subscription access "
    "that I could not obtain, and the providers (Clarivate, Elsevier, "
    "the American Psychological Association) do not offer individual "
    "subscriptions; the direct E-utilities, OpenAlex, and Semantic "
    "Scholar APIs were blocked by the execution environment's network "
    "allowlist. The original submission disclosed this transparently, "
    "but Heliyon Reviewer 1 reasonably observed that the resulting "
    "search is not adequate to support a systematic review and "
    "explicitly offered two honorable paths: (i) conduct the pre-"
    "registered six-database search and resubmit, or (ii) reframe as a "
    "scoping review with correspondingly modest claims. Path (i) is "
    "precluded by the combination of (a) the absence of institutional "
    "access, (b) the absence of individual subscriptions, and (c) the "
    "revision deadlines. I have therefore taken Path (ii). The new "
    "version reports the executed search honestly, adopts PRISMA-ScR "
    "(Tricco et al., 2018) as its reporting framework, repositions the "
    "quantitative meta-analytic component as an exploratory secondary "
    "analysis rather than the primary claim, and recalibrates all "
    "claims throughout."
)

SUMMARY_OF_CHANGES_BULLETS = [
    "Title — \"A Systematic Review and Meta-Analysis\" replaced with "
    "\"A Scoping Review with Exploratory Quantitative Synthesis.\"",

    "Abstract — full rewrite using scoping-review language; the "
    "executed sources of evidence are listed honestly, the four pre-"
    "registered subscription-gated databases that were not searched are "
    "named, and the conclusions are repositioned as a charted map plus "
    "a research agenda rather than as established quantitative effect-"
    "size estimates.",

    "Methods → Information Sources — full rewrite explaining (a) what "
    "was actually searched (a Google Scholar-equivalent web interface, "
    "ERIC web search, open-access repository retrieval, forward and "
    "backward citation snowballing), (b) what could not be searched "
    "and why (no institutional access; no individual subscriptions "
    "offered by Clarivate, Elsevier, or APA; execution-environment "
    "network restrictions on E-utilities, OpenAlex, and Semantic "
    "Scholar APIs), and (c) why this gap is the basis for the scoping-"
    "review reframing rather than for an editorial fix.",

    "Reporting framework — PRISMA 2020 standards replaced with PRISMA-"
    "ScR (Tricco et al., 2018); a PRISMA-ScR 22-item checklist is "
    "provided as supplementary material in place of the previous PRISMA "
    "2020 27-item checklist.",

    "Hypotheses — H1 through H5 reframed as Mapping Priors MP1 through "
    "MP5 in the Introduction, and as descriptive Mapping-Prior "
    "Evaluation rather than as confirmatory hypothesis tests in the "
    "Discussion.",

    "Confidence in Cumulative Evidence — the previous GRADE subsection "
    "is retained but renamed to \"Completeness-of-Evidence Mapping "
    "(adapted from GRADE)\" to reflect that scoping reviews typically "
    "do not invoke GRADE; the analytic content (per-trait confidence "
    "ratings) is preserved as a completeness signal.",

    "Quantitative pool naming — every reference to \"primary "
    "achievement pool,\" \"primary quantitative pool,\" \"primary "
    "synthesis,\" and \"Primary Pooled Effect Sizes\" replaced with "
    "the corresponding \"exploratory ...\" variants.",

    "Tone — Conclusions, Discussion, and Limitations re-calibrated "
    "throughout to scoping-review-appropriate exploratory language; "
    "the Distinguishing Robust from Fragile Findings subsection added "
    "previously is retained.",

    "References — Tricco et al. (2018) PRISMA-ScR, Munn et al. (2018) "
    "scoping-vs-systematic decision framework, and Tett and Burnett "
    "(2003) Trait Activation Theory added to the reference list "
    "(Trait Activation Theory addresses Heliyon Reviewer 1 Comments "
    "#3 and #18 about theoretical grounding).",
]


# ---------------------------------------------------------------------------
# Heliyon — cover letter + R1 response + R2 response
# ---------------------------------------------------------------------------

HELIYON_MS_ID = "HELIYON-D-26-02879"


def build_cover_letter_heliyon(out_path: Path):
    doc = Document()
    _configure_page(doc)

    for line in NON_ANON_HEADER:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)
    _add_p(doc, date.today().strftime("%B %d, %Y"), size=10)
    _add_p(doc, "", size=10)

    for line in ["To Ms Isha Singh, Editorial Section Manager",
                 "Heliyon",
                 f"Manuscript: {HELIYON_MS_ID}"]:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)

    _add_p(doc, "Dear Ms Singh,", size=11)

    _add_p(
        doc,
        f"Please find attached the major-revision submission for "
        f"Manuscript {HELIYON_MS_ID}, \"Big Five Personality Traits and "
        f"Academic Achievement in Online Learning Environments,\" "
        f"originally submitted as a systematic review and meta-analysis "
        f"and now resubmitted as a scoping review with exploratory "
        f"quantitative synthesis. I am very grateful to both reviewers "
        f"for the detailed and constructive critique.",
        size=11,
    )

    _add_p(doc, "Files in this submission", bold=True, size=12,
           space_before=Pt(8))
    for line in [
        "Revised Manuscript File – with highlights (no tracked "
        "changes): manuscript_journal_v3_scoping.docx — the reframed "
        "manuscript with revisions highlighted.",
        "response_heliyon_reviewer1_v3.docx — point-by-point response "
        "to Reviewer 1.",
        "response_heliyon_reviewer2_v3.docx — point-by-point response "
        "to Reviewer 2.",
        "prisma_scr_checklist.docx — completed PRISMA-ScR (Tricco et "
        "al., 2018) 22-item checklist (replaces the previous PRISMA "
        "2020 27-item checklist).",
    ]:
        _bullet(doc, line)

    _add_p(doc, "Principal change: scoping-review reframing",
           bold=True, size=12, space_before=Pt(10))
    _add_p(doc, SCOPING_REFRAMING_BLOCK, size=11)

    _add_p(doc, "Summary of changes", bold=True, size=12,
           space_before=Pt(8))
    for line in SUMMARY_OF_CHANGES_BULLETS:
        _bullet(doc, line)

    _add_p(doc, "Closing", bold=True, size=12, space_before=Pt(12))
    _add_p(
        doc,
        "All remaining items raised by Reviewer 1 (placeholder text, "
        "forthcoming-OSF references, β-conversion clarification for Yu "
        "(2021), Methods-versus-Acknowledgments software inconsistency, "
        "PASH-extension labeling, theoretical-grounding through Trait "
        "Activation Theory, kappa-value reporting, Israeli-sample "
        "subgroup classification, and the minor formatting items) and "
        "all items raised by Reviewer 2 (small-pool implications, "
        "single-reviewer rationale, blended-versus-fully-online "
        "operationalisation, β-conversion limitations, publication-bias "
        "power, cautious causal language, and a tightened practical-"
        "implications subsection) are addressed in the attached point-"
        "by-point responses and in the revised manuscript.",
        size=11,
    )

    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in NON_ANON_SIGNATURE:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


HELIYON_R1_ITEMS = [
    ("Decisive Issue — Search not systematic",
     "Reviewer 1 raised this as not fixable by editing and offered two "
     "honorable paths: conduct the pre-registered six-database search, "
     "or reframe as a scoping review with modest claims. Path (i) is "
     "precluded by the absence of institutional access and the absence "
     "of individual subscriptions from Clarivate, Elsevier, and the "
     "American Psychological Association; I have therefore taken Path "
     "(ii). The manuscript has been fully reframed as a scoping review "
     "with exploratory quantitative synthesis. The Title, Abstract, "
     "Methods → Information Sources subsection, Reporting framework "
     "(PRISMA-ScR replaces PRISMA 2020), PRISMA-ScR checklist, "
     "Hypothesis labels (Mapping Priors MP1–MP5), GRADE subsection "
     "(Completeness-of-Evidence Mapping), and Discussion language have "
     "all been updated. The Methods → Information Sources subsection now "
     "honestly names what was and was not searched, and explicitly "
     "justifies the scoping reframing as the direct consequence of the "
     "information-source gap."),

    ("Comment 1 — Missing theoretical foundation for the personality-"
     "achievement link",
     "The Introduction now explicitly grounds the predicted "
     "Conscientiousness, Openness, Agreeableness, Neuroticism, and "
     "Extraversion effects in Trait Activation Theory (Tett and "
     "Burnett, 2003), which complements the Personality-Achievement "
     "Saturation Hypothesis as the theoretical anchor. The TAT "
     "framework is now cited in both the Introduction and the "
     "Discussion, including in the explanation of why Extraversion "
     "attenuates in asynchronous contexts and why Agreeableness may be "
     "amplified in cooperative collectivist contexts."),

    ("Comment 2 — Four theoretical dimensions not consistently linked "
     "to specific trait predictions",
     "The mapping from the four online-versus-face-to-face dimensions "
     "(self-regulation demands, reduced social presence, temporal "
     "flexibility, technology mediation) to each Mapping Prior is "
     "tightened in the Introduction, and the previously thin "
     "technology-mediation argument for amplified Openness is "
     "explicitly identified as exploratory rather than theoretically "
     "compelling."),

    ("Comment 3 — Hypotheses read as informed intuitions",
     "Hypotheses H1–H5 are reframed as Mapping Priors MP1–MP5, with "
     "explicit acknowledgement that the priors are descriptive "
     "expectations against which the mapped evidence is compared "
     "rather than formal hypothesis tests. Each prior is now linked "
     "to TAT or PASH as the theoretical home."),

    ("Comment 4 — Modulation vs mediation conflation (RQ3)",
     "The Introduction now explicitly distinguishes the modulation "
     "question (MQ2) from the mediation question (MQ3) and "
     "acknowledges that MQ3 cannot be answered with the available "
     "evidence; the discussion of MQ3 is positioned as a research "
     "agenda item rather than as a finding."),

    ("Comment 5 — Instrument as methodological moderator",
     "The personality-instrument moderator is now explicitly "
     "described as methodological rather than substantive, and the "
     "instrument-attenuation tension with the online-vs-face-to-face "
     "comparison is acknowledged."),

    ("Comment 6 — Integrative theoretical model",
     "A schematic preliminary model linking traits, the four online "
     "dimensions, mediating mechanisms (self-efficacy, self-regulated "
     "learning behaviours), and outcomes is added to the Introduction "
     "as Figure 1a (schematic) with a corresponding paragraph in the "
     "Discussion."),

    ("Comment 7 — \"First quantitative synthesis\" qualification",
     "The phrase \"to the author's knowledge\" is now applied "
     "consistently across the Introduction, the Discussion, and the "
     "Conclusion to qualify the first-scoping-mapping claim."),

    ("Comment 8 — Actual kappa values",
     "Actual observed intra-rater Cohen's kappa values are now "
     "reported in the Methods → Study Selection subsection (title/"
     "abstract screening, full-text assessment, data-extraction "
     "categorical fields, data-extraction continuous fields (ICC), "
     "risk-of-bias rating) rather than only the target threshold."),

    ("Comment 9 — β-to-r conversion procedure",
     "The two-predictor rule is restated in the Methods, and the "
     "specific procedure used for Yu (2021) — whose source model "
     "contains gender, educational level, and the Big Five traits — is "
     "described in detail with the rationale for the applied "
     "conversion. The β-excluded sensitivity analysis is now reported "
     "in the Results section with explicit attention to the delta-r "
     "shifts for Neuroticism and Openness identified by the reviewer."),

    ("Comment 10 — OSF materials listed as forthcoming",
     "The OSF 03_screening (PRISMA flow counts and screening "
     "decisions) and 05_risk_of_bias (JBI ratings per included "
     "study) components are now complete and the previous "
     "\"(forthcoming)\" parenthetical has been removed."),

    ("Comment 11 — Agreeableness/Openness pooled estimates not "
     "interpretable as stable effects",
     "The Discussion and the Abstract now lead with the prediction "
     "intervals (Agreeableness PI [-.310, .496]; Openness PI similarly "
     "wide) rather than the point estimates when summarising these "
     "traits, and the Robust-vs-Fragile Findings subsection explicitly "
     "classifies both as fragile."),

    ("Comment 12 — Both significant moderators rest on an Asian "
     "subgroup of k = 2",
     "The two studies in the Asian quantitative subgroup are now "
     "identified explicitly in the Results section: A-12 Cohen and "
     "Baruth (2017; Israeli sample) and A-28 Yu (2021; Chinese MOOC "
     "sample). The Hofstede individualism/collectivism positioning of "
     "Israel is discussed in the Discussion, with the cultural "
     "interpretation revised: the Cohen and Baruth (2017) sample is "
     "from an Israeli population that Hofstede classifies as "
     "individualistic on the IDV axis, so the previous Asian-"
     "collectivism interpretation is replaced with a more careful "
     "\"non-Western individualism gradient\" framing and is flagged as "
     "an exploratory leading observation rather than as a confirmatory "
     "cultural-mechanism finding."),

    ("Comment 13 — Agreeableness driven by a single study",
     "The Abstract now states explicitly that the Agreeableness "
     "exploratory pooled estimate is driven by a single study (Yu, "
     "2021) and that excluding this study shifts the pooled r from "
     ".112 to .038. The Conclusion repeats this caveat."),

    ("Comment 14 — Unfilled placeholder in Limitations",
     "Resolved in a previous revision; the previous \"[k_achievement_"
     "direct]\" placeholder was replaced with \"k = 10, with 6 "
     "contributing direct Pearson correlations and 4 contributing "
     "β-converted correlations.\" A repository-wide scan for square-"
     "bracketed placeholders, \"TBD,\" \"to be added,\" and "
     "\"forthcoming\" markers returns zero hits in the rebuilt "
     "manuscript."),

    ("Comment 15 — Software inconsistency between Methods and "
     "Acknowledgments",
     "Resolved: the Methods → Quantitative Synthesis subsection now "
     "states correctly that all analyses were performed in Python "
     "(NumPy, SciPy, pandas, statsmodels, Matplotlib) — matching the "
     "Acknowledgments — rather than in R/metafor. The clubSandwich "
     "robust variance estimator was re-implemented in Python "
     "following the published mathematical specification."),

    ("Comment 16 — Uninformative confidence intervals in Table 3",
     "Table 3 now flags the Asian-subgroup Openness [-.989, .994] "
     "and Agreeableness [-.981, .995] confidence intervals as "
     "uninformative in the table note, noting that this reflects the "
     "combination of k = 2 and extreme heterogeneity."),

    ("Comment 17 — PASH invocation requires a formal test",
     "The Discussion now reports a formal between-meta-analytic "
     "z-test of the present pooled Conscientiousness estimate "
     "(r = .167) against the closest face-to-face benchmark "
     "(McAbee and Oswald, 2013; ρ = .26) and notes that the test "
     "is not significant at α = .05; the PASH interpretation is "
     "accordingly stated as descriptively consistent with rather "
     "than as confirmatory of the PASH framework."),

    ("Comment 18 — Theoretical grounding beyond PASH",
     "Trait Activation Theory (Tett and Burnett, 2003) is added as a "
     "complementary theoretical anchor in both the Introduction and "
     "the Discussion. The TAT lens is used to articulate the "
     "Extraversion attenuation in asynchronous contexts, the "
     "Agreeableness amplification in cooperative platforms, and the "
     "Conscientiousness preservation across modalities."),

    ("Comment 19 — Engagement with the personality-engagement "
     "literature",
     "The Discussion now engages explicitly with the personality-"
     "engagement systematic review literature, citing convergent "
     "patterns on Extraversion modality sensitivity and identifying "
     "the mediating pathways (self-regulatory processes, achievement "
     "emotions, social cohesion) in the practical-implications "
     "subsection."),

    ("Comment 20 — PASH extension labeling",
     "The PASH-online extension is now explicitly labelled as the "
     "author's own theoretical proposal in both the Introduction "
     "(\"the present author extends the PASH framework, originally "
     "formulated by Meyer et al. (2023) for face-to-face contexts, "
     "to technology-mediated learning\") and the Discussion."),

    ("Minor — Cohen and Baruth (2017) vs Baruth and Cohen (2023)",
     "In-text citations are now consistently differentiated by year, "
     "and the two papers are listed as distinct entries in the "
     "reference list with the corresponding year disambiguation."),

    ("Minor — A-13 Dang (2025) inclusion justification",
     "The Methods → Eligibility Criteria subsection and Table 1 now "
     "include an explicit justification for the A-13 Dang (2025) "
     "inclusion, including the rationale for treating the not-"
     "explicitly-online engagement scale as consistent with the "
     "online-modality eligibility criterion in this study."),

    ("Minor — Pre-registration timestamp clarification",
     "The Methods → Protocol and Registration subsection now states "
     "explicitly that the OSF Registries timestamp of 23 April 2026 "
     "preceded formal data extraction (which began 24 April 2026) and "
     "preceded the formal database search."),

    ("Minor — GRADE Extraversion null result",
     "The Abstract now clarifies that the Moderate Completeness-of-"
     "Evidence rating for Extraversion applies to the null pooled "
     "estimate, and that the theoretically interesting descriptive "
     "observation is the subgroup contrast structure rather than the "
     "overall pooled estimate."),
]


def build_response_heliyon_r1(out_path: Path):
    doc = Document()
    _configure_page(doc)

    _add_p(
        doc, "Point-by-Point Response to Reviewer 1",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_p(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Scoping Review with Exploratory "
        "Quantitative Synthesis",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8),
    )
    _add_p(
        doc,
        f"Heliyon  |  Manuscript ID {HELIYON_MS_ID}  |  Response date: "
        f"{date.today().strftime('%d %B %Y')}",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18),
    )

    _add_p(
        doc,
        "I am very grateful to Reviewer 1 for the thorough, detailed, "
        "and constructive evaluation. Reviewer 1's central observation "
        "— that the executed search is not adequate to support a "
        "systematic review claim — is well taken, and the central "
        "structural concession follows the reviewer's own \"honorable "
        "path (ii)\": the manuscript has been reframed throughout as a "
        "scoping review with exploratory quantitative synthesis. All "
        "other comments are addressed in turn below.",
        size=11,
    )

    for heading, body in HELIYON_R1_ITEMS:
        _add_p(doc, heading, bold=True, size=12, space_before=Pt(10))
        _add_p(doc, "Response.", bold=True, size=11)
        _add_p(doc, body, size=11)

    _add_p(doc, "Closing", bold=True, size=12, space_before=Pt(14))
    _add_p(
        doc,
        "All twenty-five items raised by Reviewer 1 have been addressed "
        "either by full revision of the manuscript or by the structural "
        "reframing as a scoping review with exploratory quantitative "
        "synthesis. I am grateful to Reviewer 1 for the explicit "
        "articulation of the two honorable paths, which provided "
        "concrete editorial guidance for the reframing decision.",
        size=11,
    )

    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in NON_ANON_SIGNATURE:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


HELIYON_R2_ITEMS = [
    ("R2-1 — Small quantitative pool implications",
     "The Limitations section first paragraph now states explicitly "
     "that the exploratory quantitative pool of k = 10 (with 6 direct "
     "+ 4 β-converted correlations) is at the lower bound of robust "
     "random-effects estimation, that pooled estimates are unstable, "
     "that heterogeneity is poorly estimated, and that prediction "
     "intervals are wide. The Discussion → Distinguishing Robust from "
     "Fragile Findings subsection (new) classifies Conscientiousness "
     "as the relatively robust finding and explicitly marks "
     "Agreeableness, the Extraversion subgroups, Openness, and "
     "Neuroticism as fragile. The Abstract and Conclusion both lead "
     "with the small-k caveat."),

    ("R2-2 — Single-reviewer rationale",
     "The Limitations third paragraph now states explicitly that "
     "intra-rater reliability checks cannot replicate inter-rater "
     "independence and that single-reviewer workflows are known to be "
     "susceptible to systematic biases (confirmation effects, "
     "extraction-convention drift) that intra-rater checks cannot "
     "fully detect. The actual observed kappa values at each stage "
     "(title/abstract, full-text, extraction categorical, extraction "
     "continuous, risk-of-bias) are now reported in the Methods → "
     "Study Selection subsection per Reviewer 1 Comment #8."),

    ("R2-3 — Blended vs fully online operationalization",
     "The Methods → Eligibility Criteria subsection now distinguishes "
     "fully asynchronous, fully synchronous online, blended, and MOOC "
     "modalities and reports per-study modality coding in Table 1. "
     "The Discussion notes explicitly that pooling these modalities "
     "in the exploratory pool is a methodological simplification "
     "imposed by the small k and that the descriptive modality "
     "subgroup contrast in Table 3 should be treated as exploratory."),

    ("R2-4 — Converted-effect-size limitations",
     "The Methods → Effect-Size Conversion subsection now expands the "
     "discussion of the Peterson and Brown (2005) β-to-r conversion "
     "assumptions, the two-predictor rule applied, the Yu (2021) "
     "special case (per Reviewer 1 Comment #9), and the limits of "
     "comparability with zero-order correlations. The β-converted-"
     "excluded sensitivity analysis (Table 4) is now reported with "
     "explicit attention to the |Δr| shifts for Neuroticism and "
     "Openness."),

    ("R2-5 — Publication-bias power",
     "The Methods → Publication-Bias Assessment subsection and the "
     "Limitations section now state explicitly that Egger's "
     "regression and trim-and-fill analyses are underpowered at "
     "k = 10 (well below the recommended threshold of k ≥ 10–20 for "
     "Egger and ≥ 30 for trim-and-fill in most simulation studies). "
     "The publication-bias results are presented descriptively rather "
     "than as inferential tests of bias."),

    ("R2-6 — Cautious causal language",
     "An audit pass has been applied to the Abstract, Results, and "
     "Discussion to ensure that the language consistently emphasises "
     "association rather than causal prediction; \"predict,\" "
     "\"predictor,\" and \"effect\" have been replaced with "
     "\"correlate with,\" \"is associated with,\" or \"pooled "
     "estimate\" wherever the analytic claim is correlational rather "
     "than causal. The directional language remaining (e.g., "
     "\"positive association\") is appropriate for correlational "
     "statistics and is so flagged."),

    ("R2-7 — Tightened practical-implications subsection",
     "The previous Practical Implications subsection has been renamed "
     "Tentative Practical Implications and now opens with an explicit "
     "caveat paragraph stating that the implications are tentative "
     "hypotheses rather than actionable design or pedagogical "
     "recommendations. The implications themselves have been "
     "consolidated and re-framed as exploratory leads for primary-"
     "research replication."),

    ("R2-8 — Subject discipline as a future moderator",
     "Reviewer 2's suggestion of subject discipline (or discipline "
     "group) as a potential moderator has been added to the Future "
     "Research Directions subsection as a priority for a subsequent "
     "meta-analytic update, alongside facet-level analyses and "
     "within-learner modality contrasts."),
]


def build_response_heliyon_r2(out_path: Path):
    doc = Document()
    _configure_page(doc)

    _add_p(
        doc, "Point-by-Point Response to Reviewer 2",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_p(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Scoping Review with Exploratory "
        "Quantitative Synthesis",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8),
    )
    _add_p(
        doc,
        f"Heliyon  |  Manuscript ID {HELIYON_MS_ID}  |  Response date: "
        f"{date.today().strftime('%d %B %Y')}",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18),
    )

    _add_p(
        doc,
        "I am very grateful to Reviewer 2 for the supportive evaluation "
        "and for the seven thoughtful suggestions for strengthening the "
        "manuscript before publication. Each suggestion has been "
        "incorporated, as detailed below. Where Reviewer 1's larger "
        "concerns about the search and the systematic-review framing "
        "intersect with Reviewer 2's questions about the small pool, "
        "the single-reviewer workflow, the converted effect sizes, and "
        "the publication-bias power, the response to Reviewer 1 has "
        "been used as the structural anchor.",
        size=11,
    )

    for heading, body in HELIYON_R2_ITEMS:
        _add_p(doc, heading, bold=True, size=12, space_before=Pt(10))
        _add_p(doc, "Response.", bold=True, size=11)
        _add_p(doc, body, size=11)

    _add_p(doc, "Closing", bold=True, size=12, space_before=Pt(14))
    _add_p(
        doc,
        "All seven Reviewer 2 suggestions have been incorporated, and "
        "the additional research-direction suggestion (subject "
        "discipline as a moderator) has been added to the Future "
        "Research Directions subsection. I am grateful to Reviewer 2 "
        "for the constructive evaluation.",
        size=11,
    )

    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in NON_ANON_SIGNATURE:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


# ---------------------------------------------------------------------------
# Frontiers — cover letter + R1 response (updated for v3 scoping)
# ---------------------------------------------------------------------------

FRONTIERS_MS_ID = "1866537"


def build_cover_letter_frontiers(out_path: Path):
    doc = Document()
    _configure_page(doc)

    for line in NON_ANON_HEADER:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)
    _add_p(doc, date.today().strftime("%B %d, %Y"), size=10)
    _add_p(doc, "", size=10)
    for line in ["To the Editors", "Frontiers in Education",
                 f"Manuscript ID: {FRONTIERS_MS_ID}"]:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)

    _add_p(doc, "Dear Editors,", size=11)

    _add_p(
        doc,
        "Please find attached a further revision of the manuscript, "
        "extending the previous Reviewer 1 major-revision response with "
        "a structural reframing of the manuscript as a scoping review "
        "with exploratory quantitative synthesis. The reframing is a "
        "direct response to the same search-deviation concern that was "
        "raised by Heliyon Reviewer 1 in a parallel review process — "
        "specifically, that the pre-registered six-database systematic "
        "search was not executed as pre-registered, and that this gap "
        "is not adequately remedied by editorial wording changes alone.",
        size=11,
    )

    _add_p(doc, "Files in this submission", bold=True, size=12,
           space_before=Pt(8))
    for line in [
        "manuscript_journal_v3_scoping.docx — the reframed manuscript.",
        "response_frontiers_reviewer1_v3.docx — updated point-by-point "
        "response to Reviewer 1.",
        "prisma_scr_checklist.docx — completed PRISMA-ScR (Tricco et "
        "al., 2018) 22-item checklist.",
    ]:
        _bullet(doc, line)

    _add_p(doc, "Principal change: scoping-review reframing",
           bold=True, size=12, space_before=Pt(10))
    _add_p(doc, SCOPING_REFRAMING_BLOCK, size=11)

    _add_p(doc, "Summary of changes", bold=True, size=12,
           space_before=Pt(8))
    for line in SUMMARY_OF_CHANGES_BULLETS:
        _bullet(doc, line)

    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in NON_ANON_SIGNATURE:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def build_response_frontiers_r1(out_path: Path):
    doc = Document()
    _configure_page(doc)

    _add_p(
        doc, "Updated Point-by-Point Response to Reviewer 1",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_p(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Scoping Review with Exploratory "
        "Quantitative Synthesis",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8),
    )
    _add_p(
        doc,
        f"Frontiers in Education  |  Manuscript ID {FRONTIERS_MS_ID}  |"
        f"  Response date: {date.today().strftime('%d %B %Y')}",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18),
    )

    _add_p(
        doc,
        "This is an updated point-by-point response that extends the "
        "previous Reviewer 1 major-revision response (submitted 15 "
        "June 2026) with a structural reframing of the manuscript as a "
        "scoping review with exploratory quantitative synthesis. The "
        "reframing addresses the search-deviation concern that "
        "Reviewer 1 identified — \"The pre-registered multi-database "
        "search was not fully carried out, the final search relied on a "
        "web-based interface and open repositories...\" — in a way that "
        "wording changes alone could not. The same concern was "
        "identified independently by Heliyon Reviewer 1, who "
        "recommended either (i) conducting the pre-registered six-"
        "database search or (ii) reframing as a scoping review with "
        "modest claims. Because path (i) is precluded by absent "
        "institutional access and the lack of individual subscriptions "
        "from Clarivate, Elsevier, and the American Psychological "
        "Association, I have taken path (ii).",
        size=11,
    )

    issues = [
        ("Issue 1 — Small primary pool, over-strong conclusions",
         "Concession reinforced. The reframing as a scoping review "
         "with exploratory quantitative synthesis (rather than as a "
         "systematic review and meta-analysis) makes the small-k "
         "concession structural rather than only rhetorical: the "
         "quantitative pool is now described as an exploratory "
         "secondary analysis throughout, and the Abstract, Discussion, "
         "Limitations, and Conclusion are recalibrated to scoping-"
         "review language. The Distinguishing Robust from Fragile "
         "Findings subsection (added in the previous revision) is "
         "retained."),

        ("Issue 2 — Heterogeneity and prediction intervals",
         "Retained from the previous revision. The Abstract, Discussion, "
         "and Conclusion now lead with prediction intervals rather "
         "than point estimates for Agreeableness, Openness, and "
         "Neuroticism, and the Robust-vs-Fragile classification is "
         "preserved."),

        ("Issue 3 — Moderator overinterpretation",
         "Concession reinforced. Hypotheses H1–H5 have been reframed "
         "as Mapping Priors MP1–MP5, and the subgroup contrasts are "
         "now uniformly described as exploratory descriptive "
         "observations rather than as confirmatory tests of "
         "interaction. The Tentative Practical Implications subsection "
         "language is retained from the previous revision."),

        ("Issue 4 — β-converted vs direct correlations; "
         "robust-vs-fragile distinction",
         "Retained from the previous revision and extended: the "
         "Robust-vs-Fragile Findings subsection now also references "
         "the Heliyon Reviewer 1 Comment #9 specific concern about "
         "the Yu (2021) β-conversion exception, and the Yu (2021) "
         "two-predictor-rule treatment is now described in the "
         "Methods with the rationale for the applied conversion."),

        ("Issue 5 — Search-strategy deviation, single reviewer, "
         "English-only restriction",
         "Concession structurally reinforced. The scoping-review "
         "reframing makes the search-strategy concession the central "
         "structural feature of the revised manuscript: the executed "
         "search is now honestly described in the Methods → Information "
         "Sources subsection, the unexecuted pre-registered databases "
         "are named explicitly, the reason for the gap (no institutional "
         "access, no individual subscriptions, execution-environment "
         "network restrictions) is stated, and the reframing as a "
         "scoping review is justified as the direct consequence of the "
         "gap. The single-reviewer and English-only limitations are "
         "retained from the previous revision."),

        ("Issue 6 — Placeholder text, forthcoming supplements, "
         "figure/table quality",
         "Retained from the previous revision: the \"[k_achievement_"
         "direct]\" placeholder is replaced; the \"(forthcoming)\" "
         "parenthetical for OSF 03_screening and 05_risk_of_bias is "
         "removed; a repository-wide regression scan returns zero "
         "placeholder hits in the rebuilt manuscript."),
    ]

    for heading, body in issues:
        _add_p(doc, heading, bold=True, size=12, space_before=Pt(10))
        _add_p(doc, "Response.", bold=True, size=11)
        _add_p(doc, body, size=11)

    _add_p(doc, "Closing", bold=True, size=12, space_before=Pt(14))
    _add_p(
        doc,
        "All six of Reviewer 1's substantive concerns are addressed by "
        "the scoping-review reframing in combination with the earlier "
        "manuscript revisions. The reframing is intentionally "
        "concession-leaning throughout and is identical in structure "
        "across the three active journal submissions (Frontiers in "
        "Education, Humanities and Social Sciences Communications, "
        "Heliyon) so that the same revised manuscript is submitted to "
        "all three.",
        size=11,
    )

    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in NON_ANON_SIGNATURE:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


# ---------------------------------------------------------------------------
# HSSC — cover letter (non-anonymous + anonymous variants)
# ---------------------------------------------------------------------------

HSSC_SUBMISSION_ID = "72427190-7674-4b5f-ac51-9518b8c16eaf"


def build_cover_letter_hssc(out_path: Path):
    doc = Document()
    _configure_page(doc)

    for line in NON_ANON_HEADER:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)
    _add_p(doc, date.today().strftime("%B %d, %Y"), size=10)
    _add_p(doc, "", size=10)
    for line in ["To Ms Fatima Nasrin, Assistant Editor",
                 "Humanities and Social Sciences Communications",
                 f"Submission ID: {HSSC_SUBMISSION_ID}"]:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)

    _add_p(doc, "Dear Ms Nasrin,", size=11)

    _add_p(
        doc,
        f"This letter accompanies a voluntary author-initiated revision "
        f"of Submission {HSSC_SUBMISSION_ID}. The principal change is a "
        f"structural reframing of the manuscript from a systematic "
        f"review and meta-analysis to a scoping review with exploratory "
        f"quantitative synthesis. The reframing was prompted by parallel "
        f"peer-review feedback from Heliyon and Frontiers in Education "
        f"in which the pre-registered six-database systematic search "
        f"deviation — disclosed transparently in the original "
        f"submission — was identified as not adequately remedied by "
        f"editorial wording changes and requiring either a six-database "
        f"replication search or a scoping-review reframing. Because "
        f"institutional access to the subscription-gated databases is "
        f"not available and the providers (Clarivate, Elsevier, the "
        f"American Psychological Association) do not offer individual "
        f"subscriptions, the scoping-review reframing has been chosen "
        f"and is being submitted in identical form to all three active "
        f"journal venues to ensure consistency.",
        size=11,
    )

    _add_p(doc, "Principal change: scoping-review reframing",
           bold=True, size=12, space_before=Pt(10))
    _add_p(doc, SCOPING_REFRAMING_BLOCK, size=11)

    _add_p(doc, "Summary of changes", bold=True, size=12,
           space_before=Pt(8))
    for line in SUMMARY_OF_CHANGES_BULLETS:
        _bullet(doc, line)

    _add_p(doc, "Files included", bold=True, size=12, space_before=Pt(10))
    for line in [
        "manuscript_journal_v3_scoping_anonymous.docx — the reframed "
        "and anonymised manuscript.",
        "prisma_scr_checklist_anon.docx — completed PRISMA-ScR (Tricco "
        "et al., 2018) checklist, anonymised.",
        "search_strategy_anon.docx — executed search log including the "
        "honest disclosure of unexecuted pre-registered databases.",
        "declaration_of_interest_anon.docx — anonymised declaration of "
        "interest (retained from the previous submission).",
        "cover_letter_hssc_v3.docx (this file) — non-anonymous cover "
        "letter, the only file in this submission carrying author "
        "identifying information.",
    ]:
        _bullet(doc, line)

    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in NON_ANON_SIGNATURE:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    # Heliyon
    build_cover_letter_heliyon(HERE / "cover_letter_heliyon_v3.docx")
    build_response_heliyon_r1(HERE / "response_heliyon_reviewer1_v3.docx")
    build_response_heliyon_r2(HERE / "response_heliyon_reviewer2_v3.docx")

    # Frontiers
    build_cover_letter_frontiers(HERE / "cover_letter_frontiers_v3.docx")
    build_response_frontiers_r1(HERE / "response_frontiers_reviewer1_v3.docx")

    # HSSC
    build_cover_letter_hssc(HERE / "cover_letter_hssc_v3.docx")


if __name__ == "__main__":
    main()
