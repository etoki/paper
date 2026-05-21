"""
build_supplements_anonymous.py — Generate double-anonymous-ready
supplementary documents for Springer HSSC technical-check revision.

Springer HSSC's minor-revision request (Submission ID 72427190-7674-
4b5f-ac51-9518b8c16eaf, 2026-05-26 deadline) requires:

  * "KINDLY DO NOT INCLUDE THE AUTHORS NAME OR DETAILS IN ... ANY OF
    THE FILES, EXCEPT OF THE COVER LETTER."

The non-anonymous PRISMA checklist and Declaration of Interest produced
by build_supplements.py both leak identifying info (author name,
affiliation, ORCID, email, OSF DOIs, GitHub URL). This script rebuilds
both as anonymised companion files for the HSSC submission.

Outputs:
  * prisma_2020_checklist_anon.docx
  * declaration_of_interest_anon.docx  (Springer wording: subject-only
    COI text, no identifying author block)

Strategy mirrors build_docx_anonymous.py: build via the canonical
supplement builders, then walk every paragraph (in body and tables)
and redact identifying tokens using the same REPLACEMENTS table.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_supplements
from build_docx_anonymous import REPLACEMENTS, rewrite_paragraph


def redact_doc(doc):
    n = 0
    for p in doc.paragraphs:
        if rewrite_paragraph(p, REPLACEMENTS):
            n += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if rewrite_paragraph(p, REPLACEMENTS):
                        n += 1
    return n


def _strip_author_block_from_doi(doc):
    """The Declaration of Interest opens with an "Author:" / "Affiliation:" /
    "ORCID:" / "Email:" identifying block. After token-level redaction those
    lines become "Author: [Author]" / "Affiliation: [Affiliation]" etc., which
    still betray the document structure. Replace the whole block with a
    single anonymised placeholder line."""
    target_prefixes = ("Author:", "Affiliation:", "ORCID:", "Email:")
    replaced = False
    for p in doc.paragraphs:
        text = "".join(r.text for r in p.runs).strip()
        if text.startswith(target_prefixes):
            for r in p.runs:
                r.text = ""
            if not replaced:
                if p.runs:
                    p.runs[0].text = (
                        "[Author identifying block removed for double-"
                        "anonymous peer review.]"
                    )
                else:
                    p.add_run(
                        "[Author identifying block removed for double-"
                        "anonymous peer review.]"
                    )
                replaced = True


def build_anonymous_prisma(out_path: Path):
    """Build the canonical PRISMA checklist, then redact."""
    build_supplements.build_prisma_checklist()
    src = build_supplements.PRISMA_OUT
    doc = Document(str(src))
    n = redact_doc(doc)
    doc.save(str(out_path))
    print(f"Wrote {out_path}  (paragraphs redacted: {n})")


EXTRA_DOI_REPLACEMENTS = [
    # Standalone "Eisuke" mentions in Author Contributions and signature
    # blocks that REPLACEMENTS misses (it only patterns "Eisuke Tokiwa").
    ("Eisuke [author's own study]", "The author"),
    ("Eisuke [Author]", "The author"),
    ("Eisuke", "The author"),
]


def build_anonymous_declaration_of_interest(out_path: Path):
    build_supplements.build_declaration_of_interest()
    src = build_supplements.DOI_OUT
    doc = Document(str(src))
    _strip_author_block_from_doi(doc)
    n = redact_doc(doc)
    # Second pass for declaration-specific residues.
    for p in doc.paragraphs:
        rewrite_paragraph(p, EXTRA_DOI_REPLACEMENTS)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rewrite_paragraph(p, EXTRA_DOI_REPLACEMENTS)
    doc.save(str(out_path))
    print(f"Wrote {out_path}  (paragraphs redacted: {n})")


def main():
    here = Path(__file__).resolve().parent
    build_anonymous_prisma(here / "prisma_2020_checklist_anon.docx")
    build_anonymous_declaration_of_interest(
        here / "declaration_of_interest_anon.docx"
    )


if __name__ == "__main__":
    main()
