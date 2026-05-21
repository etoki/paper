"""
build_docx_anonymous.py — Generate a double-anonymous-review-ready
version of manuscript_journal_v2.docx for journals such as Springer
Nature's Humanities and Social Sciences Communications (HSSC) that
require all author-identifying information to be removed from the
manuscript file at submission.

Strategy:
  1. Build the manuscript using build_docx.py's existing builders, but
     swap the APA-7 title page for a stripped placeholder title page
     so that the author name, ORCID, affiliation, and correspondence
     details never enter the document.
  2. Post-process every paragraph (in body text and tables) to replace
     identifying tokens — author name in self-citations, OSF/Preprint/
     GitHub URLs, contributor initials — with anonymised placeholders.
  3. Replace the Tokiwa (2025) reference list entry with an anonymised
     placeholder so the Reference list does not betray authorship.

The output is meant to be uploaded as the "anonymised manuscript" file
in submission systems that use double-anonymous peer review. The
non-anonymised manuscript_journal_v2.docx remains the canonical version
for Frontiers, MDPI, BMC, and Heliyon (which use single-anonymous or
open peer review).
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_docx


# ---------------------------------------------------------------------------
# Anonymised title page (replaces build_docx.build_title_page)
# ---------------------------------------------------------------------------

def build_anonymous_title_page(doc):
    """Title page with all identifying info redacted."""
    for _ in range(4):
        build_docx.add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    build_docx.add_para(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online Learning "
        "Environments: A Systematic Review and Meta-Analysis",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    build_docx.add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    build_docx.add_para(
        doc,
        "[Author identifying information removed for double-anonymous peer review]",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    build_docx.add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    build_docx.add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    build_docx.add_para(doc, "Author Note", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    build_docx.add_para(
        doc,
        "Author identifying details — name, ORCID, affiliation, and "
        "correspondence information — have been removed for double-anonymous "
        "peer review and will be reinstated upon acceptance. The research "
        "received no external funding. Conflict-of-interest, pre-registration, "
        "and open-science deposit information are provided in the Declarations "
        "section using anonymised wording; identifying URLs (OSF, preprint, "
        "code repository) will be supplied to the editorial office post-"
        "acceptance.",
        indent_first=True,
    )


# ---------------------------------------------------------------------------
# Token-level redaction
# ---------------------------------------------------------------------------

# Order matters: longer / more specific patterns first to avoid partial overlap.
REPLACEMENTS = [
    # OSF / preprint / code-repository URLs (complete URLs first)
    ("https://doi.org/10.17605/OSF.IO/E5W47",
     "[OSF preregistration DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/79M5J",
     "[OSF project DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/7FRGH",
     "[OSF protocol DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/UVJDY",
     "[OSF search DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/YH28G",
     "[OSF screening DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/5UW98",
     "[OSF extraction DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/3XYNE",
     "[OSF risk-of-bias DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/Q4XKB",
     "[OSF analysis DOI — provided post-acceptance]"),
    ("https://doi.org/10.17605/OSF.IO/XZ6PT",
     "[OSF index DOI — provided post-acceptance]"),
    ("https://doi.org/10.21203/rs.3.rs-9513298",
     "[Preprint DOI — provided post-acceptance]"),
    ("https://github.com/etoki/paper",
     "[Code repository — provided post-acceptance]"),
    ("https://osf.io/e5w47/",
     "[OSF preregistration link — provided post-acceptance]"),
    ("10.17605/OSF.IO/E5W47",
     "[OSF preregistration DOI — provided post-acceptance]"),
    ("10.17605/OSF.IO/79M5J",
     "[OSF project DOI — provided post-acceptance]"),
    ("10.17605/OSF.IO/7FRGH",
     "[OSF protocol DOI]"),
    ("10.17605/OSF.IO/UVJDY",
     "[OSF search DOI]"),
    ("10.17605/OSF.IO/YH28G",
     "[OSF screening DOI]"),
    ("10.17605/OSF.IO/5UW98",
     "[OSF extraction DOI]"),
    ("10.17605/OSF.IO/3XYNE",
     "[OSF risk-of-bias DOI]"),
    ("10.17605/OSF.IO/Q4XKB",
     "[OSF analysis DOI]"),
    ("10.17605/OSF.IO/XZ6PT",
     "[OSF index DOI]"),
    ("https://orcid.org/0009-0009-7124-6669",
     "[ORCID — provided post-acceptance]"),

    # Self-citation patterns (multi-token before single-token)
    ("A-25 Tokiwa, 2025", "A-25 [author's own study, 2025]"),
    ("A-25 Tokiwa", "A-25 [author's own study]"),
    ("(Tokiwa, 2025)", "(the author's own prior primary study)"),
    ("Tokiwa (2025)", "the author's own prior primary study"),
    ("Tokiwa, 2025", "the author's own prior primary study, 2025"),
    ("Tokiwa is in the", "the author's own study is in the"),
    ("Tokiwa", "[author's own study]"),

    # Personal identifiers (after self-citation patterns are handled)
    ("Eisuke Tokiwa", "[Author]"),
    ("eisuke.tokiwa@sunblaze.jp", "[Email]"),
    ("0009-0009-7124-6669", "[ORCID]"),
    ("SUNBLAZE Co., Ltd.", "[Affiliation]"),
    ("SUNBLAZE", "[Affiliation]"),

    # Authors' contributions initials
    ("(ET)", "(the author)"),
]


def rewrite_paragraph(p, replacements):
    full = "".join(r.text for r in p.runs)
    new = full
    for old, repl in replacements:
        new = new.replace(old, repl)
    if new == full:
        return False
    first_run = p.runs[0] if p.runs else None
    for r in p.runs[1:]:
        r.text = ""
    if first_run is not None:
        first_run.text = new
    else:
        p.add_run(new)
    return True


def redact_doc(doc, replacements):
    n_changed = 0
    for p in doc.paragraphs:
        if rewrite_paragraph(p, replacements):
            n_changed += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if rewrite_paragraph(p, replacements):
                        n_changed += 1
    return n_changed


def replace_data_availability_section(doc):
    """Rewrite the paragraph that follows the 'Data Availability Statement'
    heading with the Springer HSSC-compliant anonymised version (lists what
    has been shared as anonymised supplementary files; explains that the
    identifying URLs are withheld only for double-anonymous peer review)."""
    new_text = (
        "All materials underpinning this systematic review and meta-analysis "
        "have been made available to the editorial office and referees as "
        "anonymised supplementary files with this submission. Specifically: "
        "(a) the completed PRISMA 2020 27-item checklist with section-level "
        "locations (prisma_2020_checklist_anon.docx); (b) the full Boolean "
        "search strategy, database and source list, query-level hit counts, "
        "and the documented deviation from the pre-registered search plan "
        "(search_strategy_anon.docx); (c) the PRISMA 2020 flow diagram "
        "(reproduced as Figure 1 of the main manuscript); (d) the eligibility "
        "criteria, screening decisions, and full-text exclusion reasons "
        "(reported in the Methods and Results sections of the manuscript); "
        "(e) the data-extraction table (reproduced as Table 1 of the "
        "manuscript) listing every included study's country, modality, "
        "sample size, personality instrument, outcome operationalisation, "
        "era, region, JBI risk-of-bias score, and inclusion status; (f) the "
        "synthesis tables (Tables 2-5 of the manuscript: pooled effects, "
        "moderator analyses, sensitivity analyses, and adapted GRADE "
        "ratings); and (g) the anonymised Declaration of Interest "
        "(declaration_of_interest_anon.docx). Beyond the materials uploaded "
        "here, the full pre-registration record on a public open-science "
        "registry, the seven-component open-science project (protocol, "
        "search log, screening ledger, data-extraction CSV, JBI risk-of-bias "
        "spreadsheet, analysis code and outputs, and article-level DOI "
        "index), and the version-controlled code repository hosting the "
        "full Python analysis pipeline have all been permanently deposited "
        "under stable DOIs and are publicly accessible. Those DOIs are "
        "withheld at this stage solely because their URLs identify the "
        "author and will be supplied to the editorial office post-"
        "acceptance for inclusion in the final published version. Nothing "
        "has been redacted from these deposits; the only restriction is "
        "the temporary withholding of identifying URLs for double-anonymous "
        "review."
    )
    found = False
    for p in doc.paragraphs:
        text = "".join(r.text for r in p.runs).strip()
        if text == "Data Availability Statement":
            found = True
            continue
        if found and text:
            first = p.runs[0] if p.runs else None
            for r in p.runs[1:]:
                r.text = ""
            if first is not None:
                first.text = new_text
            else:
                p.add_run(new_text)
            return True
    return False


def replace_tokiwa_reference(doc):
    """Replace the Tokiwa (2025) reference list entry with an anonymised
    placeholder so the Reference list does not betray authorship."""
    n = 0
    for p in doc.paragraphs:
        text = "".join(r.text for r in p.runs)
        if "[author's own study], E. (2025)" in text or text.lstrip().startswith("[author's own study], E."):
            for r in p.runs:
                r.text = ""
            placeholder = (
                "[Author's own primary study (2025) — full reference removed for "
                "double-anonymous review; details available to the editorial office "
                "on request.]"
            )
            if p.runs:
                p.runs[0].text = placeholder
            else:
                p.add_run(placeholder)
            n += 1
    return n


# ---------------------------------------------------------------------------
# Anonymous build entry
# ---------------------------------------------------------------------------

def build_anonymous_manuscript(output_path):
    doc = Document()
    build_docx.configure_page(doc)
    build_docx.configure_styles(doc)

    build_anonymous_title_page(doc)
    doc.add_page_break()
    build_docx.build_declarations(doc, for_journal=True)
    doc.add_page_break()

    build_docx.build_abstract(doc)
    build_docx.build_intro_part1(doc)
    build_docx.build_intro_part2(doc)
    build_docx.build_intro_part3(doc)
    build_docx.build_methods_part1(doc)
    build_docx.build_methods_part2(doc)
    build_docx.build_methods_part3(doc)
    build_docx.build_methods_part4(doc)
    build_docx.build_methods_part5(doc)
    build_docx.build_results_part1(doc)
    build_docx.build_results_part2(doc)
    build_docx.build_results_part3(doc)
    build_docx.build_discussion_part1(doc)
    build_docx.build_discussion_part2(doc)
    build_docx.build_discussion_part3(doc)
    build_docx.build_conclusion(doc)
    build_docx.build_references(doc)

    build_docx.build_table1_characteristics(doc)
    build_docx.build_table2_pooled(doc)
    build_docx.build_table3_moderators(doc)
    build_docx.build_table4_sensitivity(doc)
    build_docx.build_table5_grade(doc)
    build_docx.build_figure1_prisma(doc)
    build_docx.build_forest_plots(doc)
    build_docx.build_funnel_plots(doc)

    n_changed = redact_doc(doc, REPLACEMENTS)
    n_refs = replace_tokiwa_reference(doc)
    da_replaced = replace_data_availability_section(doc)

    build_docx.italicize_stats_in_doc(doc)

    doc.save(output_path)
    print(f"Wrote {output_path}")
    print(f"  paragraphs redacted: {n_changed}")
    print(f"  Tokiwa reference entries replaced: {n_refs}")
    print(f"  Data Availability section rewritten: {da_replaced}")


def main():
    here = Path(__file__).resolve().parent
    out = here / "manuscript_journal_v2_anonymous.docx"
    build_anonymous_manuscript(str(out))


if __name__ == "__main__":
    main()
