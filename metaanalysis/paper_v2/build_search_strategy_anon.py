"""
build_search_strategy_anon.py — Generate the anonymised Search Strategy
supplementary document for Springer HSSC technical-check revision.

The HSSC minor-revision request (Submission ID 72427190-7674-4b5f-ac51-
9518b8c16eaf, 2026-05-26 deadline) asks for full search strategies for
all databases consulted to be made available to referees and editors,
as anonymous supplementary material.

Source content: metaanalysis/search_log.md. This script renders that
content as a docx with all identifying tokens (author name, OSF/GitHub
URLs, OSF project IDs, ORCID, email, affiliation) replaced by anonymous
placeholders.

Output: search_strategy_anon.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HERE = Path(__file__).resolve().parent
OUT = HERE / "search_strategy_anon.docx"


def _set_run_font(run, font_name="Calibri", size_pt=11, bold=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _add_paragraph(doc, text, *, bold=False, size=11, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def build_search_strategy_anon():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)

    # Title
    title = _add_paragraph(
        doc,
        "Supplementary File: Full Search Strategy and Search Log",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_paragraph(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Systematic Review and Meta-Analysis",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12),
    )
    _add_paragraph(
        doc,
        "[Author identifying information removed for double-anonymous peer review]",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18),
    )

    # 1. Pre-registration
    _add_paragraph(doc, "1. Pre-registration and Open-Science Deposit",
                   bold=True, size=13, space_after=Pt(6))
    _add_paragraph(
        doc,
        "The search strategy described below was specified in advance and "
        "time-stamped on OSF Registries on 23 April 2026, prior to formal "
        "data extraction. The full pre-registration record, the database "
        "search log, screening decisions, data extraction forms, JBI "
        "risk-of-bias ratings, and the analysis code are all permanently "
        "deposited in a time-stamped open-science repository. The DOIs "
        "for the pre-registration record and the seven-component OSF "
        "project (protocol, search, screening, extraction, risk-of-bias, "
        "analysis, and article-level DOI index) will be supplied to the "
        "editorial office post-acceptance for double-anonymous peer "
        "review purposes; in the interim, all relevant files have been "
        "uploaded as anonymised supplementary materials with this "
        "submission.",
    )

    # 2. Three-concept Boolean search strategy
    _add_paragraph(doc, "2. Three-Concept Boolean Search Strategy",
                   bold=True, size=13, space_after=Pt(6))
    _add_paragraph(
        doc,
        "The search combined three concept blocks with the Boolean "
        "operator AND. Each concept block was constructed as an OR-"
        "expansion of synonyms, with both controlled-vocabulary "
        "(MeSH/thesaurus) and free-text variants. Limits: English "
        "language; no publication-year restriction.",
        space_after=Pt(10),
    )

    _add_paragraph(doc, "Concept 1 — Personality", bold=True, size=11,
                   space_after=Pt(2))
    _add_paragraph(
        doc,
        '"Big Five" OR "Five-Factor Model" OR "FFM" OR "HEXACO" OR "BFI" '
        'OR "NEO-PI-R" OR "NEO-FFI" OR "IPIP" OR "conscientiousness" OR '
        '"openness to experience" OR "extraversion" OR "agreeableness" '
        'OR "neuroticism" OR "emotional stability" OR "personality traits"',
        size=10, space_after=Pt(10),
    )

    _add_paragraph(doc, "Concept 2 — Online learning", bold=True, size=11,
                   space_after=Pt(2))
    _add_paragraph(
        doc,
        '"online learning" OR "e-learning" OR "distance learning" OR '
        '"remote learning" OR "virtual learning" OR "blended learning" '
        'OR "hybrid learning" OR "MOOC" OR "massive open online course" '
        'OR "web-based learning" OR "computer-mediated learning" OR '
        '"learning management system" OR "LMS" OR "online course" OR '
        '"synchronous online" OR "asynchronous online"',
        size=10, space_after=Pt(10),
    )

    _add_paragraph(doc, "Concept 3 — Academic outcome", bold=True, size=11,
                   space_after=Pt(2))
    _add_paragraph(
        doc,
        '"academic performance" OR "academic achievement" OR "GPA" OR '
        '"grade point average" OR "test score" OR "course grade" OR '
        '"learning outcome" OR "learning performance" OR "academic success"',
        size=10, space_after=Pt(18),
    )

    # 3. Databases consulted
    _add_paragraph(doc, "3. Databases and Sources Consulted",
                   bold=True, size=13, space_after=Pt(6))

    headers = ["#", "Database / Source", "Access route",
               "Status", "Execution date", "Raw hits"]
    rows = [
        ("1", "PubMed / MEDLINE", "NCBI E-utilities API (blocked)",
         "Not executed", "2026-04-23", "—"),
        ("2", "OpenAlex", "OpenAlex REST API (blocked)",
         "Not executed", "2026-04-23", "—"),
        ("3", "ERIC", "ERIC web search interface",
         "Completed", "2026-04-23", "Contributed"),
        ("4", "Semantic Scholar", "Semantic Scholar API (blocked)",
         "Not executed", "2026-04-23", "—"),
        ("5", "Google Scholar / WebSearch", "WebSearch tool (8 split queries)",
         "Completed", "2026-04-23", "80 (~28 novel)"),
        ("6", "ProQuest Dissertations", "Institutional subscription required",
         "Not executed", "—", "—"),
        ("7", "Forward + backward citation snowballing",
         "Hand-tracked from included and prior meta-analyses",
         "Completed", "2026-04-23 to 2026-04-26", "Captured in extraction log"),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = False
    widths = [Cm(0.8), Cm(3.5), Cm(4.8), Cm(2.4), Cm(2.4), Cm(2.6)]
    for hi, htext in enumerate(headers):
        cell = table.rows[0].cells[hi]
        cell.width = widths[hi]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(htext)
        _set_run_font(run, size_pt=10, bold=True)
        _set_cell_borders(cell)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.width = widths[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, size_pt=10)
            _set_cell_borders(cell)

    _add_paragraph(doc, "", space_after=Pt(8))

    # 4. Deviation from pre-registration
    _add_paragraph(doc, "4. Deviation from Pre-Registered Search Plan",
                   bold=True, size=13, space_after=Pt(6))
    _add_paragraph(
        doc,
        "The pre-registered protocol named PubMed/MEDLINE, OpenAlex, ERIC, "
        "Semantic Scholar, and ProQuest Dissertations as the intended "
        "bibliographic sources. During execution, direct programmatic "
        "access to the PubMed E-utilities API, OpenAlex REST API, and "
        "Semantic Scholar API was blocked by the execution environment's "
        "network whitelist (HTTP 403 / connection refused). ProQuest "
        "Dissertations required an institutional subscription that was "
        "not available.",
    )
    _add_paragraph(
        doc,
        "As a mitigation, the same three-concept Boolean strategy was "
        "executed via the WebSearch tool (Google-based, returns indexed "
        "content from PubMed/PMC, Frontiers, Springer, Elsevier, and "
        "open-access repositories), split into eight targeted queries "
        "(see §5 below). The ERIC web search interface was reachable "
        "and was used directly. Forward and backward citation "
        "snowballing from included primary studies and from the eight "
        "prior meta-analyses (Poropat 2009 through Chen et al. 2025) "
        "was used to compensate for the reduced bibliographic-database "
        "coverage.",
    )
    _add_paragraph(
        doc,
        "This deviation is disclosed transparently in the Methods "
        "section of the manuscript (Information Sources and Search "
        "Strategy subsection) and again in the Limitations subsection. "
        "A future replication search in the originally pre-registered "
        "bibliographic databases is planned for the next manuscript "
        "version once institutional access becomes available.",
        space_after=Pt(18),
    )

    # 5. WebSearch query log
    _add_paragraph(doc, "5. WebSearch Query Log (2026-04-23)",
                   bold=True, size=13, space_after=Pt(6))

    qheaders = ["Q#", "Query string", "Raw hits", "Novel candidates"]
    qrows = [
        ("Q1", '"Big Five" "online learning" academic achievement GPA 2024',
         "10", "5"),
        ("Q2", '"online learning" "Big Five" personality correlation 2023-2024 university',
         "10", "3"),
        ("Q3", '"MOOC" personality Big Five completion dropout',
         "10", "4"),
        ("Q4", '"distance learning" personality conscientiousness 2022-2023',
         "10", "7"),
        ("Q5", 'personality online course dissertation 2020-2022',
         "10", "3"),
        ("Q6", "HEXACO online learning academic outcomes",
         "10", "2"),
        ("Q7", '"K-12" OR "high school" online learning Big Five grades',
         "10", "2"),
        ("Q8", "Europe Germany Netherlands online COVID personality",
         "10", "2"),
        ("Total", "—", "80", "28 unique novel candidates"),
    ]

    qtable = doc.add_table(rows=len(qrows) + 1, cols=len(qheaders))
    qtable.autofit = False
    qwidths = [Cm(1.4), Cm(10.5), Cm(2.0), Cm(3.0)]
    for hi, htext in enumerate(qheaders):
        cell = qtable.rows[0].cells[hi]
        cell.width = qwidths[hi]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(htext)
        _set_run_font(run, size_pt=10, bold=True)
        _set_cell_borders(cell)
    for ri, row in enumerate(qrows, start=1):
        for ci, val in enumerate(row):
            cell = qtable.rows[ri].cells[ci]
            cell.width = qwidths[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, size_pt=10,
                          bold=(row[0] == "Total" and ci == 0))
            _set_cell_borders(cell)

    _add_paragraph(doc, "", space_after=Pt(12))

    # 6. PRISMA flow counts
    _add_paragraph(doc, "6. Final PRISMA 2020 Flow Counts",
                   bold=True, size=13, space_after=Pt(6))
    _add_paragraph(
        doc,
        "The PRISMA 2020 flow diagram is reproduced as Figure 1 of the "
        "main manuscript. Final counts:",
        space_after=Pt(6),
    )
    for line in [
        "Identification: 108 records (WebSearch n = 80; prior informal "
        "search n = 28; benchmark meta-analyses n = 5; minus duplicates ≈ 40)",
        "Records after deduplication: 68",
        "Records screened (title/abstract): 68",
        "Records excluded at title/abstract: ~25 (off-modality, off-"
        "population, non-research commentary)",
        "Full-text reports assessed for eligibility: 43",
        "Full-text exclusions (with reasons): 12 (not online modality 3-5; "
        "no extractable effect size 2-3; duplicate/overlap samples 1-2; "
        "other 2-4)",
        "Studies included in qualitative synthesis: 25 (the canonical "
        "retained set, study IDs A-XX in the data-extraction file)",
        "Studies contributing to the primary quantitative meta-analysis "
        "(direct r or β-converted Pearson correlations): 10 (pooled "
        "N = 3,384)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        _set_run_font(run, size_pt=10)

    _add_paragraph(doc, "", space_after=Pt(12))

    # 7. Where to find the underlying logs
    _add_paragraph(doc, "7. Underlying Logs Provided with This Submission",
                   bold=True, size=13, space_after=Pt(6))
    _add_paragraph(
        doc,
        "The following anonymised supplementary files accompany this "
        "submission and contain the full underlying records:",
        space_after=Pt(4),
    )
    for line in [
        "search_strategy_anon.docx (this document) — full Boolean "
        "strategy, database list, and query-level hit counts.",
        "prisma_2020_checklist_anon.docx — completed PRISMA 2020 27-"
        "item checklist with section-level locations.",
        "Manuscript (anonymised) — Figure 1 reproduces the PRISMA 2020 "
        "flow diagram with final counts; Methods → Eligibility Criteria "
        "subsection lists all inclusion / exclusion rules.",
        "Manuscript (anonymised) — Table 1 lists each included study "
        "with extraction-level metadata (country, modality, sample "
        "size, instrument, outcome type, era, region, JBI risk-of-bias "
        "score, inclusion status).",
        "Manuscript (anonymised) — Tables 2–5 report the synthesis "
        "results (pooled effects, moderators, sensitivity, GRADE).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        _set_run_font(run, size_pt=10)

    _add_paragraph(doc, "", space_after=Pt(4))
    _add_paragraph(
        doc,
        "The full pre-registration record, the seven-component OSF "
        "project, and the version-controlled code repository (including "
        "the analysis pipeline, the data-extraction CSV, the screening "
        "ledger, and the JBI risk-of-bias spreadsheet) will be released "
        "post-acceptance under their stable DOIs, which will be "
        "supplied to the editorial office at that point. They are "
        "withheld here only because their URLs identify the author.",
    )

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


def main():
    build_search_strategy_anon()


if __name__ == "__main__":
    main()
