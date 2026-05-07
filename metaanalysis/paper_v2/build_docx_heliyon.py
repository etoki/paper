"""
build_docx_heliyon.py — Heliyon (Cell Press) numbered reference style.

Generates manuscript_journal_heliyon_v2.docx using the same source
content as build_docx.py but rewritten in the journal's required
numbered reference style:

  * In-text citations: APA "Author (Year)" / "(Author, Year)" patterns
    are converted to bracketed numbers, e.g. "Smith et al. [3]" or
    "[3, 5, 7]". Multi-citation parentheticals collapse into a single
    bracket group.
  * References section: numbered list ordered by first appearance in
    the body, rendered with the same italic-aware formatter used by
    build_docx.py (so journal/book titles remain italicized).

Implementation strategy: the manuscript is first built using
build_docx.py's existing builders (which produce APA citations), then
a post-process pass rewrites every paragraph's text and finally
appends the numbered References section in first-appearance order.

Limitations to review:
  * Citation rewriting collapses run-level formatting inside affected
    paragraphs to a single run. Statistical italics (r, p, N, k, etc.)
    are reapplied at the end via build_docx.italicize_stats_in_doc.
  * Edge-case patterns (e.g. parenthetical group with explanatory text
    such as "(see Smith, 2020)") are preserved and only the citation
    portion is replaced inside the original parentheses.
  * Citation keys are matched on (FirstAuthorSurname, Year). Unknown
    keys are left unchanged so the regression is visible to a reviewer.
"""

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_docx
from references_data import REFERENCES


# ---------------------------------------------------------------------------
# Citation key extraction from references_data.py
# ---------------------------------------------------------------------------

# Unicode-aware surname character class. [^\W\d_] = any letter (incl. Unicode),
# excluding digits and underscores. Allows hyphens and apostrophes inside
# surnames (e.g., O'Brien, Smith-Jones).
_SURNAME = r"[A-ZÀ-ÖØ-Ý][^\W\d_,&()'\s]*(?:[\-'ʼ][A-ZÀ-ÖØ-Ý][^\W\d_,&()'\s]*)*"

FIRST_SURNAME_RE = re.compile(rf"^({_SURNAME})(?:,| ?&| and )", re.UNICODE)
YEAR_RE = re.compile(r"\((\d{4}[a-z]?)\)")

# Capture every author surname appearing in the pre-year segment of a
# reference entry — surnames are followed by ", I." (initials) or ", &" or
# " and " or ", ".
ALL_AUTHORS_RE = re.compile(
    rf"\b({_SURNAME})(?=,\s*[A-ZÀ-ÖØ-Ý]\.|\s*&|\s+and\b)",
    re.UNICODE,
)


def _strip_diacritics(s: str) -> str:
    """ASCII-fold diacritics (Alkış -> Alkis, Bahçekapılı -> Bahcekapili)."""
    import unicodedata
    nfkd = unicodedata.normalize("NFKD", s)
    ascii_only = "".join(c for c in nfkd if not unicodedata.combining(c))
    # Special-case Turkish dotless ı / dotted İ which NFKD does not strip
    return (
        ascii_only.replace("ı", "i").replace("İ", "I")
                  .replace("ş", "s").replace("Ş", "S")
                  .replace("ğ", "g").replace("Ğ", "G")
                  .replace("ç", "c").replace("Ç", "C")
                  .replace("ø", "o").replace("Ø", "O")
                  .replace("æ", "ae").replace("Æ", "Ae")
    )


def extract_first_key(ref_text: str):
    """Extract (FirstAuthor surname, Year) citation key from an APA entry."""
    m1 = FIRST_SURNAME_RE.match(ref_text)
    m2 = YEAR_RE.search(ref_text)
    if m1 and m2:
        return (m1.group(1), m2.group(1))
    return None


def extract_all_keys(ref_text: str):
    """Extract every (surname, year) tuple a reader might use to cite this entry,
    including co-authors who may appear as the head of an "et al." narrative."""
    m_year = YEAR_RE.search(ref_text)
    if not m_year:
        return []
    year = m_year.group(1)
    pre = ref_text[: m_year.start()]
    surnames = ALL_AUTHORS_RE.findall(pre)
    # Also include the very last author before "(YYYY)" — typically rendered
    # as "..., & Surname, I. I.". The trailing surname has no following "&" so
    # ALL_AUTHORS_RE misses it; capture it separately.
    tail_re = re.compile(rf",\s*&\s*({_SURNAME}),\s*[A-ZÀ-ÖØ-Ý]\.", re.UNICODE)
    tail = tail_re.findall(pre)
    surnames.extend(tail)
    # Single-author papers ("Surname, I. (YYYY)") — capture standalone surname.
    solo_re = re.compile(rf"^({_SURNAME}),\s*[A-ZÀ-ÖØ-Ý]\.\s*\(", re.UNICODE)
    m_solo = solo_re.match(ref_text)
    if m_solo:
        surnames.append(m_solo.group(1))
    return [(s, year) for s in dict.fromkeys(surnames)]  # dedup, preserve order


REF_KEYS = [extract_first_key(r) for r in REFERENCES]

# Build (surname, year) -> ref_idx map. First-author keys win over co-author
# keys to avoid Wang(2023) resolving to Tlili et al. (2023) just because
# "Wang, H." is a co-author there. Diacritic-folded variants are also added
# so the body's "Alkis (2018)" matches references' "Alkış, N., ..."
KEY_TO_REFIDX = {}

# Pass 1: register first-author keys (highest priority).
for i, ref in enumerate(REFERENCES):
    fk = extract_first_key(ref)
    if fk:
        KEY_TO_REFIDX[fk] = i
        folded = (_strip_diacritics(fk[0]), fk[1])
        KEY_TO_REFIDX.setdefault(folded, i)

# Pass 2: register co-author keys as fallback (don't override first-author).
for i, ref in enumerate(REFERENCES):
    for key in extract_all_keys(ref):
        KEY_TO_REFIDX.setdefault(key, i)
        folded = (_strip_diacritics(key[0]), key[1])
        KEY_TO_REFIDX.setdefault(folded, i)


# ---------------------------------------------------------------------------
# In-text citation patterns
# ---------------------------------------------------------------------------

NARRATIVE_RE = re.compile(
    rf"\b({_SURNAME})"
    rf"(?:\s+(?:and|&)\s+{_SURNAME})?"
    r"(\s+et\s+al\.)?"
    r"(?:'s|’s)?"
    r"\s+\((\d{4}[a-z]?)\)",
    re.UNICODE,
)

PAREN_RE = re.compile(r"\(([^()]*?\d{4}[a-z]?[^()]*?)\)")
INNER_CITATION_RE = re.compile(
    rf"\b({_SURNAME})"
    rf"(?:\s+(?:and|&)\s+{_SURNAME})?"
    r"(?:\s+et\s+al\.)?"
    r",\s*(\d{4}[a-z]?)",
    re.UNICODE,
)


class CitationNumbering:
    def __init__(self):
        self.cited_order = []
        self.refidx_to_num = {}
        self.unknown_keys = []

    def number_for_key(self, key):
        idx = KEY_TO_REFIDX.get(key)
        if idx is None:
            # Fall back to diacritic-folded surname lookup.
            folded = (_strip_diacritics(key[0]), key[1])
            idx = KEY_TO_REFIDX.get(folded)
        if idx is None:
            if key not in self.unknown_keys:
                self.unknown_keys.append(key)
            return None
        if idx not in self.refidx_to_num:
            self.cited_order.append(idx)
            self.refidx_to_num[idx] = len(self.cited_order)
        return self.refidx_to_num[idx]


def replace_narrative(text, numbering):
    def sub(m):
        surname = m.group(1)
        et_al = m.group(2) or ""
        year = m.group(3)
        n = numbering.number_for_key((surname, year))
        if n is None:
            return m.group(0)
        return f"{m.group(0).split('(')[0].rstrip()} [{n}]"
    return NARRATIVE_RE.sub(sub, text)


def replace_parenthetical(text, numbering):
    def outer_sub(m):
        inner = m.group(1)
        nums = []
        placeholders = {}

        def inner_sub(im):
            surname = im.group(1)
            year = im.group(2)
            n = numbering.number_for_key((surname, year))
            if n is None:
                return im.group(0)
            nums.append(n)
            placeholder = f"__N{n}__"
            placeholders[placeholder] = n
            return placeholder

        new_inner = INNER_CITATION_RE.sub(inner_sub, inner)
        if not nums:
            return m.group(0)

        residue = re.sub(r"__N\d+__", "", new_inner)
        residue = re.sub(r"[;,\s]+", "", residue)
        if not residue:
            return f"[{', '.join(str(n) for n in nums)}]"

        merged = new_inner
        for ph, n in placeholders.items():
            merged = merged.replace(ph, f"[{n}]", 1)
        return f"({merged})"

    return PAREN_RE.sub(outer_sub, text)


def convert_text(text, numbering):
    text = replace_narrative(text, numbering)
    text = replace_parenthetical(text, numbering)
    return text


# ---------------------------------------------------------------------------
# Document rewrite
# ---------------------------------------------------------------------------

def rewrite_paragraph(p, numbering):
    full = "".join(r.text for r in p.runs)
    new = convert_text(full, numbering)
    if new == full:
        return
    first_run = p.runs[0] if p.runs else None
    for r in p.runs[1:]:
        r.text = ""
    if first_run is not None:
        first_run.text = new
    else:
        p.add_run(new)


def rewrite_doc(doc, numbering):
    for p in doc.paragraphs:
        rewrite_paragraph(p, numbering)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rewrite_paragraph(p, numbering)


# ---------------------------------------------------------------------------
# Numbered References section (Cell Press / Vancouver-like)
# ---------------------------------------------------------------------------

def build_numbered_references_section(doc, numbering):
    doc.add_page_break()
    heading = doc.add_paragraph("References", style="Heading 1")
    build_docx.set_double_space(heading)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    build_docx.set_cell_font(heading, bold=True)

    if not numbering.cited_order:
        return

    italic_pattern = re.compile(r"(<i>.*?</i>)", re.DOTALL)

    for n, idx in enumerate(numbering.cited_order, start=1):
        ref_text = REFERENCES[idx]
        prefixed = f"{n}. {ref_text}"

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 2.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)

        for part in italic_pattern.split(prefixed):
            if not part:
                continue
            if part.startswith("<i>") and part.endswith("</i>"):
                run = p.add_run(part[3:-4])
                run.italic = True
                build_docx._set_run_font(run)
            else:
                run = p.add_run(part)
                build_docx._set_run_font(run)


# ---------------------------------------------------------------------------
# Heliyon build entry
# ---------------------------------------------------------------------------

def build_heliyon_manuscript(output_path):
    doc = Document()
    build_docx.configure_page(doc)
    build_docx.configure_styles(doc)
    build_docx.build_title_page(doc)
    doc.add_page_break()
    build_docx.build_declarations(doc, for_journal=True)
    doc.add_page_break()
    build_docx.build_abstract(doc)
    build_docx.build_intro_part1(doc)
    build_docx.build_intro_part2(doc)
    build_docx.build_intro_part3(doc)
    build_docx.build_methods_part1(doc)
    build_docx.build_methods_part2(doc)
    build_docx.build_methods_part3(doc)
    build_docx.build_methods_part4(doc)
    build_docx.build_methods_part5(doc)
    build_docx.build_results_part1(doc)
    build_docx.build_results_part2(doc)
    build_docx.build_results_part3(doc)
    build_docx.build_discussion_part1(doc)
    build_docx.build_discussion_part2(doc)
    build_docx.build_discussion_part3(doc)
    build_docx.build_conclusion(doc)

    numbering = CitationNumbering()
    rewrite_doc(doc, numbering)

    build_numbered_references_section(doc, numbering)

    build_docx.build_table1_characteristics(doc)
    build_docx.build_table2_pooled(doc)
    build_docx.build_table3_moderators(doc)
    build_docx.build_table4_sensitivity(doc)
    build_docx.build_table5_grade(doc)
    build_docx.build_figure1_prisma(doc)
    build_docx.build_forest_plots(doc)
    build_docx.build_funnel_plots(doc)

    rewrite_doc(doc, numbering)

    build_docx.italicize_stats_in_doc(doc)

    doc.save(output_path)
    print(f"Wrote {output_path}")
    print(f"  cited references: {len(numbering.cited_order)} / {len(REFERENCES)} total")
    if numbering.unknown_keys:
        print(f"  ⚠ unknown citation keys (not in references_data.py):")
        for k in numbering.unknown_keys:
            print(f"     - {k[0]} ({k[1]})")
    uncited = [i for i in range(len(REFERENCES)) if i not in numbering.refidx_to_num]
    if uncited:
        print(f"  ⚠ references in references_data.py never cited in body:")
        for i in uncited:
            key = REF_KEYS[i]
            print(f"     - {key[0] if key else '?'} ({key[1] if key else '?'})")


def main():
    here = Path(__file__).resolve().parent
    out = here / "manuscript_journal_heliyon_v2.docx"
    build_heliyon_manuscript(str(out))


if __name__ == "__main__":
    main()
