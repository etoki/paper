"""Generate journal-tailored cover letters for v2 manuscript.

Targets:

- Frontiers in Education — author-initiated resubmission of the v2
  refactor (already under the journal's submission system; editor invited
  resubmission after the author flagged corrections)
- Education Sciences (MDPI) — ~$1,800 ≈ 27万円
- Systematic Reviews (BMC) — ~$2,545 ≈ 38万円
- Heliyon (Elsevier) — ~$1,890 ≈ 28万円

Each cover letter is tailored to the journal's scope, audience, and
editorial culture. All share the same factual core (preregistration,
preprint, OSF deposits, COI).
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADER_LINES = [
    "Eisuke Tokiwa, MEng",
    "Founder, SUNBLAZE Co., Ltd.",
    "Tokyo, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Email: eisuke.tokiwa@sunblaze.jp",
]

SIGNATURE_LINES = [
    "Eisuke Tokiwa",
    "Founder, SUNBLAZE Co., Ltd.",
    "Tokyo, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Email: eisuke.tokiwa@sunblaze.jp",
]

# Shared factual blocks (identical across journals)
SHARED_OPEN_SCIENCE = (
    "__Open science.__ "
    "Pre-registration: https://doi.org/10.17605/OSF.IO/E5W47 (OSF Registries, "
    "registered 23 April 2026, prior to formal data extraction). "
    "Preprint: https://doi.org/10.21203/rs.3.rs-9513298 (Research Square v1, "
    "posted 27 April 2026). Data, analysis code, search logs, screening "
    "decisions, risk-of-bias ratings, and supplementary materials are "
    "publicly available on the OSF project (https://doi.org/10.17605/OSF.IO/79M5J), "
    "comprising seven separately DOI-tagged components covering protocol, "
    "search, screening, extraction, risk of bias, analysis, and the article-"
    "level DOI index. A version-controlled mirror is maintained at "
    "https://github.com/etoki/paper."
)

SHARED_COI = (
    "__Conflict of interest.__ "
    "I declare no financial conflicts of interest. One of my own prior primary "
    "studies (Tokiwa, 2025, manuscript in preparation) is potentially eligible "
    "for inclusion; this is addressed transparently through a pre-specified "
    "sensitivity analysis excluding the author's own study, which leaves the "
    "primary conclusions unchanged (|Δr| < .001 for every trait, because that "
    "study did not contribute extractable zero-order correlations to the "
    "primary quantitative pool)."
)

SHARED_NONDUP = (
    "The manuscript has not been published elsewhere and is not under "
    "consideration by another journal. The Research Square preprint (v1) is "
    "explicitly disclosed above; any substantive revisions during peer review "
    "will be reflected in a versioned preprint update under the same DOI stem. "
    "I am the sole author and accept full accountability for all aspects of "
    "the work."
)

# ---------------------------------------------------------------------------
# Frontiers in Education (author-initiated resubmission)
# ---------------------------------------------------------------------------
FE_ADDRESSEE = [
    "To the Editors,",
    "Frontiers in Education",
]

FE_BODY = [
    "Dear Editors,",

    ("Thank you for your message inviting resubmission. Following my "
     "earlier submission of this manuscript to Frontiers in Education, "
     "I conducted an extensive self-audit and substantive refactor of the "
     "analysis pipeline and reporting. I contacted the editorial office "
     "to flag that corrections were warranted, and per your reply I am "
     "now resubmitting the corrected version (v2) for consideration."),

    ("__Summary of changes since the initial submission.__ "
     "(1) The full quantitative pipeline was re-run after re-curating the "
     "primary study set, yielding a transparent and now-canonical record: "
     "31 primary studies catalogued at full-text assessment, 25 retained "
     "for qualitative synthesis, and 10 contributing direct or β-converted "
     "Pearson correlations to the primary quantitative pool (pooled "
     "N = 3,384). All Tables 2–5 and the body-text effect sizes are "
     "regenerated from the canonical pipeline (REML τ² with Hartung-Knapp-"
     "Sidik-Jonkman CIs). "
     "(2) An adapted GRADE assessment was added (Table 5), yielding "
     "Moderate confidence for Conscientiousness and Extraversion, and Low "
     "for Openness, Agreeableness, and Neuroticism. "
     "(3) Seven pre-specified sensitivity analyses (COI exclusion, "
     "Peterson-Brown β-converted exclusion, RoB < 5 exclusion, leave-one-"
     "out, alternative τ² estimators, HEXACO crosswalk variants, small-"
     "sample exclusion) and a publication-bias suite (Egger, trim-and-"
     "fill, p-curve) were completed and reported. "
     "(4) The full PRISMA 2020 checklist and a Declaration of Interest "
     "form are attached as supplementary files. "
     "(5) Reference list and citation entries were audited end-to-end "
     "(authors, years, volumes, pages, and DOIs cross-validated against "
     "Crossref); fabricated or imprecise bibliographic fields identified "
     "during the audit were corrected. "
     "(6) An automated hallucination-check protocol (T1–T9) was developed "
     "to verify every numerical claim in the manuscript against the "
     "canonical analysis outputs; all checks pass on the resubmitted "
     "version."),

    ("__Why the work belongs in Frontiers in Education.__ "
     "Frontiers in Education has an established record of publishing "
     "rigorously preregistered, openly reported reviews on technology-"
     "supported learning and individual differences in education. The "
     "manuscript reports the first quantitative meta-analytic synthesis "
     "of Big Five personality traits and academic achievement dedicated "
     "specifically to online learning environments, with two pre-"
     "registered moderator findings (Extraversion × Region and "
     "Extraversion × Outcome Type) that are directly relevant to readers "
     "designing online and blended curricula across diverse cultural and "
     "assessment contexts."),

    ("__Headline findings.__ "
     "Random-effects meta-analysis with REML estimation and Hartung-"
     "Knapp-Sidik-Jonkman CI adjustment yielded Conscientiousness "
     "r = .167, 95% CI [.089, .243]; Agreeableness r = .112; Openness "
     "r = .086; Neuroticism r = .018; Extraversion r = .002. Two pre-"
     "registered moderator effects were highly significant: Extraversion "
     "× Region (Asian r = −.131 vs. non-Asian r = .050; "
     "Q_between = 46.43, p < .001) and Extraversion × Outcome Type "
     "(objective r = −.038 vs. self-rated r = .117; Q_between = 17.30, "
     "p < .001). Conscientiousness remains the strongest Big Five "
     "predictor in online environments but at attenuated magnitude "
     "relative to face-to-face benchmarks (ρ ≈ .22–.28 in eight prior "
     "meta-analyses)."),

    SHARED_OPEN_SCIENCE,

    SHARED_COI,

    SHARED_NONDUP,

    "Thank you for accommodating this author-initiated correction and "
    "resubmission. I am happy to provide any further information the "
    "editorial team may require.",

    "Sincerely,",
]

# ---------------------------------------------------------------------------
# Education Sciences (MDPI)
# ---------------------------------------------------------------------------
ES_ADDRESSEE = [
    "To the Editors,",
    "Education Sciences (MDPI)",
]

ES_BODY = [
    "Dear Editors,",

    "I am pleased to submit “Big Five Personality Traits and Academic "
    "Achievement in Online Learning Environments: A Systematic Review and "
    "Meta-Analysis” for consideration as a Review article in Education "
    "Sciences. The manuscript reports a pre-registered, PRISMA 2020-compliant "
    "quantitative synthesis at the intersection of educational psychology, "
    "personality science, and online/distance learning research.",

    ("__Background and gap.__ "
     "Conscientiousness has been established as the strongest Big Five "
     "predictor of academic achievement across eight prior meta-analyses "
     "(Poropat, 2009 through Chen et al., 2025), but those syntheses pooled "
     "face-to-face, blended, and online samples without testing learning "
     "modality as a moderator. With online and blended instruction now "
     "constituting a substantial share of post-secondary education globally, "
     "an online-specific quantitative synthesis is needed to determine "
     "whether established face-to-face benchmarks generalize to technology-"
     "mediated contexts."),

    ("__Methods and findings.__ "
     "The review followed a pre-registered protocol (OSF Registries E5W47, "
     "registered 23 April 2026) and PRISMA 2020 reporting standards. From "
     "31 primary studies catalogued at full-text assessment, 25 were retained "
     "for qualitative synthesis and 10 contributed direct or β-converted "
     "Pearson correlations to the primary quantitative pool (pooled "
     "N = 3,384). Random-effects meta-analysis with REML estimation and "
     "Hartung-Knapp-Sidik-Jonkman confidence-interval adjustment yielded "
     "Conscientiousness r = .167, 95% CI [.089, .243]; Agreeableness "
     "r = .112; Openness r = .086; Neuroticism r = .018; Extraversion "
     "r = .002. Two pre-registered moderator effects were highly significant: "
     "Extraversion × Region (Asian r = −.131 vs. non-Asian r = .050; "
     "Q_between = 46.43, p < .001) and Extraversion × Outcome Type "
     "(objective r = −.038 vs. self-rated r = .117; Q_between = 17.30, "
     "p < .001), indicating that the Extraversion–achievement association "
     "shifts negatively in Asian samples and in objective achievement outcomes. "
     "GRADE confidence ratings ranged from Moderate (Conscientiousness, "
     "Extraversion) to Low (Openness, Agreeableness, Neuroticism). Seven "
     "pre-specified sensitivity analyses confirm robustness."),

    ("__Fit with Education Sciences.__ "
     "Education Sciences has an established record of publishing rigorous "
     "systematic reviews and meta-analyses on technology-supported learning, "
     "personality and individual differences in education, and post-pandemic "
     "educational change. The manuscript's combination of pre-registered "
     "methodology, transparent open-science deposits, and online-specific "
     "focus aligns directly with the journal's mission of methodologically "
     "robust, openly reported education research. Findings are directly "
     "relevant to readers designing online and blended curricula and to "
     "researchers refining personality-context interaction theory."),

    SHARED_OPEN_SCIENCE,

    SHARED_COI,

    SHARED_NONDUP,

    "All MDPI submission requirements are addressed: ORCID provided in the "
    "header, structured abstract conformant to journal guidelines, "
    "supplementary materials linked from the OSF deposit, and ICMJE "
    "authorship criteria met (sole author).",

    "Thank you for considering this manuscript. I look forward to your "
    "decision.",

    "Sincerely,",
]

# ---------------------------------------------------------------------------
# Systematic Reviews (BMC)
# ---------------------------------------------------------------------------
SR_ADDRESSEE = [
    "To the Editors,",
    "Systematic Reviews (BMC)",
]

SR_BODY = [
    "Dear Editors,",

    "I am pleased to submit “Big Five Personality Traits and Academic "
    "Achievement in Online Learning Environments: A Systematic Review and "
    "Meta-Analysis” for consideration in Systematic Reviews. The manuscript "
    "reports a pre-registered, PRISMA 2020-compliant systematic review and "
    "meta-analysis with a fully reproducible analytic pipeline.",

    ("__Why this synthesis is needed.__ "
     "Eight prior meta-analyses have established Conscientiousness as the "
     "strongest Big Five predictor of academic achievement (ρ ≈ "
     ".19–.28; Poropat, 2009 through Chen et al., 2025). However, every "
     "one of these syntheses has aggregated face-to-face, blended, and "
     "online samples without testing learning modality as a substantive "
     "moderator. As post-secondary instruction is increasingly delivered "
     "online, the question of whether face-to-face benchmarks generalize to "
     "technology-mediated environments has been empirically open and "
     "societally consequential. The present review is, to my knowledge, the "
     "first quantitative synthesis dedicated specifically to online learning "
     "environments."),

    ("__Methodological rigor (highlights for Systematic Reviews readers).__"
     " (1) Protocol pre-registered on OSF Registries (E5W47) on 23 April "
     "2026, prior to formal data extraction; PRISMA 2020-compliant reporting "
     "with the completed checklist deposited as supplementary material. "
     "(2) Effect-size synthesis on the Fisher z scale with REML τ² "
     "estimation and Hartung-Knapp-Sidik-Jonkman confidence-interval "
     "adjustment to control Type I error in small-k scenarios. (3) Risk of "
     "bias assessed with the Joanna Briggs Institute 8-item checklist for "
     "analytical cross-sectional studies; intra-rater reliability target "
     "κ ≥ 0.80 met for screening, full-text assessment, extraction, and "
     "RoB rating. (4) Seven pre-specified sensitivity analyses (COI "
     "exclusion, Peterson-Brown β-converted exclusion, RoB < 5 "
     "exclusion, leave-one-out, alternative τ² estimators, HEXACO "
     "crosswalk variants, small-sample exclusion) all confirm robustness "
     "of the primary conclusions. (5) Publication-bias assessment using "
     "Egger's regression, Duval and Tweedie's trim-and-fill, and Simonsohn "
     "et al.'s p-curve. (6) Certainty of evidence rated using an "
     "adaptation of GRADE for correlational syntheses, yielding ratings "
     "ranging from Moderate (Conscientiousness, Extraversion) to Low "
     "(Openness, Agreeableness, Neuroticism). (7) Three pre-registered "
     "deviations are transparently disclosed (database access limitations, "
     "moderator quantitative-vs-narrative reduction with k-per-level "
     "documentation, and the post-registration addition of one benchmark "
     "meta-analysis to the introductory comparison set)."),

    ("__Principal results.__ "
     "Across 31 catalogued primary studies (25 retained for qualitative "
     "synthesis, 10 contributing to the primary quantitative pool; pooled "
     "N = 3,384), Conscientiousness emerged as the strongest predictor "
     "(r = .167, 95% CI [.089, .243]) but at attenuated magnitude relative "
     "to face-to-face benchmarks (ρ ≈ .22–.28). Two pre-"
     "registered moderator effects were highly significant: Extraversion × "
     "Region (Q_between = 46.43, p < .001) and Extraversion × Outcome Type "
     "(Q_between = 17.30, p < .001), the latter being, to my knowledge, the "
     "first meta-analytic evidence that the choice of objective vs. self-"
     "rated outcome systematically alters the Extraversion–achievement "
     "association in online learning environments."),

    ("__Fit with Systematic Reviews.__ "
     "Systematic Reviews is the natural venue for a methodologically rigorous, "
     "preregistered, PRISMA 2020-compliant synthesis with a fully reproducible "
     "analytic pipeline. The manuscript's strict adherence to systematic "
     "review methodology, its transparent disclosure of pre-registered "
     "deviations, and its publicly deposited data + code + search logs are "
     "directly aligned with the journal's editorial mission."),

    SHARED_OPEN_SCIENCE,

    SHARED_COI,

    SHARED_NONDUP,

    "Thank you for considering this manuscript. I look forward to your "
    "decision.",

    "Sincerely,",
]

# ---------------------------------------------------------------------------
# Heliyon (Elsevier)
# ---------------------------------------------------------------------------
HE_ADDRESSEE = [
    "To the Editors,",
    "Heliyon",
]

HE_BODY = [
    "Dear Editors,",

    "I am pleased to submit “Big Five Personality Traits and Academic "
    "Achievement in Online Learning Environments: A Systematic Review and "
    "Meta-Analysis” for consideration in Heliyon, with the suggested "
    "section “Education” or “Psychology”.",

    ("__Significance and novelty.__ "
     "Eight prior meta-analyses have established Conscientiousness as the "
     "strongest Big Five predictor of academic achievement (ρ ≈ "
     ".19–.28), but no prior synthesis has tested learning modality "
     "as a substantive moderator. With online and blended instruction now "
     "constituting a substantial share of post-secondary education, the "
     "generalizability of face-to-face benchmarks to technology-mediated "
     "environments has been an empirically open question. This pre-"
     "registered, PRISMA 2020-compliant review is, to my knowledge, the "
     "first quantitative meta-analytic synthesis dedicated to online "
     "learning environments."),

    ("__Summary of methods and findings.__ "
     "The review followed a pre-registered protocol (OSF Registries E5W47, "
     "registered 23 April 2026, prior to formal data extraction). From 31 "
     "primary studies catalogued at full-text assessment, 25 were retained "
     "for qualitative synthesis and 10 contributed direct or β-converted "
     "Pearson correlations to the primary quantitative pool (total pooled "
     "N = 3,384). Pooled estimates from random-effects REML meta-analysis "
     "with Hartung-Knapp-Sidik-Jonkman adjustment: Conscientiousness "
     "r = .167, 95% CI [.089, .243]; Agreeableness r = .112; Openness "
     "r = .086; Neuroticism r = .018; Extraversion r = .002. Two pre-"
     "registered moderator effects were highly significant: Extraversion × "
     "Region (Asian r = −.131 vs. non-Asian r = .050; Q_between = 46.43, "
     "p < .001) and Extraversion × Outcome Type (objective r = −.038 vs. "
     "self-rated r = .117; Q_between = 17.30, p < .001). GRADE confidence "
     "ratings range from Moderate (Conscientiousness, Extraversion) to Low "
     "(Openness, Agreeableness, Neuroticism). Seven pre-specified "
     "sensitivity analyses confirm robustness."),

    ("__Fit with Heliyon.__ "
     "Heliyon's broad multidisciplinary scope is well suited to a manuscript "
     "that sits at the intersection of educational psychology, personality "
     "science, and online/distance learning research. Several primary studies "
     "in the synthesis are themselves Elsevier titles, including The Internet "
     "and Higher Education (Abe, 2020), Personality and Individual Differences "
     "(Quigley et al., 2022), and Computers in Human Behavior (Cohen & "
     "Baruth, 2017). Heliyon's commitment to publishing methodologically "
     "rigorous research regardless of perceived novelty premium is directly "
     "aligned with this manuscript's emphasis on transparent reporting and "
     "open data."),

    SHARED_OPEN_SCIENCE,

    SHARED_COI,

    SHARED_NONDUP,

    "All Elsevier submission requirements are addressed: ORCID, structured "
    "abstract, declarations section, data availability statement linked to "
    "the OSF deposit, and CRediT statement (sole author).",

    "Thank you for considering this submission. I look forward to your "
    "decision.",

    "Sincerely,",
]


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------
def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    """Add a paragraph; supports inline bold via __segment__ markers."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15

    parts = text.split("__")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    return p


def build_document(addressee, body, out_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for line in HEADER_LINES:
        add_paragraph(doc, line, space_after=Pt(0))
    add_paragraph(doc, "", space_after=Pt(0))

    add_paragraph(doc, date.today().strftime("%B %d, %Y"))
    add_paragraph(doc, "", space_after=Pt(0))

    for line in addressee:
        add_paragraph(doc, line, space_after=Pt(0))
    add_paragraph(doc, "", space_after=Pt(0))

    for paragraph in body:
        add_paragraph(doc, paragraph)

    for line in SIGNATURE_LINES:
        add_paragraph(doc, line, space_after=Pt(0))

    doc.save(str(out_path))


def main():
    here = Path(__file__).resolve().parent
    targets = [
        ("frontiers_in_education", FE_ADDRESSEE, FE_BODY),
        ("education_sciences_mdpi", ES_ADDRESSEE, ES_BODY),
        ("systematic_reviews_bmc", SR_ADDRESSEE, SR_BODY),
        ("heliyon", HE_ADDRESSEE, HE_BODY),
    ]
    for slug, addressee, body in targets:
        out = here / f"cover_letter_{slug}.docx"
        build_document(addressee, body, out)
        print(f"Wrote {out}")


if __name__ == "__main__":
    main()
