"""
Build the anonymised HSSC response cover letter as a PDF.

Springer Nature HSSC uses double-anonymous peer review and warns that
the point-by-point response file "may be shown to reviewers as part of
our assessment". The submission portal also accepts only PDF for that
field. This script:

  1. Builds the standard HSSC cover letter via build_cover_letter.py.
  2. Strips identifying tokens (author name, ORCID, email, affiliation,
     OSF/preprint/code URLs) from header, body, and signature.
  3. Saves the redacted version as cover_letter_hssc_springer_anon.docx.
  4. Converts to PDF via mammoth (docx -> HTML) and weasyprint
     (HTML -> PDF), producing cover_letter_hssc_springer_anon.pdf —
     the file to upload to the "Point-by-point response (optional)"
     slot of the HSSC submission portal.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import build_cover_letter
from build_docx_anonymous import REPLACEMENTS, rewrite_paragraph


ANON_HEADER = [
    "[Author identifying header — name, affiliation, ORCID, email — removed",
    "for double-anonymous peer review; will be reinstated upon acceptance.]",
]

ANON_SIGNATURE = [
    "[Author signature removed for double-anonymous peer review.]",
]


def build_anonymous_hssc_docx(out_path: Path) -> None:
    """Render the HSSC cover letter with anonymised header / signature, then
    redact identifying tokens from the body using the same replacement table
    as the manuscript anonymiser."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for line in ANON_HEADER:
        build_cover_letter.add_paragraph(doc, line, space_after=Pt(0))
    build_cover_letter.add_paragraph(doc, "", space_after=Pt(0))

    from datetime import date
    build_cover_letter.add_paragraph(doc, date.today().strftime("%B %d, %Y"))
    build_cover_letter.add_paragraph(doc, "", space_after=Pt(0))

    for line in build_cover_letter.HSSC_ADDRESSEE:
        build_cover_letter.add_paragraph(doc, line, space_after=Pt(0))
    build_cover_letter.add_paragraph(doc, "", space_after=Pt(0))

    for paragraph in build_cover_letter.HSSC_BODY:
        build_cover_letter.add_paragraph(doc, paragraph)

    for line in ANON_SIGNATURE:
        build_cover_letter.add_paragraph(doc, line, space_after=Pt(0))

    n = 0
    for p in doc.paragraphs:
        if rewrite_paragraph(p, REPLACEMENTS):
            n += 1

    doc.save(str(out_path))
    print(f"Wrote {out_path}  (paragraphs redacted: {n})")


def convert_to_pdf(docx_path: Path) -> Path:
    """Convert a .docx to .pdf via mammoth (docx -> HTML) and weasyprint
    (HTML -> PDF). Avoids the LibreOffice/Java dependency, which fails to
    load .docx files in this sandbox."""
    import mammoth
    import weasyprint

    pdf_path = docx_path.with_suffix(".pdf")
    with open(docx_path, "rb") as f:
        html_body = mammoth.convert_to_html(f).value
    full_html = (
        '<!DOCTYPE html><html><head><meta charset="utf-8"><style>'
        "@page { size: letter; margin: 2cm 2.5cm; }"
        'body { font-family: "DejaVu Sans", Arial, sans-serif; '
        "font-size: 11pt; line-height: 1.4; }"
        "p { margin: 6pt 0; }"
        "strong, b { font-weight: bold; }"
        "</style></head><body>" + html_body + "</body></html>"
    )
    weasyprint.HTML(string=full_html).write_pdf(str(pdf_path))
    print(f"Wrote {pdf_path}")
    return pdf_path


def main():
    here = Path(__file__).resolve().parent
    docx_path = here / "cover_letter_hssc_springer_anon.docx"
    build_anonymous_hssc_docx(docx_path)
    convert_to_pdf(docx_path)


if __name__ == "__main__":
    main()
