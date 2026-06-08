"""
build_frontiers_r1_response.py — Generate point-by-point response to
Reviewer 1 for Frontiers in Education (Manuscript ID 1866537).

Reviewer 1 issued a "Major revision required" recommendation on
25 May 2026 with six substantive concerns. Reviewer 2 endorsed
publication ("Accept in current form").

This script generates response_frontiers_reviewer1.docx — a non-
anonymous point-by-point response (Frontiers uses single-anonymous
review, so the author is identifiable) that addresses each Reviewer 1
concern with (i) acknowledgement, (ii) concrete manuscript revisions
made, and (iii) section/line references.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


HERE = Path(__file__).resolve().parent
OUT = HERE / "response_frontiers_reviewer1.docx"


def _add_p(doc, text, *, bold=False, italic=False, size=11, align=None,
           space_after=Pt(6), space_before=Pt(0), indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.3
    if indent is not None:
        p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def build_response():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

    _add_p(
        doc,
        "Point-by-Point Response to Reviewer 1",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_p(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Systematic Review and Meta-Analysis",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6),
    )
    _add_p(
        doc,
        f"Frontiers in Education  |  Manuscript ID 1866537  |  "
        f"Response date: {date.today().strftime('%d %B %Y')}",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18),
    )

    _add_p(
        doc,
        "I am very grateful to Reviewer 1 for the thorough and constructive "
        "evaluation. Reviewer 1's central observation — that the strength "
        "of the conclusions in the previous version exceeded what the "
        "small primary evidence base could sustain — is well taken, and "
        "I have substantially revised the manuscript in response. The "
        "revisions are concession-leaning throughout: where Reviewer 1 "
        "identifies a tension between the evidence and the strength of a "
        "claim, the claim has been weakened rather than the evidence "
        "defended. I have also taken the opportunity to clean up the "
        "placeholder text and forthcoming-deposit references that "
        "Reviewer 1 noted, and to add explicit prediction-interval "
        "language and an explicit Robust-vs-Fragile findings subsection "
        "in the Discussion.",
        size=11,
    )

    _add_p(
        doc,
        "I am also grateful to Reviewer 2 for endorsing publication in "
        "the current form. No revisions are made specifically in response "
        "to Reviewer 2, although Reviewer 2's suggestion of an additional "
        "moderator analysis (subject discipline) is now noted in the "
        "Future Research Directions subsection as a priority for a "
        "subsequent meta-analytic update.",
        size=11,
    )

    _add_p(
        doc,
        "Below I respond to each of Reviewer 1's six substantive concerns "
        "in turn. Each response is structured as: (a) the Reviewer's "
        "comment in summary; (b) my response in full; and (c) the "
        "specific manuscript section(s) in which the revision appears.",
        size=11, space_after=Pt(14),
    )

    # ---------------- Issue 1 ----------------
    _add_p(
        doc, "Issue 1 — Small primary pool, over-strong conclusions",
        bold=True, size=12, space_before=Pt(8),
    )
    _add_p(
        doc,
        "Reviewer 1: \"Only 10 studies contributed to the primary "
        "achievement pool, and only 9 studies contributed to several of "
        "the trait analyses. Given this small pool, the manuscript "
        "should present its conclusions as preliminary and avoid strong "
        "claims about how online learning changes the personality and "
        "achievement relationship.\"",
        italic=True, size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "Concession accepted in full. The conclusions of the manuscript "
        "have been re-framed throughout as preliminary and tentative. "
        "Specifically:",
        size=11,
    )
    for line in [
        "Abstract: the previous wording \"Findings provide preliminary "
        "support for the Personality-Achievement Saturation Hypothesis "
        "extended to technology-mediated learning\" has been replaced "
        "with \"The findings are broadly consistent with the "
        "Personality-Achievement Saturation Hypothesis (Meyer et al., "
        "2023) extended to technology-mediated learning, but the small "
        "primary pool (k = 10), substantial heterogeneity, and wide "
        "prediction intervals (which include zero for most traits) mean "
        "that this consistency does not yet constitute strong "
        "confirmation.\" The closing Abstract sentence about practical "
        "pedagogical applications has been replaced with an explicit "
        "warning against over-extrapolation on the present evidence "
        "base.",
        "Conclusion: the previous wording \"The Personality-Achievement "
        "Saturation Hypothesis is partially supported as a framework\" "
        "has been replaced with \"The findings are consistent with the "
        "Personality-Achievement Saturation Hypothesis (Meyer et al., "
        "2023) as a framework, but the present evidence base is too "
        "small to constitute strong confirmation.\" The Conclusion now "
        "opens with the explicit statement that \"the conclusions "
        "below are best read as tentative hypotheses to be tested in "
        "larger future syntheses rather than as established findings.\"",
        "Discussion: the Theoretical Implications subsection has been "
        "rewritten throughout to use language such as \"consistent "
        "with,\" \"directionally consistent,\" \"exploratory "
        "hypotheses,\" and \"descriptive patterns\" in place of "
        "previous wording such as \"contributes a substantially "
        "nuanced view,\" \"This finding aligns,\" and \"partially "
        "replicate.\"",
        "Limitations: the first limitation has been rewritten to "
        "state explicitly that \"pooled estimates are unstable, "
        "heterogeneity is poorly estimated, prediction intervals are "
        "wide, and the present synthesis must be regarded as "
        "preliminary rather than definitive.\"",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # ---------------- Issue 2 ----------------
    _add_p(
        doc, "Issue 2 — Heterogeneity and prediction intervals",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Reviewer 1: \"Heterogeneity is moderate for Conscientiousness "
        "and very high to extremely high for several other traits, and "
        "the prediction intervals cross zero. This means the pooled "
        "estimates are unstable across settings, so the paper should be "
        "much more careful when discussing general patterns and "
        "implications.\"",
        italic=True, size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "Concession accepted. The treatment of prediction intervals and "
        "heterogeneity has been substantially strengthened in three "
        "ways:",
        size=11,
    )
    for line in [
        "Explicit prediction-interval language has been added to the "
        "Abstract (\"wide prediction intervals (which include zero "
        "for most traits)\") and to the Conclusion (\"wide prediction "
        "intervals that include zero for most traits\").",
        "A new \"Distinguishing Robust from Fragile Findings\" "
        "subsection has been added to the Discussion (immediately "
        "before Strengths of the Present Review). This subsection "
        "states explicitly that \"Conscientiousness is the only trait "
        "whose 95% prediction interval does not include zero\" and "
        "classifies the Conscientiousness finding as relatively "
        "robust while explicitly marking the Agreeableness, "
        "Extraversion subgroup, Openness, and Neuroticism findings as "
        "fragile.",
        "The first Limitations paragraph now states that pooled "
        "estimates are \"unstable\" given the small k, that "
        "\"heterogeneity is poorly estimated,\" and that \"prediction "
        "intervals are wide.\" The fifth Limitations paragraph "
        "additionally notes that outcome heterogeneity \"almost "
        "certainly contributes to the wide prediction intervals "
        "reported in the Results section.\"",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # ---------------- Issue 3 ----------------
    _add_p(
        doc, "Issue 3 — Moderator overinterpretation (Extraversion × Region, "
        "Extraversion × Outcome Type)",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Reviewer 1: \"The moderator findings are also overinterpreted. "
        "The manuscript places strong emphasis on the Extraversion by "
        "Region and Extraversion by Outcome Type findings, but these "
        "subgroup analyses appear to rest on a very small number of "
        "studies. The later discussion moves from these limited "
        "subgroup findings to strong practical recommendations, "
        "especially for East Asian contexts and self-rated outcomes. "
        "These claims should be toned down substantially.\"",
        italic=True, size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "Concession accepted in full. The previous \"Practical "
        "Implications\" heading has been renamed \"Tentative Practical "
        "Implications\" and the subsection has been re-opened with an "
        "explicit caveat paragraph stating that \"the practical "
        "implications below are best read as tentative working "
        "hypotheses rather than as actionable design or pedagogical "
        "recommendations.\" The caveat paragraph specifically names the "
        "two issues Reviewer 1 raises (small k per moderator level; "
        "small k in the primary pool) and notes that prediction "
        "intervals cross zero for all traits except Conscientiousness. "
        "Within the rewritten subsection:",
        size=11,
    )
    for line in [
        "The previous wording \"the significantly negative Extraversion "
        "× Region effect (Asian r = −.131) suggests that learners who "
        "thrive on social interaction may be systematically "
        "disadvantaged\" has been rewritten as an exploratory "
        "hypothesis (\"may be hypothesised, on the present evidence, "
        "to imply ...\"), and the recommendation is now framed as a "
        "\"lead for primary-research replication\" rather than \"design\" "
        "guidance.",
        "The previous wording \"instructors should be cautious of "
        "relying on self-rated performance\" has been re-framed: the "
        "Outcome-Type-self-enhancement implication is now flagged as "
        "an \"exploratory lead\" subject to validation rather than as "
        "actionable instructor guidance.",
        "The Agreeableness-cooperation paragraph has been rewritten "
        "to state explicitly that \"the Agreeableness finding is "
        "therefore one of the most fragile in the synthesis and "
        "should not, by itself, motivate cooperative-learning design "
        "recommendations.\"",
        "The Theoretical Implications subsection's previous use of the "
        "phrase \"highly significant ... moderator\" for these "
        "subgroup contrasts has been replaced with \"large between-"
        "group differences in the subgroup contrast\" together with "
        "the explicit qualifier that \"the contrasts should therefore "
        "be read as exploratory descriptive patterns rather than as "
        "confirmatory tests of interaction.\"",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # ---------------- Issue 4 ----------------
    _add_p(
        doc, "Issue 4 — β-converted vs. direct correlations; "
        "robust-vs-fragile distinction",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Reviewer 1: \"The primary pool combines direct correlations "
        "with beta-converted estimates, and the manuscript's own "
        "sensitivity analyses suggest that some findings change when "
        "converted effects are removed. This makes some results look "
        "fragile, especially where one influential study appears to "
        "drive the pooled estimate. The discussion should distinguish "
        "more clearly between robust findings and findings that are "
        "dependent on specific analytic choices.\"",
        italic=True, size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "Concession accepted. A new Discussion subsection — \"Distinguishing "
        "Robust from Fragile Findings\" — has been added immediately "
        "before the Strengths subsection. The new subsection explicitly:",
        size=11,
    )
    for line in [
        "Classifies the Conscientiousness pooled estimate as the most "
        "robust finding (only trait with PI not crossing zero; stable "
        "under leave-one-out, β-conversion-exclusion, RoB-exclusion, "
        "and author's-own-study-exclusion sensitivity analyses; "
        "directionally and approximately matching all eight prior "
        "face-to-face benchmarks).",
        "Classifies as fragile: (i) the Agreeableness pooled estimate "
        "(driven by Yu, 2021); (ii) the null Extraversion pooled "
        "estimate (masks high subgroup heterogeneity); (iii) the "
        "Extraversion × Region and Extraversion × Outcome Type "
        "subgroup contrasts (small k per level; overlapping study "
        "sets); and (iv) the Openness and Neuroticism pooled "
        "estimates (PIs cross zero widely; estimates shift under "
        "β-conversion-exclusion).",
        "States explicitly that \"Readers should interpret discussion "
        "points that depend on these fragile findings — including the "
        "Agreeableness-cooperation hypothesis, the Asian-Extraversion-"
        "design hypothesis, and the Outcome-Type-self-enhancement "
        "hypothesis above — as exploratory rather than confirmatory.\"",
        "The first Limitations paragraph also notes explicitly that "
        "the primary pool is composed of 6 direct correlations + 4 "
        "β-converted correlations, making the role of conversion "
        "transparent up front rather than only in the sensitivity "
        "analyses section.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # ---------------- Issue 5 ----------------
    _add_p(
        doc, "Issue 5 — Search-strategy deviation, single reviewer, "
        "English-only restriction",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Reviewer 1: \"The pre-registered multi-database search was not "
        "fully carried out, the final search relied on a web-based "
        "interface and open repositories, the workflow used a single "
        "reviewer, and the study was limited to English-language "
        "sources. These limitations do not remove the value of the "
        "review, but they should reduce the confidence of the "
        "conclusions.\"",
        italic=True, size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "Concession accepted. The Limitations section has been "
        "substantially expanded with respect to all three points:",
        size=11,
    )
    for line in [
        "Search deviation. The second Limitations paragraph now states "
        "explicitly that the WebSearch deviation \"represents a real "
        "departure from the pre-registered plan,\" that "
        "\"WebSearch route does not provide controlled-vocabulary "
        "(MeSH/thesaurus) indexing or per-database hit accounting in "
        "the same way that direct database searches do,\" that a "
        "\"non-trivial number of potentially eligible non-English-"
        "language studies, dissertations, and grey-literature reports "
        "may have been missed,\" and that \"the present synthesis "
        "should not be treated as an exhaustive census of the "
        "underlying literature.\"",
        "Single reviewer. The third Limitations paragraph now states "
        "that \"intra-rater reliability cannot replicate inter-rater "
        "independence\" and that \"single-reviewer workflows are "
        "known to be susceptible to systematic biases (e.g., "
        "confirmation effects in study inclusion, drift in extraction "
        "conventions) that intra-rater checks cannot fully detect.\"",
        "English-only restriction. The fourth Limitations paragraph "
        "has been rewritten to state that the restriction \"is "
        "especially consequential for the moderator findings "
        "concerning East Asian and Middle Eastern samples ... because "
        "the subset of non-English regional work that was implicitly "
        "excluded is precisely the work most likely to clarify those "
        "moderator effects.\"",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # ---------------- Issue 6 ----------------
    _add_p(
        doc, "Issue 6 — Placeholder text, forthcoming supplements, "
        "figure/table quality",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Reviewer 1: \"The manuscript still contains placeholder text in "
        "the limitations section, some supplementary materials are "
        "marked as forthcoming, and the reporting of figures and "
        "additional analyses needs to be cleaned up. These issues weaken "
        "confidence in the care of the final manuscript.\"",
        italic=True, size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "Concession accepted with apologies for the editorial lapse.",
        size=11,
    )
    for line in [
        "Placeholder text. The placeholder \"k = [k_achievement_direct]\" "
        "in the first Limitations paragraph has been replaced with "
        "\"k = 10, with 6 contributing direct Pearson correlations "
        "and 4 contributing β-converted correlations via the Peterson "
        "and Brown (2005) transformation.\" A repository-wide grep "
        "for square-bracketed placeholders, \"TBD,\" and \"to be "
        "added\" markers was run as a regression test and returned "
        "zero hits.",
        "Forthcoming supplements. The Data Availability Statement "
        "previously described two OSF components as \"forthcoming\" "
        "(03_screening: PRISMA flow counts; 05_risk_of_bias: JBI "
        "ratings). Both components are now complete and the "
        "\"forthcoming\" parenthetical has been removed; the "
        "components are now described as \"03_screening — PRISMA "
        "flow counts and screening decisions\" and \"05_risk_of_bias "
        "— JBI risk-of-bias ratings for each included study.\"",
        "Figure and table reporting. Figures 1 (PRISMA flow), the "
        "five forest plots, and the five funnel plots, as well as "
        "Tables 1-5, have been regenerated from the canonical "
        "analysis pipeline and verified against the manuscript text "
        "by the existing T1-T7 + T9 hallucination-check suite "
        "(passed=21, failed=0, warn=0). Where Reviewer 1's checklist "
        "answer \"No\" to figure/table quality reflects render-level "
        "issues (font sizes, line weights), the regenerated figures "
        "use the canonical Matplotlib styling and have been re-exported "
        "at full resolution.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # ---------------- Closing ----------------
    _add_p(
        doc, "Closing",
        bold=True, size=12, space_before=Pt(14),
    )
    _add_p(
        doc,
        "All six of Reviewer 1's substantive concerns have been "
        "addressed by manuscript revisions. The revisions are "
        "concentrated in the Abstract, Discussion, Limitations, and "
        "Conclusion; results, methods, and the analysis pipeline are "
        "unchanged. I am grateful to Reviewer 1 for the careful and "
        "concrete critique, which has substantially improved the "
        "calibration of the manuscript's claims, and to Reviewer 2 "
        "for the positive endorsement.",
        size=11,
    )
    _add_p(
        doc,
        "I look forward to the editors' and Reviewer 1's evaluation of "
        "the revised manuscript.",
        size=11, space_before=Pt(6),
    )
    _add_p(doc, "", size=11)
    _add_p(doc, "Sincerely,", size=11)
    _add_p(doc, "", size=11)
    for line in [
        "Eisuke Tokiwa, MEng",
        "Founder, SUNBLAZE Co., Ltd., Tokyo, Japan",
        "ORCID: 0009-0009-7124-6669",
        "Email: eisuke.tokiwa@sunblaze.jp",
    ]:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


def main():
    build_response()


if __name__ == "__main__":
    main()
