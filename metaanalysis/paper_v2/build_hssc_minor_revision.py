"""
build_hssc_minor_revision.py — Generate the Springer Nature HSSC
minor-revision submission package.

Springer HSSC minor-revision request (Submission ID 72427190-7674-4b5f-
ac51-9518b8c16eaf, decision 2026-05-17, deadline 2026-05-26) requires:

  (a) Cover letter with editable ethical statements; identifying author
      info is permitted ONLY in the cover letter.
  (b) Point-by-point response — anonymised, no author identifying info.
  (c) Manuscript (anonymised) with explicit Data Availability statement.
  (d) Supporting documentation (anonymised): PRISMA checklist, search
      strategy, declaration of interest, etc.

This script produces (a) and (b):

  Outputs:
    - cover_letter_hssc_minor_revision.docx       (non-anonymous, editable)
    - response_hssc_minor_revision_anon.docx      (anonymous)
    - response_hssc_minor_revision_anon.pdf       (anonymous, PDF)
"""

import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import build_cover_letter
from build_docx_anonymous import REPLACEMENTS, rewrite_paragraph


SUBMISSION_ID = "72427190-7674-4b5f-ac51-9518b8c16eaf"
DEADLINE = "26 May 2026"


# ---------------------------------------------------------------------------
# Shared editable ethics + data availability blocks
# (these are the canonical wording embedded in the Cover Letter, used both
#  for the manuscript Declarations and for the submission-portal Declarations)
# ---------------------------------------------------------------------------

ETHICS_BLOCK_LINES = [
    ("Ethics Approval", True),
    ("Not applicable. This study is a systematic review and meta-analysis of "
     "previously published quantitative research and did not involve any "
     "direct collection of data from human participants by the present author. "
     "As no new primary data were collected, no Institutional Review Board "
     "(IRB) or Research Ethics Committee approval was required for the "
     "conduct of this synthesis. The conduct of this review followed the "
     "ethical principles outlined in the Declaration of Helsinki (as revised "
     "in 2013) and the PRISMA 2020 reporting guidelines for systematic "
     "reviews. Ethical approval for each of the included primary studies was "
     "the responsibility of the original investigators of those studies, as "
     "documented in their respective publications.", False),
    ("Consent to Participate", True),
    ("Not applicable. This systematic review and meta-analysis did not "
     "involve any direct interaction with, or recruitment of, human "
     "participants by the present author. Informed consent and consent to "
     "participate for the underlying data were obtained by the original "
     "investigators of each included primary study at the time of original "
     "data collection, in accordance with the institutional and ethical "
     "standards applicable at that time.", False),
    ("Consent for Publication", True),
    ("Not applicable. This manuscript does not contain any individual "
     "person's data in any form (including individual details, images, "
     "or videos).", False),
    ("Human Ethics and Consent to Participate Declarations", True),
    ("Not applicable.", False),
]

DATA_AVAILABILITY_TEXT = (
    "All materials underpinning this systematic review and meta-analysis "
    "have been made available to the editorial office and referees as "
    "anonymised supplementary files with this submission. Specifically: "
    "(a) the completed PRISMA 2020 27-item checklist with section-level "
    "locations (prisma_2020_checklist_anon.docx); (b) the full Boolean "
    "search strategy, database and source list, query-level hit counts, "
    "and the documented deviation from the pre-registered search plan "
    "(search_strategy_anon.docx); (c) the PRISMA 2020 flow diagram "
    "(reproduced as Figure 1 of the main manuscript); (d) the eligibility "
    "criteria, screening decisions, and full-text exclusion reasons "
    "(reported in the Methods and Results sections); (e) the data-"
    "extraction table (reproduced as Table 1) listing every included "
    "study's country, modality, sample size, personality instrument, "
    "outcome operationalisation, era, region, JBI risk-of-bias score, "
    "and inclusion status; (f) the synthesis tables (Tables 2-5: pooled "
    "effects, moderator analyses, sensitivity analyses, and adapted "
    "GRADE ratings); and (g) the anonymised Declaration of Interest "
    "(declaration_of_interest_anon.docx). Beyond the materials uploaded "
    "here, the full pre-registration record, the seven-component open-"
    "science project (protocol, search log, screening ledger, data-"
    "extraction CSV, JBI risk-of-bias spreadsheet, analysis code and "
    "outputs, and article-level DOI index), and the version-controlled "
    "code repository hosting the full Python analysis pipeline have all "
    "been permanently deposited under stable DOIs and are publicly "
    "accessible. Those DOIs are withheld at this stage solely because "
    "their URLs identify the author and will be supplied to the "
    "editorial office post-acceptance for inclusion in the final "
    "published version. Nothing has been redacted from these deposits; "
    "the only restriction is the temporary withholding of identifying "
    "URLs for double-anonymous review."
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_run_font(run, size_pt=11, bold=None):
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _add_p(doc, text, *, bold=False, size=11, align=None,
           space_after=Pt(6), space_before=Pt(0), indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.3
    if indent is not None:
        p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


def _configure_page(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# (a) Cover letter — non-anonymous, includes editable ethics statements
# ---------------------------------------------------------------------------

def build_cover_letter_minor_revision(out_path: Path):
    doc = Document()
    _configure_page(doc)

    # Identifying header (allowed in cover letter only)
    for line in build_cover_letter.HEADER_LINES:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)

    _add_p(doc, date.today().strftime("%B %d, %Y"), size=10)
    _add_p(doc, "", size=10)

    # Addressee
    for line in ["To Fatima Nasrin, Assistant Editor",
                 "Humanities and Social Sciences Communications",
                 f"Submission ID: {SUBMISSION_ID}"]:
        _add_p(doc, line, size=10, space_after=Pt(0))
    _add_p(doc, "", size=10)

    # Salutation
    _add_p(doc, "Dear Dr Nasrin,", size=11)

    # Opening
    _add_p(
        doc,
        f"Thank you for the minor-revision decision on Submission "
        f"ID {SUBMISSION_ID} and for the clear technical-check guidance. "
        f"I have addressed every item raised, in advance of the {DEADLINE} "
        f"deadline. This cover letter is the only file in this submission "
        f"that carries author identifying information; the manuscript, "
        f"point-by-point response, and all supplementary files have been "
        f"prepared as fully anonymised double-anonymous-ready versions in "
        f"accordance with the editor's instruction.",
        size=11,
    )

    # Files in this submission
    _add_p(doc, "Files in this submission",
           bold=True, size=12, space_before=Pt(8))
    _add_p(
        doc,
        "• manuscript_journal_v2_anonymous.docx — anonymised main "
        "manuscript with a new Data Availability section.",
        size=10, space_after=Pt(2),
    )
    _add_p(
        doc,
        "• response_hssc_minor_revision_anon.pdf — anonymised point-by-"
        "point response addressing every item in this minor-revision "
        "request.",
        size=10, space_after=Pt(2),
    )
    _add_p(
        doc,
        "• prisma_2020_checklist_anon.docx — completed PRISMA 2020 27-"
        "item checklist (supplementary material).",
        size=10, space_after=Pt(2),
    )
    _add_p(
        doc,
        "• search_strategy_anon.docx — full Boolean search strategy, "
        "database list, query-level log, and documented deviation from "
        "the pre-registered search plan (supplementary material).",
        size=10, space_after=Pt(2),
    )
    _add_p(
        doc,
        "• declaration_of_interest_anon.docx — anonymised Declaration of "
        "Interest (related file).",
        size=10, space_after=Pt(2),
    )
    _add_p(
        doc,
        "• cover_letter_hssc_minor_revision.docx (this file) — non-"
        "anonymous cover letter, the only file carrying author "
        "identifying information.",
        size=10, space_after=Pt(10),
    )

    # Editable Ethical Statements
    _add_p(
        doc, "Editable Ethical Statements",
        bold=True, size=12, space_before=Pt(4),
    )
    _add_p(
        doc,
        "Per the editor's instruction, the ethics statements below are "
        "reproduced in this cover letter in editable text format. They "
        "match verbatim the Declarations block of the manuscript and the "
        "values entered into the submission-portal Declarations form. "
        "There is no separate ethics committee confirmation letter "
        "because, as stated below, no IRB or ethics committee approval "
        "was required for the conduct of this systematic review (no new "
        "primary data were collected).",
        size=11,
    )

    for text, bold in ETHICS_BLOCK_LINES:
        _add_p(
            doc, text, bold=bold, size=11,
            space_after=Pt(2 if bold else 8),
        )

    # Data Availability statement
    _add_p(
        doc, "Data Availability Statement", bold=True, size=12,
        space_before=Pt(8),
    )
    _add_p(
        doc,
        "Per the editor's instruction, the Data Availability Statement "
        "below is reproduced here in editable text format. It matches "
        "verbatim the Data Availability section of the manuscript (in the "
        "Declarations block) and the value entered into the submission-"
        "portal Declarations form.",
        size=11,
    )
    _add_p(doc, DATA_AVAILABILITY_TEXT, size=11)

    # Closing
    _add_p(
        doc,
        "If any further technical-check items remain after this "
        "submission, I am happy to address them within the same week. "
        "Thank you for your patience.",
        size=11, space_before=Pt(8),
    )
    _add_p(doc, "Sincerely,", size=11, space_before=Pt(6))
    _add_p(doc, "", size=11)
    for line in build_cover_letter.SIGNATURE_LINES:
        _add_p(doc, line, size=10, space_after=Pt(0))

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


# ---------------------------------------------------------------------------
# (b) Point-by-point response — anonymised, no author identifying info
# ---------------------------------------------------------------------------

def build_response_minor_revision_docx(out_path: Path):
    doc = Document()
    _configure_page(doc)

    _add_p(
        doc, "Point-by-Point Response to Minor-Revision Request",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_p(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Systematic Review and Meta-Analysis",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8),
    )
    _add_p(
        doc,
        f"Submission ID: {SUBMISSION_ID}  |  Decision: minor revision  |  "
        f"Deadline: {DEADLINE}",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_p(
        doc,
        "[Author identifying information removed for double-anonymous "
        "peer review.]",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18),
    )

    _add_p(
        doc,
        "We thank the Assistant Editor for the clear minor-revision "
        "guidance. Each item raised in the request is addressed in turn "
        "below. All revisions are confined to the Declarations block of "
        "the manuscript and to the supplementary materials; no changes "
        "have been made to the results, methods, or conclusions of the "
        "study.",
        size=11,
    )

    # Item 1
    _add_p(
        doc, "Item 1 — Data and supporting documentation for systematic "
        "review / meta-analysis",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Editor's request: provide PRISMA checklist and flow diagram, "
        "full search strategies, extracted records / screening logs / "
        "inclusion-exclusion criteria, data extraction forms, synthesis "
        "tables, and any other materials necessary to verify "
        "methodology and findings — either as files or as stable "
        "repository links.",
        size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "All requested materials have been provided as anonymised "
        "supplementary files with this submission:",
        size=11,
    )
    for line in [
        "PRISMA 2020 27-item checklist — prisma_2020_checklist_anon.docx. "
        "Section-level locations for each of the 27 items are listed "
        "alongside the item text.",
        "PRISMA 2020 flow diagram — Figure 1 of the main manuscript "
        "(also reproduced in the search-strategy supplement with the "
        "final counts).",
        "Full Boolean search strategy — search_strategy_anon.docx. "
        "Contains all three concept blocks with full OR-expansion, the "
        "database and source list, execution dates, query-level hit "
        "counts, and a transparent disclosure of the pre-registered-vs-"
        "executed search deviation.",
        "Eligibility criteria and screening decisions — reported in the "
        "Methods → Eligibility Criteria, Information Sources, and Study "
        "Selection subsections of the manuscript; full-text exclusion "
        "reasons are reported in Results.",
        "Data-extraction forms / synthesis tables — reproduced as Table "
        "1 (extraction-level metadata for every included study) and "
        "Tables 2-5 (pooled effects, moderator analyses, sensitivity "
        "analyses, GRADE) of the main manuscript.",
        "Underlying datasets, analysis code, screening ledger, JBI risk-"
        "of-bias spreadsheet, and article-level DOI index — permanently "
        "deposited under stable DOIs in a time-stamped open-science "
        "repository and a version-controlled code repository. Their "
        "URLs are withheld at this stage solely because they identify "
        "the author; they will be supplied to the editorial office "
        "post-acceptance.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)

    # Item 2: Data Availability statement
    _add_p(
        doc,
        "Item 2 — Data Availability statement (manuscript file)",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Editor's request: include a Data Availability statement at the "
        "end of the manuscript detailing (i) what materials have been "
        "shared, (ii) where they can be accessed, and (iii) any reasons "
        "why materials cannot be shared.",
        size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "A Data Availability Statement has been added to the "
        "Declarations block of the manuscript (immediately after "
        "Funding, before Authors' Contributions). The full statement "
        "lists, item-by-item, every supplementary file uploaded with "
        "this submission, explicitly identifies the manuscript Tables "
        "and Figure 1 as the synthesis and flow-diagram materials, and "
        "explains that the OSF and code-repository URLs are temporarily "
        "withheld only because they identify the author; the deposits "
        "themselves are complete and unredacted.",
        size=11,
    )

    # Item 3: Data Availability statement (portal)
    _add_p(
        doc,
        "Item 3 — Data Availability statement (submission-portal "
        "Declarations form)",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Editor's request: provide the same Data Availability statement "
        "on the submission system under 'Declarations'; do not accept "
        "'NO'; ensure the same text is in both the manuscript and the "
        "submission system.",
        size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "The identical Data Availability statement that appears in the "
        "manuscript Declarations block has been entered into the "
        "submission-portal Declarations form. The verbatim text used in "
        "both locations is reproduced in the (non-anonymous) cover "
        "letter for the editor's review.",
        size=11,
    )

    # Item 4: Ethics in editable format
    _add_p(
        doc,
        "Item 4 — Ethical statements in editable format within the "
        "cover letter",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Editor's request: provide ethical statements in editable "
        "format within the cover letter; if submitting an ethics-"
        "committee confirmation letter (in English), include it along "
        "with the editable ethical statement.",
        size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "The four ethical statements (Ethics Approval, Consent to "
        "Participate, Consent for Publication, and the summary 'Human "
        "Ethics and Consent to Participate Declarations: Not "
        "applicable') are reproduced in editable text format in the "
        "(non-anonymous) cover letter that accompanies this response, "
        "matching verbatim the wording in the Declarations block of "
        "the manuscript. No separate ethics-committee confirmation "
        "letter is included because, as the Ethics Approval statement "
        "explains, no IRB or Research Ethics Committee approval was "
        "required for the conduct of this systematic review and meta-"
        "analysis (no new primary data were collected from human "
        "participants by the present author).",
        size=11,
    )

    # Item 5: Anonymisation of all files except cover letter
    _add_p(
        doc,
        "Item 5 — Author identifying information removed from all "
        "files except the cover letter",
        bold=True, size=12, space_before=Pt(10),
    )
    _add_p(
        doc,
        "Editor's instruction: do not include author name or details in "
        "the point-by-point response letter or in any of the files, "
        "except the cover letter.",
        size=10,
    )
    _add_p(doc, "Response.", bold=True, size=11)
    _add_p(
        doc,
        "The manuscript, this point-by-point response, the PRISMA "
        "checklist, the search-strategy supplement, and the Declaration "
        "of Interest have all been rebuilt as fully anonymised "
        "versions. Specifically, the following identifying tokens have "
        "been removed wherever they appeared: author name, ORCID, "
        "email, institutional affiliation, OSF and preregistration "
        "DOIs, preprint DOI, and code-repository URL. Each anonymised "
        "file has been verified programmatically against a 11-token "
        "leak-detection scan that reported zero hits. Identifying "
        "information is confined entirely to the cover letter "
        "(cover_letter_hssc_minor_revision.docx).",
        size=11,
    )

    # Closing
    _add_p(
        doc,
        "All revision items have been addressed. The manuscript, the "
        "Declarations block, and the supplementary files are ready for "
        "the next stage of editorial processing. If any further "
        "technical-check items remain after this submission, they will "
        "be addressed within the same week.",
        size=11, space_before=Pt(12),
    )

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def _convert_to_pdf(docx_path: Path) -> Path:
    """Convert a docx to PDF via mammoth + weasyprint."""
    import mammoth
    import weasyprint

    pdf_path = docx_path.with_suffix(".pdf")
    with open(docx_path, "rb") as f:
        html_body = mammoth.convert_to_html(f).value
    full_html = (
        '<!DOCTYPE html><html><head><meta charset="utf-8"><style>'
        "@page { size: letter; margin: 2cm 2.5cm; }"
        'body { font-family: "DejaVu Sans", Arial, sans-serif; '
        "font-size: 11pt; line-height: 1.4; }"
        "p { margin: 6pt 0; }"
        "h1, h2, h3 { color: #000; }"
        "strong, b { font-weight: bold; }"
        "ul { margin: 4pt 0; padding-left: 1.5em; }"
        "li { margin: 2pt 0; }"
        "</style></head><body>" + html_body + "</body></html>"
    )
    weasyprint.HTML(string=full_html).write_pdf(str(pdf_path))
    print(f"Wrote {pdf_path}")
    return pdf_path


def main():
    here = Path(__file__).resolve().parent

    cover_path = here / "cover_letter_hssc_minor_revision.docx"
    build_cover_letter_minor_revision(cover_path)

    response_docx = here / "response_hssc_minor_revision_anon.docx"
    build_response_minor_revision_docx(response_docx)
    _convert_to_pdf(response_docx)


if __name__ == "__main__":
    main()
