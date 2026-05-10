"""Generate Heliyon-required supplementary documents.

Heliyon submission additionally requires:
  1. PRISMA 2020 Checklist  (Page et al., 2021; BMJ 372:n71)
  2. Declaration of Interest Statement (Elsevier standard form)

Both are produced as standalone .docx files in paper_v2/, alongside the
manuscript, cover letter, and split files.

The PRISMA 2020 checklist is the 27-item main checklist. For each item, the
"Location where item is reported" column points to the section/subsection of
manuscript_journal_v2.docx (page numbers are deliberately avoided because
they shift between the journal-formatted and reviewer-formatted PDFs).
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
PRISMA_OUT = HERE / "prisma_2020_checklist.docx"
DOI_OUT = HERE / "declaration_of_interest.docx"


# ---------------------------------------------------------------------------
# Shared formatting helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, font_name="Calibri", size_pt=10, bold=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_para_font(paragraph, font_name="Calibri", size_pt=10, bold=None):
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold)


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _configure_page(doc, landscape=False):
    s = doc.sections[0]
    if landscape:
        s.page_width, s.page_height = s.page_height, s.page_width
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)


def _add_paragraph(doc, text, *, bold=False, size=10, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


# ---------------------------------------------------------------------------
# PRISMA 2020 Checklist
# ---------------------------------------------------------------------------
# Each tuple: (section, item_number, item_text, location_reported)
# Sections follow Page, M.J., et al. (2021). The PRISMA 2020 statement: an
# updated guideline for reporting systematic reviews. BMJ, 372, n71.
PRISMA_ITEMS = [
    # ---- TITLE ------------------------------------------------------------
    ("TITLE", "", "", ""),
    ("Title", "1",
     "Identify the report as a systematic review.",
     "Title page — title explicitly identifies the report as “A Systematic "
     "Review and Meta-Analysis.”"),

    # ---- ABSTRACT ---------------------------------------------------------
    ("ABSTRACT", "", "", ""),
    ("Abstract", "2",
     "See the PRISMA 2020 for Abstracts checklist.",
     "Abstract section — structured abstract reports background, methods "
     "(PRISMA 2020 compliance, REML/HKSJ random-effects synthesis, JBI risk "
     "of bias, GRADE), results (k = 25 qualitative / 10 quantitative; pooled "
     "N = 3,384; trait-by-trait pooled r and 95% CI; two pre-registered "
     "moderator effects), and conclusions. Pre-registration DOI is referenced "
     "in the Author Note and Declarations."),

    # ---- INTRODUCTION ------------------------------------------------------
    ("INTRODUCTION", "", "", ""),
    ("Rationale", "3",
     "Describe the rationale for the review in the context of existing "
     "knowledge.",
     "Introduction → “Personality and Academic Achievement in Face-to-Face "
     "Contexts” and “The Shift to Online Learning Environments” subsections. "
     "Eight prior meta-analyses (Poropat 2009 through Chen et al. 2025) are "
     "summarized; the absence of an online-modality moderator is the stated "
     "gap."),
    ("Objectives", "4",
     "Provide an explicit statement of the objective(s) or question(s) the "
     "review addresses.",
     "Introduction → “The Present Review” subsection — five pre-registered "
     "research questions (RQ1–RQ5) and the PICOS framework are stated "
     "explicitly."),

    # ---- METHODS ----------------------------------------------------------
    ("METHODS", "", "", ""),
    ("Eligibility criteria", "5",
     "Specify the inclusion and exclusion criteria for the review and how "
     "studies were grouped for the syntheses.",
     "Methods → “Eligibility Criteria” subsection — full PICOS eligibility "
     "rules, including population, exposure (Big Five / HEXACO instruments), "
     "comparator, primary and secondary outcomes, study design, and "
     "publication-type restrictions. Grouping for synthesis is by Big Five "
     "trait."),
    ("Information sources", "6",
     "Specify all databases, registers, websites, organisations, reference "
     "lists and other sources searched or consulted to identify studies. "
     "Specify the date when each source was last searched or consulted.",
     "Methods → “Information Sources and Search Strategy” subsection — six "
     "pre-registered databases (PubMed/MEDLINE, PsycINFO, ERIC, Web of "
     "Science, Scopus, ProQuest), Google Scholar, and forward/backward "
     "citation snowballing. The protocol-level access deviation and "
     "mitigation are disclosed; full search log with execution dates is "
     "deposited at https://doi.org/10.17605/OSF.IO/UVJDY."),
    ("Search strategy", "7",
     "Present the full search strategies for all databases, registers and "
     "websites, including any filters and limits used.",
     "Methods → “Information Sources and Search Strategy” subsection — the "
     "three-concept Boolean strategy (personality × online learning × "
     "academic outcome) is reproduced in full. The complete query syntax, "
     "filters, and per-source hit counts are deposited as search_log.md on "
     "OSF (02_search component, https://doi.org/10.17605/OSF.IO/UVJDY)."),
    ("Selection process", "8",
     "Specify the methods used to decide whether a study met the inclusion "
     "criteria of the review, including how many reviewers screened each "
     "record and each report retrieved, whether they worked independently, "
     "and if applicable, details of automation tools used in the process.",
     "Methods → “Study Selection” subsection — single reviewer (ET) under a "
     "pre-specified intra-rater reliability protocol with 10% (title/abstract) "
     "and 20% (full-text) re-screening after a ≥7-day wash-out, target "
     "Cohen's κ ≥ 0.80; Zotero used for de-duplication. The single-reviewer "
     "deviation is transparently disclosed."),
    ("Data collection process", "9",
     "Specify the methods used to collect data from reports, including how "
     "many reviewers collected data from each report, whether they worked "
     "independently, any processes for obtaining or confirming data from "
     "study investigators, and if applicable, details of automation tools "
     "used in the process.",
     "Methods → “Data Extraction” subsection — single reviewer extraction "
     "with 10% double-extraction and target ICC(2,1) ≥ 0.90 / κ ≥ 0.80. "
     "Author-contact protocol (up to two attempts, ≥2-week interval) for "
     "missing statistics. No automation tools."),
    ("Data items", "10a",
     "List and define all outcomes for which data were sought. Specify "
     "whether all results that were compatible with each outcome domain in "
     "each study were sought (e.g. for all measures, time points, analyses), "
     "and if not, the methods used to decide which results to collect.",
     "Methods → “Eligibility Criteria” and “Data Extraction” subsections — "
     "primary outcome (academic achievement: GPA, course grade, exam, "
     "composite); secondary outcomes (satisfaction, engagement, learning "
     "behavior, dropout/persistence). For multi-result studies, the most "
     "inclusive effect size was retained in the primary pool, with "
     "sensitivity analysis using RVE / three-level models."),
    ("Data items", "10b",
     "List and define all other variables for which data were sought (e.g. "
     "participant and intervention characteristics, funding sources). "
     "Describe any assumptions made about any missing or unclear information.",
     "Methods → “Data Extraction” subsection — full extraction-form "
     "specification (sample, learning context, personality measurement, "
     "outcome measurement, effect-size statistics). Missing-information "
     "policy: corresponding-author contact, conversion rules where possible, "
     "otherwise exclusion with logged reason."),
    ("Study risk of bias assessment", "11",
     "Specify the methods used to assess risk of bias in the included "
     "studies, including details of the tool(s) used, how many reviewers "
     "assessed each study and whether they worked independently, and if "
     "applicable, details of automation tools used in the process.",
     "Methods → “Risk of Bias Assessment” subsection — Joanna Briggs "
     "Institute (JBI) Critical Appraisal Checklist for Analytical Cross-"
     "Sectional Studies, 8 items, single reviewer with 20% re-assessment, "
     "target κ ≥ 0.80. No automation."),
    ("Effect measures", "12",
     "Specify for each outcome the effect measure(s) (e.g. risk ratio, mean "
     "difference) used in the synthesis or presentation of results.",
     "Methods → “Effect Size Metric and Transformations” subsection — "
     "Pearson r as the primary metric; analyses on the Fisher z scale with "
     "back-transformation to r for reporting. Conversion rules for β, d, "
     "η², F, and t are pre-specified."),
    ("Synthesis methods", "13a",
     "Describe the processes used to decide which studies were eligible for "
     "each synthesis (e.g. tabulating the study intervention characteristics "
     "and comparing against the planned groups for each synthesis).",
     "Methods → “Eligibility Criteria,” “Study Selection,” and “Statistical "
     "Analysis and Synthesis Strategy” subsections; Table 1 catalogues all "
     "31 records with eligibility, retention, and trait-by-trait "
     "contribution."),
    ("Synthesis methods", "13b",
     "Describe any methods required to prepare the data for presentation or "
     "synthesis, such as handling of missing summary statistics, or data "
     "conversions.",
     "Methods → “Effect Size Metric and Transformations” subsection — "
     "Fisher z transformation, conversion of β / d / η² / F / t to r, "
     "Neuroticism sign alignment for Emotional Stability reports, and the "
     "HEXACO-to-Big-Five crosswalk (Ashton & Lee, 2007) with sensitivity "
     "variants."),
    ("Synthesis methods", "13c",
     "Describe any methods used to tabulate or visually display results of "
     "individual studies and syntheses.",
     "Methods → “Statistical Analysis and Synthesis Strategy” and "
     "“Publication Bias” subsections — forest plots (Figures 2–6), funnel "
     "plots (Figures 7–11), and tables (Tables 1–5). Visualizations produced "
     "in ggplot2."),
    ("Synthesis methods", "13d",
     "Describe any methods used to synthesize results and provide a "
     "rationale for the choice(s). If meta-analysis was performed, describe "
     "the model(s), method(s) to identify the presence and extent of "
     "statistical heterogeneity, and software package(s) used.",
     "Methods → “Statistical Analysis and Synthesis Strategy” subsection — "
     "random-effects meta-analysis with REML τ² estimation, Hartung-Knapp-"
     "Sidik-Jonkman 95% CI adjustment, Cochran's Q, I², τ², τ, 95% prediction "
     "intervals. Software: R ≥ 4.3.0 with metafor, clubSandwich, dmetar, "
     "ggplot2 (deposited renv.lock and sessionInfo on OSF)."),
    ("Synthesis methods", "13e",
     "Describe any methods used to explore possible causes of heterogeneity "
     "among study results (e.g. subgroup analysis, meta-regression).",
     "Methods → “Moderator and Sensitivity Analyses” subsection — nine "
     "pre-specified moderators (5 categorical: modality, education level, "
     "region, era, outcome type, instrument; 3 continuous: publication year, "
     "log-N, RoB score), tested via mixed-effects meta-regression with "
     "Wald-type QM tests; Holm-Bonferroni correction within trait. Pre-"
     "registered restriction to k ≥ 10 per level for quantitative testing, "
     "with the remaining moderators reported narratively (disclosed as a "
     "deviation)."),
    ("Synthesis methods", "13f",
     "Describe any sensitivity analyses conducted to assess robustness of "
     "the synthesized results.",
     "Methods → “Moderator and Sensitivity Analyses” subsection — seven "
     "pre-specified sensitivity analyses (RoB < 5 exclusion; author's-own-"
     "study exclusion; converted-effect-size exclusion; small-sample "
     "exclusion; HEXACO-mapping variants; leave-one-out with Cook's "
     "distance; DerSimonian-Laird vs. REML). Results in Results → "
     "“Sensitivity and Robustness” and Table 4."),
    ("Reporting bias assessment", "14",
     "Describe any methods used to assess risk of bias due to missing "
     "results in a synthesis (arising from reporting biases).",
     "Methods → “Publication Bias” subsection — funnel plots, Egger's "
     "regression, Peters' regression, Duval & Tweedie trim-and-fill (as "
     "sensitivity check, not point-estimate correction), and Simonsohn et "
     "al. p-curve. Grey-literature inclusion partially mitigates file-drawer "
     "effects."),
    ("Certainty assessment", "15",
     "Describe any methods used to assess certainty (or confidence) in the "
     "body of evidence for an outcome.",
     "Methods → “Confidence in Cumulative Evidence (GRADE)” subsection — "
     "GRADE adaptation for observational correlational syntheses (Schünemann "
     "et al., 2019); five downgrade domains (RoB, inconsistency, "
     "indirectness, imprecision, publication bias) and two upgrade "
     "considerations (large magnitude, dose-response). Minimum effect of "
     "interest |r| = .10."),

    # ---- RESULTS ----------------------------------------------------------
    ("RESULTS", "", "", ""),
    ("Study selection", "16a",
     "Describe the results of the search and selection process, from the "
     "number of records identified in the search to the number of studies "
     "included in the review, ideally using a flow diagram.",
     "Results → “Study Selection” subsection and Figure 1 (PRISMA 2020 flow "
     "diagram). Counts: 31 catalogued at full-text → 25 retained for "
     "qualitative synthesis → 10 contributing to the primary quantitative "
     "achievement pool."),
    ("Study selection", "16b",
     "Cite studies that might appear to meet the inclusion criteria, but "
     "which were excluded, and explain why they were excluded.",
     "Results → “Study Selection” subsection and Table 1 — six post-"
     "eligibility exclusions are individually disclosed with reasons (A-09, "
     "A-10, A-16: face-to-face on secondary review; A-24: effect size not "
     "extractable; A-05: sample overlap with A-04; A-27: PDF unavailable). "
     "Five title-stage exclusions for non-Big-Five frameworks (MBTI, "
     "Proactive Personality, TAM, TUE) are summarized in the same "
     "subsection."),
    ("Study characteristics", "17",
     "Cite each included study and present its characteristics.",
     "Results → “Characteristics of Included Studies” subsection and "
     "Table 1 — full per-study characteristics (sample, country, modality, "
     "instrument, outcome, N, year, era)."),
    ("Risk of bias in studies", "18",
     "Present assessments of risk of bias for each included study.",
     "Results → “Risk of Bias Across Included Studies” subsection. Per-"
     "study JBI 8-item ratings are deposited as Table S2 (supplementary) on "
     "OSF (05_risk_of_bias component, "
     "https://doi.org/10.17605/OSF.IO/3XYNE)."),
    ("Results of individual studies", "19",
     "For all outcomes, present, for each study: (a) summary statistics for "
     "each group (where appropriate) and (b) an effect estimate and its "
     "precision (e.g. confidence/credible interval), ideally using "
     "structured tables or plots.",
     "Forest plots (Figures 2–6) display per-study r and 95% CI for each of "
     "the five Big Five traits. Per-study extracted statistics are deposited "
     "in data_extraction.csv on OSF (04_extraction component, "
     "https://doi.org/10.17605/OSF.IO/5UW98)."),
    ("Results of syntheses", "20a",
     "For each synthesis, briefly summarise the characteristics and risk of "
     "bias among contributing studies.",
     "Results → “Characteristics of Included Studies,” “Risk of Bias Across "
     "Included Studies,” and the trait-specific subsections of “Primary "
     "Pooled Effects.”"),
    ("Results of syntheses", "20b",
     "Present results of all statistical syntheses conducted. If meta-"
     "analysis was done, present for each the summary estimate and its "
     "precision (e.g. confidence/credible interval) and measures of "
     "statistical heterogeneity. If comparing groups, describe the direction "
     "of the effect.",
     "Results → “Primary Pooled Effects” subsection and Table 2 — pooled "
     "r, 95% HKSJ CI, prediction interval, Q, I², τ², τ for each Big Five "
     "trait."),
    ("Results of syntheses", "20c",
     "Present results of all investigations of possible causes of "
     "heterogeneity among study results.",
     "Results → “Moderator Analyses” subsection and Table 3 — Q_between, "
     "QM, R², and uncorrected/Holm-Bonferroni-corrected p-values for each "
     "tested moderator within each trait."),
    ("Results of syntheses", "20d",
     "Present results of all sensitivity analyses conducted to assess the "
     "robustness of the synthesized results.",
     "Results → “Sensitivity and Robustness” subsection and Table 4 — full "
     "comparison of pooled r and 95% CI under each of the seven sensitivity "
     "specifications (RoB < 5 exclusion, author's-own-study exclusion, "
     "converted-effect-size exclusion, small-sample exclusion, HEXACO-"
     "mapping variants, leave-one-out, alternative τ² estimator)."),
    ("Reporting biases", "21",
     "Present assessments of risk of bias due to missing results (arising "
     "from reporting biases) for each synthesis assessed.",
     "Results → “Publication Bias” subsection — Egger's regression, Peters' "
     "regression, trim-and-fill, and p-curve per trait. Funnel plots "
     "(Figures 7–11)."),
    ("Certainty of evidence", "22",
     "Present assessments of certainty (or confidence) in the body of "
     "evidence for each outcome assessed.",
     "Results → “GRADE Confidence Ratings” subsection and Table 5 — per-"
     "trait GRADE rating (Moderate for Conscientiousness and Extraversion; "
     "Low for Openness, Agreeableness, Neuroticism) with downgrade "
     "rationales."),

    # ---- DISCUSSION -------------------------------------------------------
    ("DISCUSSION", "", "", ""),
    ("Discussion", "23a",
     "Provide a general interpretation of the results in the context of "
     "other evidence.",
     "Discussion → “Principal Findings in Context” subsection — "
     "interpretation against the eight prior face-to-face meta-analyses and "
     "the Personality-Achievement Saturation Hypothesis (PASH)."),
    ("Discussion", "23b",
     "Discuss any limitations of the evidence included in the review.",
     "Discussion → “Limitations of the Evidence Base” subsection."),
    ("Discussion", "23c",
     "Discuss any limitations of the review processes used.",
     "Discussion → “Limitations of the Review Process” subsection — "
     "single-reviewer workflow, search-database access deviation, narrowed "
     "moderator set, reliance on grey literature, and language restriction."),
    ("Discussion", "23d",
     "Discuss implications of the results for practice, policy, and future "
     "research.",
     "Discussion → “Implications for Practice and Future Research” "
     "subsection and Conclusion."),

    # ---- OTHER INFORMATION -----------------------------------------------
    ("OTHER INFORMATION", "", "", ""),
    ("Registration and protocol", "24a",
     "Provide registration information for the review, including the "
     "register name and registration number, or state that the review was "
     "not registered.",
     "Declarations → “Pre-registration”; Author Note; and Methods → "
     "“Pre-registration and Reporting Standards.” OSF Registries, "
     "https://osf.io/e5w47/, DOI: 10.17605/OSF.IO/E5W47, registered "
     "23 April 2026."),
    ("Registration and protocol", "24b",
     "Indicate where the review protocol can be accessed, or state that a "
     "protocol was not prepared.",
     "Declarations → “Pre-registration” and “Availability of data and "
     "material”; Methods → “Pre-registration and Reporting Standards.” "
     "Protocol deposited at https://doi.org/10.17605/OSF.IO/7FRGH "
     "(01_protocol component)."),
    ("Registration and protocol", "24c",
     "Describe and explain any amendments to information provided at "
     "registration or in the protocol.",
     "Methods → “Deviations from the Pre-registered Protocol” subsection — "
     "three deviations are itemized (search-database access, "
     "moderator quantitative-vs-narrative reduction with k-per-level "
     "documentation, post-registration addition of Chen et al. 2025 to the "
     "benchmark corpus only, not the synthesis pool)."),
    ("Support", "25",
     "Describe sources of financial or non-financial support for the "
     "review, and the role of the funders or sponsors in the review.",
     "Declarations → “Funding”; cover letter. No external funding."),
    ("Competing interests", "26",
     "Declare any competing interests of review authors.",
     "Declarations → “Conflict of Interest”; Declaration of Interest "
     "Statement (separate file); cover letter. One author-own primary "
     "study (Tokiwa, 2025) is potentially eligible; addressed by a pre-"
     "specified sensitivity analysis (|Δr| < .001 across traits)."),
    ("Availability of data, code and other materials", "27",
     "Report which of the following are publicly available and where they "
     "can be found: template data collection forms; data extracted from "
     "included studies; data used for all analyses; analytic code; any "
     "other materials used in the review.",
     "Declarations → “Availability of data and material”; Methods → "
     "“Software and Reproducibility.” All materials are deposited on the "
     "OSF project (https://doi.org/10.17605/OSF.IO/79M5J) across seven "
     "DOI-tagged components: 01_protocol, 02_search, 03_screening, "
     "04_extraction, 05_risk_of_bias, 06_analysis, 07_pdf_index. "
     "Version-controlled mirror at https://github.com/etoki/paper."),
]


def build_prisma_checklist():
    doc = Document()
    _configure_page(doc, landscape=True)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("PRISMA 2020 Checklist")
    _set_run_font(run, size_pt=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    run = sub.add_run(
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Systematic Review and Meta-Analysis"
    )
    _set_run_font(run, size_pt=11, bold=False)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run(
        "Adapted from Page, M. J., McKenzie, J. E., Bossuyt, P. M., Boutron, "
        "I., Hoffmann, T. C., Mulrow, C. D., et al. (2021). The PRISMA 2020 "
        "statement: an updated guideline for reporting systematic reviews. "
        "BMJ, 372, n71. doi: 10.1136/bmj.n71. "
        "Companion PRISMA 2020 Abstract checklist content is integrated into "
        "the structured Abstract of the manuscript."
    )
    _set_run_font(run, size_pt=9, bold=False)

    # ---- Table ----
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    # Column widths (sum ~ 26 cm for landscape A4 minus margins ~25 cm)
    widths = [Cm(4.0), Cm(1.4), Cm(9.0), Cm(11.0)]
    for col, w in zip(table.columns, widths):
        for cell in col.cells:
            cell.width = w

    headers = ["Section / Topic", "#", "Checklist item",
               "Location where item is reported"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = widths[i]
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        _set_run_font(run, size_pt=10, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(hdr[i], "D9D9D9")
        _set_cell_borders(hdr[i])

    for section, num, item, location in PRISMA_ITEMS:
        if num == "" and item == "":
            # Section banner row
            row = table.add_row().cells
            for i, c in enumerate(row):
                c.width = widths[i]
                _set_cell_borders(c)
                _set_cell_shading(c, "F2F2F2")
            row[0].text = ""
            p = row[0].paragraphs[0]
            run = p.add_run(section)
            _set_run_font(run, size_pt=10, bold=True)
            # Merge across all 4 columns
            row[0].merge(row[1]).merge(row[2]).merge(row[3])
            continue

        row = table.add_row().cells
        for i, c in enumerate(row):
            c.width = widths[i]
            _set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Section / Topic
        p = row[0].paragraphs[0]
        run = p.add_run(section)
        _set_run_font(run, size_pt=9, bold=False)

        # Item number
        p = row[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(num)
        _set_run_font(run, size_pt=9, bold=True)

        # Checklist item
        p = row[2].paragraphs[0]
        run = p.add_run(item)
        _set_run_font(run, size_pt=9, bold=False)

        # Location reported
        p = row[3].paragraphs[0]
        run = p.add_run(location)
        _set_run_font(run, size_pt=9, bold=False)

    # ---- Footer note ----
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(6)
    run = foot.add_run(
        "Page numbers are not used because they shift between the "
        "journal-formatted PDF and the reviewer (double-spaced) PDF; locations "
        "are referenced by section and subsection headings instead. The full "
        "PRISMA 2020 statement and Explanation & Elaboration document are "
        "available at https://www.prisma-statement.org."
    )
    _set_run_font(run, size_pt=9, bold=False)

    doc.save(str(PRISMA_OUT))
    print(f"Wrote {PRISMA_OUT}")


# ---------------------------------------------------------------------------
# Declaration of Interest Statement (Elsevier standard form)
# ---------------------------------------------------------------------------
def build_declaration_of_interest():
    doc = Document()
    _configure_page(doc, landscape=False)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Declaration of Interest Statement")
    _set_run_font(run, size_pt=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    run = sub.add_run(
        "Big Five Personality Traits and Academic Achievement in Online "
        "Learning Environments: A Systematic Review and Meta-Analysis"
    )
    _set_run_font(run, size_pt=11, bold=False)

    # ---- Author block ----
    for line, bold in [
        ("Author: Eisuke Tokiwa, MEng", True),
        ("Affiliation: Founder, SUNBLAZE Co., Ltd., Tokyo, Japan", False),
        ("ORCID: 0009-0009-7124-6669", False),
        ("Email: eisuke.tokiwa@sunblaze.jp", False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        _set_run_font(run, size_pt=11, bold=bold)

    doc.add_paragraph()  # spacer

    # ---- Statement (Elsevier standard wording, adapted) ----
    _add_paragraph(
        doc,
        "The author whose name is listed immediately above certifies that "
        "they have NO affiliations with or involvement in any organization "
        "or entity with any financial interest (such as honoraria; "
        "educational grants; participation in speakers' bureaus; membership, "
        "employment, consultancies, stock ownership, or other equity "
        "interest; and expert testimony or patent-licensing arrangements), "
        "or non-financial interest (such as personal or professional "
        "relationships, affiliations, knowledge or beliefs) in the subject "
        "matter or materials discussed in this manuscript, with the single "
        "exception described below.",
        size=11,
        space_after=Pt(8),
    )

    _add_paragraph(
        doc,
        "Disclosed potentially competing interest:",
        size=11, bold=True, space_after=Pt(4),
    )
    _add_paragraph(
        doc,
        "One of the author's prior primary studies — Tokiwa, E. (2025), "
        "\"Who excels in online learning in Japan?\", Frontiers in "
        "Psychology, 16, Article 1420996 "
        "(https://doi.org/10.3389/fpsyg.2025.1420996; CC BY, open access) "
        "— is potentially eligible for inclusion in the present "
        "systematic review and meta-analysis. To address this transparently "
        "and a priori, a sensitivity analysis excluding the author's own "
        "study was pre-specified in the publicly time-stamped protocol "
        "(OSF Registries, https://doi.org/10.17605/OSF.IO/E5W47, "
        "registered 23 April 2026, prior to formal data extraction). The "
        "sensitivity analysis result, reported in the Results section and "
        "in Table 4 of the manuscript, leaves the primary trait-level "
        "conclusions unchanged: |Δr| < .001 for every Big Five trait, "
        "because Tokiwa (2025) reported test-completion outcomes without "
        "extractable trait-by-achievement zero-order Pearson correlations "
        "and therefore did not contribute effect sizes to the primary "
        "quantitative pool. The study is, however, retained in the "
        "qualitative-synthesis set (k = 25) and described narratively in "
        "the corresponding subsection.",
        size=11, space_after=Pt(8),
    )

    _add_paragraph(
        doc,
        "Funding:",
        size=11, bold=True, space_after=Pt(4),
    )
    _add_paragraph(
        doc,
        "This research did not receive any specific grant from funding "
        "agencies in the public, commercial, or not-for-profit sectors.",
        size=11, space_after=Pt(8),
    )

    _add_paragraph(
        doc,
        "Employment and other potentially competing roles:",
        size=11, bold=True, space_after=Pt(4),
    )
    _add_paragraph(
        doc,
        "The author is the founder of SUNBLAZE Co., Ltd. (Tokyo, Japan). "
        "SUNBLAZE had no role in the design or conduct of the review; in "
        "the collection, management, analysis, or interpretation of the "
        "data; in the preparation, review, or approval of the manuscript; "
        "or in the decision to submit the manuscript for publication. The "
        "author has no patents, products under development, or marketed "
        "products related to the subject matter of this review.",
        size=11, space_after=Pt(8),
    )

    _add_paragraph(
        doc,
        "Use of generative-AI tools in manuscript preparation:",
        size=11, bold=True, space_after=Pt(4),
    )
    _add_paragraph(
        doc,
        "Generative-AI assistants (Anthropic Claude) were used during "
        "manuscript preparation for routine writing-aid tasks (language "
        "polishing, table formatting, and Python/R code scaffolding for "
        "the figure and table builders deposited on OSF). All scientific "
        "content, including the search strategy, screening decisions, "
        "data extraction, statistical analyses, interpretation of results, "
        "and conclusions, is the work of the author, who reviewed and "
        "verified all AI-assisted output and takes full responsibility "
        "for the integrity of the manuscript.",
        size=11, space_after=Pt(8),
    )

    _add_paragraph(
        doc,
        "Authorship and CRediT statement (sole author):",
        size=11, bold=True, space_after=Pt(4),
    )
    _add_paragraph(
        doc,
        "Eisuke Tokiwa: Conceptualization, Methodology, Investigation, "
        "Data Curation, Formal Analysis, Software, Validation, "
        "Visualization, Writing — Original Draft, Writing — Review & "
        "Editing, Project Administration, Funding Acquisition (none "
        "obtained). The author meets all four ICMJE authorship criteria "
        "and accepts full accountability for all aspects of the work, "
        "including its accuracy and integrity.",
        size=11, space_after=Pt(12),
    )

    # ---- Signature ----
    _add_paragraph(doc, "Signed,", size=11, space_after=Pt(2))
    _add_paragraph(doc, "Eisuke Tokiwa", size=11, bold=True, space_after=Pt(0))
    _add_paragraph(
        doc,
        "Founder, SUNBLAZE Co., Ltd. — Tokyo, Japan",
        size=11, space_after=Pt(0),
    )
    _add_paragraph(
        doc,
        "ORCID: 0009-0009-7124-6669 — Email: eisuke.tokiwa@sunblaze.jp",
        size=11, space_after=Pt(8),
    )

    from datetime import date
    _add_paragraph(
        doc,
        f"Date: {date.today().strftime('%B %d, %Y')}",
        size=11, space_after=Pt(0),
    )

    doc.save(str(DOI_OUT))
    print(f"Wrote {DOI_OUT}")


def main():
    build_prisma_checklist()
    build_declaration_of_interest()


if __name__ == "__main__":
    main()
