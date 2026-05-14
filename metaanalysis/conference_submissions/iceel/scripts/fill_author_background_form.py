"""Fill in the ICEEL Author-Background Form for paper EL2036.

The form is a 1-table, 8-row, 5-col layout. Row 0 is the header; the author
goes in row 1. Single-author submission, so rows 2-7 stay blank.
"""
from pathlib import Path

from docx import Document

SRC = Path("metaanalysis/conference_submissions/iceel/Author-Background Form.docx")
OUT = Path("metaanalysis/conference_submissions/iceel/Author-Background Form FILLED.docx")

VALUES = {
    "name_aff_country": "Mr. Eisuke Tokiwa, SUNBLAZE Co., Ltd., Japan",
    "email": "eisuke.tokiwa@sunblaze.jp",
    "prefix": "PhD Candidate",
    "research_field": "Educational Psychology; Online Learning; Big Five Personality; Meta-Analysis; Cross-Cultural Psychology",
    "website": "https://orcid.org/0009-0009-7124-6669",
}


def set_cell(cell, text):
    """Replace all paragraphs in a cell with a single paragraph of `text`,
    preserving the cell's paragraph style if possible."""
    # Wipe existing paragraphs
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    cell.add_paragraph(text)


def main():
    doc = Document(str(SRC))
    table = doc.tables[0]
    row = table.rows[1]
    set_cell(row.cells[0], VALUES["name_aff_country"])
    set_cell(row.cells[1], VALUES["email"])
    set_cell(row.cells[2], VALUES["prefix"])
    set_cell(row.cells[3], VALUES["research_field"])
    set_cell(row.cells[4], VALUES["website"])
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    # Verify
    out_doc = Document(str(OUT))
    r1 = out_doc.tables[0].rows[1]
    for ci, cell in enumerate(r1.cells):
        print(f"  [1][{ci}] {repr(cell.text.strip()[:90])}")


if __name__ == "__main__":
    main()
