"""Build ECEL 2026 ACPI submission package for paper EL-095.

Reads metaanalysis/conference_submissions/ecel/full_paper.md and produces:

  EL-095-Tokiwa.docx  — paper with author details (Calibri 10pt, A4, ACPI styles)
  EL-095.docx          — blinded version (author info removed; self-citation kept)

ACPI specs (from Style-model-paper-Jan_2025-2.pdf):
  - Calibri 10pt throughout
  - A4 page; 2.54 cm margins
  - Title in Title Case, <= 12 words
  - Abstract 300-500 words
  - Total <= 5000 words, <= 10 pages
  - Headings: Heading 1 (top sections), Heading 2 (subsections), Heading 3 (sub-sub)
  - References in Harvard style (we use author-year inline citations + alphabetical list)

Run:
  python3 metaanalysis/conference_submissions/ecel/scripts/build_acpi_submission.py
"""
from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[4]
ECEL = ROOT / "metaanalysis" / "conference_submissions" / "ecel"
TEMPLATE = ECEL / "_acpi_template.docx"
SRC_MD = ECEL / "full_paper.md"
FIG = ECEL / ".." / "figures" / "prisma_flow_ecel.png"

OUT_FULL = ECEL / "EL-095-Tokiwa.docx"
OUT_BLIND = ECEL / "EL-095.docx"

# 12-word title required by ECEL form
TITLE = "Modality matters for Extraversion: A Big Five meta-regression in online learning environments"

AUTHOR_BLOCK = [
    "Eisuke Tokiwa",
    "SUNBLAZE Co., Ltd., Tokyo, Japan",
    "eisuke.tokiwa@sunblaze.jp",
]

ABSTRACT_PREFIX = "Abstract: "
KEYWORDS_PREFIX = "Keywords: "
KEYWORDS = "meta-analysis, Big Five, online learning, learning modality, Extraversion"


# ----------------------------------------------------------------------
# Markdown parsing (minimal, paper-specific)
# ----------------------------------------------------------------------
def load_md_sections():
    """Return ordered list of (kind, level, text) tuples.

    kinds: 'title', 'h1', 'h2', 'h3', 'p', 'table', 'figure', 'list'
    """
    raw = SRC_MD.read_text(encoding="utf-8")
    lines = raw.splitlines()
    # Strip frontmatter: everything between the title (#) and the first
    # `## ` heading is metadata (Author, Affiliation, ORCID, Email,
    # Target venue, Target length, Manuscript draft, ...). Drop it.
    title_seen = False
    skip_until_first_h2 = False
    filtered: list[str] = []
    for ln in lines:
        if ln.startswith("# ") and not title_seen:
            title_seen = True
            skip_until_first_h2 = True
            filtered.append(ln)
            continue
        if skip_until_first_h2:
            if ln.startswith("## "):
                skip_until_first_h2 = False
                filtered.append(ln)
            # else: drop frontmatter line silently
            continue
        filtered.append(ln)
    lines = filtered
    out = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        # Headings
        if line.startswith("# "):
            out.append(("title", 0, line[2:].strip()))
            i += 1; continue
        if line.startswith("### "):
            out.append(("h3", 3, line[4:].strip()))
            i += 1; continue
        if line.startswith("## "):
            out.append(("h2", 2, line[3:].strip()))
            i += 1; continue
        # Skip horizontal rules, blank lines, and frontmatter lines
        # (Author / Affiliation / ORCID / Email / Target venue /
        # Manuscript draft etc.). These appear before the first "##"
        # heading in our markdown source; here we just blacklist the
        # `**Key**:` shape.
        FRONTMATTER_KEYS = (
            "**Author",
            "**Affiliation",
            "**ORCID",
            "**Email",
            "**Target venue",
            "**Manuscript draft",
            "**Status",
            "**Word count",
            "**Date",
        )
        if line.strip() in ("---", "") or any(line.startswith(k) for k in FRONTMATTER_KEYS):
            i += 1; continue
        # Image (PRISMA figure)
        if line.startswith("!["):
            # capture caption from inside the brackets
            m = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            caption = m.group(1) if m else ""
            out.append(("figure", 0, caption))
            i += 1; continue
        # Markdown table block: header line | --- | --- ...
        if line.startswith("|") and i + 1 < n and re.match(r"^\|\s*[-:|\s]+\|\s*$", lines[i+1]):
            tbl = [line]
            i += 1
            tbl.append(lines[i])
            i += 1
            while i < n and lines[i].startswith("|"):
                tbl.append(lines[i])
                i += 1
            out.append(("table", 0, "\n".join(tbl)))
            continue
        # List item
        if line.lstrip().startswith(("- ", "* ", "1. ")):
            li_block = [line]
            i += 1
            while i < n and lines[i].lstrip().startswith(("- ", "* ", "1. ", "  ")):
                li_block.append(lines[i])
                i += 1
            out.append(("list", 0, "\n".join(li_block)))
            continue
        # Paragraph — collect until blank
        para = [line]
        i += 1
        while i < n and lines[i].strip() and not (lines[i].startswith(("#", "|", "!", "- ", "* "))):
            para.append(lines[i])
            i += 1
        text = " ".join(p.strip() for p in para).strip()
        if text:
            out.append(("p", 0, text))
    return out


# ----------------------------------------------------------------------
# Inline formatting: convert simple markdown to plain-text-with-runs.
# We strip **bold** and _italic_ marks; ACPI says "do not use bold for
# emphasis", italics are acceptable. To keep it simple, convert **x**
# to plain text and _x_ to italic runs.
# ----------------------------------------------------------------------
INLINE_RE = re.compile(r"(\*\*([^*]+)\*\*)|(\*([^*]+)\*)")


def add_runs(paragraph, text, font_name="Calibri", font_size=10):
    """Add text to paragraph, converting **bold** and *italic* markers."""
    # ACPI: avoid bold for emphasis. So **x** -> plain, *x* -> italic.
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = font_name
            r.font.size = Pt(font_size)
        if m.group(2) is not None:
            # **bold** -> plain (ACPI rule)
            r = paragraph.add_run(m.group(2))
        else:
            r = paragraph.add_run(m.group(4))
            r.italic = True
        r.font.name = font_name
        r.font.size = Pt(font_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = font_name
        r.font.size = Pt(font_size)


# ----------------------------------------------------------------------
# Section heading text rewrite for Heading 1 — strip the leading number
# (the model's Heading 1 already implies numbering, so we keep our
# numbering: "1. Introduction" stays as the literal text).
# ----------------------------------------------------------------------
def add_heading(doc, text, level):
    # The template was converted from .doc by LibreOffice; python-docx
    # cannot resolve the heading style by display name "Heading 1" because
    # the WP styleId is "Heading1" (no space) and there is no matching
    # name<->id alias entry. We look the style up by id and assign it
    # via the paragraph's style property directly.
    style_id = {1: "Heading1", 2: "Heading2", 3: "Heading3"}[level]
    style = doc.styles[style_id]
    p = doc.add_paragraph()
    p.style = style
    add_runs(p, text)
    return p


# ----------------------------------------------------------------------
# Table builder
# ----------------------------------------------------------------------
def add_table_from_md(doc, md_block):
    rows = [r for r in md_block.splitlines() if r.strip().startswith("|")]
    if len(rows) < 2:
        return
    sep_idx = next((i for i, r in enumerate(rows) if re.match(r"^\|\s*[-:|\s]+\|\s*$", r)), 1)
    headers = [c.strip() for c in rows[0].strip("|").split("|")]
    body = []
    for r in rows[sep_idx + 1:]:
        cells = [c.strip() for c in r.strip("|").split("|")]
        body.append(cells)
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(body), cols=ncols)
    # The LibreOffice-converted template doesn't ship the standard table
    # styles, so leave the default Normal Table style in place.
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, h)
    for i, row in enumerate(body, start=1):
        for j in range(ncols):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, row[j] if j < len(row) else "")
    # Add blank paragraph after the table for spacing
    doc.add_paragraph()


def add_figure(doc, caption):
    """Insert PRISMA figure with caption."""
    if FIG.exists():
        doc.add_picture(str(FIG))
    # Caption (Calibri 10pt, "Figure 1." prefix already in caption text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, caption)


def add_list(doc, md_block):
    for line in md_block.splitlines():
        s = line.lstrip()
        if not s:
            continue
        marker = s[:2]
        text = s[2:].strip() if marker in ("- ", "* ") else re.sub(r"^\d+\.\s+", "", s)
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, text)


# ----------------------------------------------------------------------
# Build docx for a given mode: "full" (author info) or "blind"
# ----------------------------------------------------------------------
def set_page_margins(doc, cm=2.54):
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def add_title_and_author(doc, mode):
    # Re-use the existing first paragraph (already in the body) as the title
    # to avoid the "no Heading 1 style" issue from python-docx after the
    # body has been emptied.
    if doc.paragraphs:
        p = doc.paragraphs[0]
    else:
        p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    # Authors (full mode only)
    if mode == "full":
        for line in AUTHOR_BLOCK:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, line)
    # Spacer
    doc.add_paragraph()


# For blind mode: replace author identifiers in body text and references
# so a reviewer cannot trivially identify the author. ACPI requires "no
# author details", and while reviewers in our small research community
# may still guess from the topic, removing the literal name + DOI is
# standard practice. We map the author's own Frontiers article to a
# bracketed placeholder.
BLIND_REPLACEMENTS = [
    # Specific references first (longest match wins; order matters)
    ("Tokiwa, E. (2025). Who excels in online learning in Japan? Frontiers in Psychology, 16, Article 1420996. https://doi.org/10.3389/fpsyg.2025.1420996",
     "[Author's own primary study, 2025] (full reference withheld for double-blind review)"),
    ("Tokiwa, E. (2025).", "[Author], (2025)."),
    ("(A-25 Tokiwa", "(A-25 [Author's prior study]"),
    ("A-25 Tokiwa", "A-25 [Author's prior study]"),
    ("Tokiwa (2025)", "[Author] (2025)"),
    ("(Tokiwa, 2025)", "([Author], 2025)"),
    ("(Tokiwa 2025)", "([Author], 2025)"),
    ("10.3389/fpsyg.2025.1420996", "[DOI withheld for double-blind review]"),
]


def apply_blind_replacements(text: str) -> str:
    for old, new in BLIND_REPLACEMENTS:
        text = text.replace(old, new)
    return text


def build(mode: str, out_path: Path):
    sections = load_md_sections()
    if mode == "blind":
        # Mutate paragraph text in-place
        sections = [(k, l, apply_blind_replacements(t)) for (k, l, t) in sections]
    doc = Document(str(TEMPLATE))
    # The template ships with one empty paragraph; leave the body element
    # in place (its presence is required for some style lookups) but
    # remove all but the first paragraph (which we'll repurpose for the
    # title). Removing all paragraphs breaks `add_paragraph(style=...)`
    # lookups in python-docx 1.x.
    paras = list(doc.paragraphs)
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # Clear the first paragraph's content so we can re-use it as the title
    first = paras[0] if paras else None
    if first is not None:
        for r in list(first.runs):
            r.text = ""
    set_page_margins(doc)

    # ----- Title + author block -----
    add_title_and_author(doc, mode)

    # ----- Now stream sections from md, but skip:
    #   - the source title (we replaced with 12-word ECEL title)
    #   - the "**Author**:" line (handled via AUTHOR_BLOCK)
    #   - acknowledgments in blind mode
    # ------------------------------------------------------------------
    in_abstract = False
    in_keywords = False
    skip_section = False  # True when we're inside a section to skip in blind mode

    for kind, level, text in sections:
        if kind == "title":
            # original markdown title — replaced
            continue
        if kind == "h2" and text.startswith("Abstract"):
            in_abstract = True
            in_keywords = False
            skip_section = False
            continue  # we'll inline "Abstract: ..." into the next paragraph
        if kind == "h2" and (text.startswith("Acknowledgments") or text.startswith("Acknowledgements")):
            # Skip acknowledgments in blind mode
            skip_section = (mode == "blind")
            if not skip_section:
                add_heading(doc, text, 1)
            continue
        if kind in ("h2", "h3"):
            # End of acknowledgments-skip when next heading hits
            skip_section = False
            in_abstract = False
        if skip_section:
            continue
        if kind == "h2":
            add_heading(doc, text, 1)
            continue
        if kind == "h3":
            add_heading(doc, text, 2)
            continue
        if kind == "p":
            # Detect Keywords paragraph (starts with **Keywords**:) — we
            # bake the keywords explicitly to enforce the 5-keyword version.
            if text.lower().startswith("keywords"):
                p = doc.add_paragraph()
                r = p.add_run(KEYWORDS_PREFIX)
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10)
                add_runs(p, KEYWORDS)
                in_abstract = False
                continue
            if in_abstract:
                # Strip any leading "Abstract:" already in text
                txt = re.sub(r"^Abstract:?\s*", "", text).strip()
                p = doc.add_paragraph()
                r = p.add_run(ABSTRACT_PREFIX)
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10)
                add_runs(p, txt)
                in_abstract = False
                continue
            p = doc.add_paragraph()
            add_runs(p, text)
            continue
        if kind == "table":
            add_table_from_md(doc, text)
            continue
        if kind == "figure":
            add_figure(doc, text)
            continue
        if kind == "list":
            add_list(doc, text)
            continue

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main():
    build("full", OUT_FULL)
    build("blind", OUT_BLIND)


if __name__ == "__main__":
    main()
