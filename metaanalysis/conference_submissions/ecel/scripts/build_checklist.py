"""Pre-fill the ECEL EL-095 Copyright Checklist for the user.

Auto-fills the non-signature fields in Copyright-checklist-ACPIL_2026.docx
and saves the result as EL-095-Checklist.docx. The user still needs to
add their signature to the bottom Signature cell (hand-sign on print, or
paste a signature image, or type the name as a typed-signature) before
attaching to the submission email.
"""
from pathlib import Path

from docx import Document

ECEL = Path(__file__).resolve().parents[1]
SRC = ECEL / "Copyright-checklist-ACPIL_2026.docx"
OUT = ECEL / "EL-095-Checklist.docx"


def set_cell(cell, text):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    cell.add_paragraph(text)


def main():
    doc = Document(str(SRC))
    info = doc.tables[0]
    # Rows 0-2 are horizontally merged (gridSpan=2): label + value live
    # in the same cell, with the user expected to type the value after
    # the label and trailing colon. We preserve the label and append.
    def append_after_label(cell, value):
        # The cell already contains a paragraph with text like
        # "Conference Name & Paper Reference: ". Append the value to the
        # first non-empty paragraph's last run.
        for p in cell.paragraphs:
            if p.runs and p.text.strip():
                p.add_run(value)
                return
        # Fall back: add a fresh paragraph
        cell.add_paragraph(value)

    append_after_label(info.rows[0].cells[0], "ECEL 2026 / EL-095")
    append_after_label(info.rows[1].cells[0],
        "Modality matters for Extraversion: A Big Five meta-regression in online learning environments")
    append_after_label(info.rows[2].cells[0],
        "Eisuke Tokiwa, eisuke.tokiwa@sunblaze.jp")
    # Rows 4-13: Yes/No confirmations -> "Yes"
    # Row 14: visa requirement -> "No"
    yes_rows = [4, 5, 6, 7, 8, 9, 10, 11, 12, 13]
    for ri in yes_rows:
        set_cell(info.rows[ri].cells[1], "Yes")
    set_cell(info.rows[14].cells[1], "No (virtual attendance, no visa required)")

    # Signature table (Name / Email / Signature placeholder / Date)
    sig = doc.tables[1]
    # Row 0: cells [0]=label 'Name', [1]=value, [2]=label 'Email', [3]=value
    set_cell(sig.rows[0].cells[1], "Eisuke Tokiwa")
    set_cell(sig.rows[0].cells[3], "eisuke.tokiwa@sunblaze.jp")
    # Row 1: cells [0]=label 'Signature', [1]=value (leave blank for user
    # to hand-sign or paste image), [2]=label 'Date', [3]=value
    set_cell(sig.rows[1].cells[3], "2026-05-17")
    # Signature cell left empty for manual signing
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print()
    print("Manual step remaining:")
    print("  Open the file in Word, add your signature to the Signature")
    print("  cell of the bottom table (hand-sign on print, paste a")
    print("  signature image, or type your name as a typed signature),")
    print("  then save before attaching to the submission email.")


if __name__ == "__main__":
    main()
