"""Generate the HSSC submission cover letter as a Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADER_LINES = [
    "Eisuke Tokiwa, MEng",
    "Founder, SUNBLAZE Co., Ltd.",
    "Tokyo, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Email: eisuke.tokiwa@sunblaze.jp",
]

ADDRESSEE_LINES = [
    "To the Editors,",
    "Humanities and Social Sciences Communications",
]

BODY_PARAGRAPHS = [
    "Dear Editors,",

    "I am pleased to submit “Big Five Personality Traits and Academic "
    "Achievement in Online Learning Environments: A Systematic Review and "
    "Meta-Analysis” for consideration as an Article in Humanities and "
    "Social Sciences Communications.",

    ("__Context and importance.__ "
     "Eight prior meta-analyses (Poropat, 2009; McAbee & Oswald, 2013; "
     "Vedel, 2014; Stajkovic et al., 2018; Mammadov, 2022; Zell & Lesick, "
     "2022; Meyer et al., 2023; Chen et al., 2025) have established "
     "Conscientiousness as the strongest Big Five predictor of academic "
     "achievement (ρ ≈ .19–.28), with smaller but reliable "
     "Openness effects. None of these syntheses, however, has tested "
     "learning modality—online, blended, or face-to-face—as a "
     "substantive moderator. As post-secondary instruction is increasingly "
     "delivered online, the question of whether established "
     "personality–achievement relationships transfer to "
     "technology-mediated environments is empirically open and societally "
     "consequential."),

    ("__Contribution.__ "
     "This pre-registered (OSF Registries e5w47, registered 23 April 2026) "
     "PRISMA 2020-compliant review is, to my knowledge, the first "
     "quantitative synthesis of this question. Drawing on 31 primary "
     "studies (10 contributing direct correlations to the primary pool; "
     "pooled N = 3,384), I report online-specific pooled estimates for "
     "each Big Five trait (Conscientiousness r = .167, 95% CI "
     "[.089, .243]; Agreeableness r = .112; Openness r = .086; Neuroticism "
     "r = .018; Extraversion r = .002), and document two highly "
     "significant pre-registered moderator effects: Extraversion × "
     "Region (Asian r = −.131 vs. non-Asian r = .050; "
     "Q_between = 46.43, p < .001) and Extraversion × Outcome Type "
     "(objective r = −.038 vs. self-rated r = .117; Q_between = "
     "17.30, p < .001). The findings preliminarily support a "
     "technology-mediated extension of the Personality–Achievement "
     "Saturation Hypothesis and document novel cultural- and "
     "outcome-dependent shifts in the Extraversion–achievement link."),

    ("__Fit with the journal.__ "
     "The work sits at the interdisciplinary intersection of educational "
     "psychology, individual differences, and the social transformation "
     "of learning under technology mediation—questions that extend "
     "beyond any single subfield. The findings are directly relevant to "
     "educators designing post-pandemic online curricula, to psychologists "
     "refining theories of personality–context interaction, and to "
     "policymakers concerned with cultural variation in technology-"
     "mediated education—audiences that align with the broad "
     "social-science readership of Humanities and Social Sciences "
     "Communications."),

    ("__Open science.__ "
     "Pre-registration: https://doi.org/10.17605/OSF.IO/E5W47. "
     "Preprint: https://doi.org/10.21203/rs.3.rs-9513298/v1 (Research "
     "Square v1, posted 27 April 2026). Data, analysis code, search "
     "logs, and supplementary materials: https://doi.org/10.17605/OSF.IO/79M5J "
     "(seven OSF components covering protocol, search, screening, "
     "extraction, risk of bias, analysis, and the article-level DOI "
     "index). A version-controlled mirror is maintained at "
     "https://github.com/etoki/paper."),

    ("__Conflict of interest.__ "
     "I declare no financial conflicts. One of my own previously "
     "published primary studies — Tokiwa, E. (2025), “Who excels "
     "in online learning in Japan?”, Frontiers in Psychology, 16, "
     "1420996 (CC BY, https://doi.org/10.3389/fpsyg.2025.1420996) — is "
     "potentially eligible for inclusion; this is addressed through a "
     "pre-specified sensitivity analysis excluding the author’s "
     "own study, which leaves the primary conclusions unchanged."),

    ("The manuscript has not been published elsewhere and is not under "
     "consideration by another journal. The submission has been prepared "
     "in compliance with the journal’s double-anonymous peer review "
     "policy: identifying information has been removed from the main "
     "manuscript and supplementary files, and the author contribution "
     "statement is provided through the submission portal rather than in "
     "the manuscript itself. I am the sole author and accept full "
     "accountability for all aspects of the work."),

    "Thank you for considering this submission.",

    "Sincerely,",
]


def add_paragraph(doc, text, bold_segments=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_after=Pt(6)):
    """Add a paragraph; supports inline bold via __segment__ markers."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15

    parts = text.split("__")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    return p


def build_document(out_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for line in HEADER_LINES:
        add_paragraph(doc, line, space_after=Pt(0))
    add_paragraph(doc, "", space_after=Pt(0))

    add_paragraph(doc, date.today().strftime("%B %d, %Y"))
    add_paragraph(doc, "", space_after=Pt(0))

    for line in ADDRESSEE_LINES:
        add_paragraph(doc, line, space_after=Pt(0))
    add_paragraph(doc, "", space_after=Pt(0))

    for paragraph in BODY_PARAGRAPHS:
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "Eisuke Tokiwa", space_after=Pt(0))
    add_paragraph(doc, "Founder, SUNBLAZE Co., Ltd.", space_after=Pt(0))
    add_paragraph(doc, "Tokyo, Japan", space_after=Pt(0))
    add_paragraph(doc, "ORCID: 0009-0009-7124-6669", space_after=Pt(0))
    add_paragraph(doc, "Email: eisuke.tokiwa@sunblaze.jp", space_after=Pt(0))

    doc.save(str(out_path))


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "cover_letter_hssc.docx"
    build_document(out)
    print(f"Wrote {out}")
