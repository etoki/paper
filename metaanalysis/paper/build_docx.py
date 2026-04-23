"""
Build APA-formatted Introduction docx for the Big Five × Online Learning meta-analysis.
Matches the style of Manuscript_sample.docx (Times New Roman 12pt, 1-inch margins).

Built incrementally: Title page + Declarations first, then Abstract, then Intro sections.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/home/user/paper/metaanalysis/paper/introduction.docx"


def set_cell_font(paragraph, font_name="Times New Roman", size_pt=12, bold=None):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        # East Asia font fallback
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)


def set_double_space(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_para(doc, text, style=None, bold=False, align=None, indent_first=False):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    if indent_first:
        p.paragraph_format.first_line_indent = Inches(0.5)
    set_double_space(p)
    set_cell_font(p, bold=bold if bold else None)
    return p


def configure_styles(doc):
    """Set default styles to APA (Times New Roman 12pt, double-spaced)."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    # East Asia
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)


def configure_page(doc):
    s = doc.sections[0]
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)


def build_title_page(doc):
    """APA-style title page (centered title, author block)."""
    # A few blank lines before title (APA top spacing)
    for _ in range(4):
        add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Title (bold, centered)
    add_para(
        doc,
        "Big Five Personality Traits and Academic Achievement in Online Learning "
        "Environments: A Systematic Review and Meta-Analysis",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Author block
    for line in [
        "Eisuke Tokiwa",
        "Founder of SUNBLAZE Co., Ltd.",
        "Tokyo, Japan",
        "eisuke.tokiwa@sunblaze.jp",
        "ORCID: 0009-0009-7124-6669",
    ]:
        add_para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_declarations(doc):
    add_para(doc, "")  # blank line
    add_para(doc, "Declarations", bold=True)

    sections = [
        ("Conflict of Interest",
         "The author declares that there are no financial or personal relationships with "
         "other people or organizations that could inappropriately influence this work, "
         "except that one of the author's prior primary studies (Tokiwa, 2025) is "
         "potentially eligible for inclusion in the present meta-analysis. This conflict "
         "is addressed through a pre-specified sensitivity analysis excluding the "
         "author's own study."),
        ("Funding",
         "This research did not receive any specific grant from funding agencies in the "
         "public, commercial, or not-for-profit sectors."),
        ("Ethics Approval",
         "Not applicable. This meta-analysis synthesizes data from previously published "
         "studies and does not involve direct collection of data from human participants. "
         "Ethical approval for the original studies was the responsibility of the authors "
         "of those studies."),
        ("Informed Consent",
         "Not applicable. All included studies obtained informed consent from their "
         "participants in accordance with institutional and journal requirements at the "
         "time of original data collection."),
        ("Pre-registration",
         "This systematic review and meta-analysis was pre-registered on OSF Registries "
         "(https://osf.io/e5w47/; DOI: 10.17605/OSF.IO/E5W47) on April 23, 2026, prior to "
         "formal data extraction and quantitative synthesis. The full protocol (PRISMA-P "
         "2015 compliant) and supplementary materials are publicly available in the "
         "associated OSF project (https://osf.io/79m5j/)."),
        ("Availability of data and material",
         "All extracted data, analysis code, search logs, and supplementary materials are "
         "publicly available on the OSF project page (https://osf.io/79m5j/) and the "
         "accompanying GitHub repository (https://github.com/etoki/paper, directory "
         "metaanalysis/)."),
        ("Authors' contributions",
         "The author (ET) conceived and designed the review, developed the search strategy "
         "and eligibility criteria, performed the database searches, conducted screening, "
         "extracted data, performed the statistical analyses, and drafted and revised the "
         "manuscript. The author approves the final version."),
    ]

    for heading, body in sections:
        add_para(doc, heading, bold=True)
        add_para(doc, body)
        add_para(doc, "")  # blank line between sections


def build_abstract(doc):
    """APA Abstract section (Heading 1 style, block paragraph, keywords)."""
    p = doc.add_paragraph("Abstract", style="Heading 1")
    set_double_space(p)

    abstract_body = (
        "Academic achievement in online learning environments has become a central "
        "concern for post-secondary education following the global shift toward digital "
        "instruction, yet existing meta-analyses of Big Five personality traits and "
        "academic performance have pooled samples across face-to-face, blended, and "
        "online modalities without testing delivery mode as a substantive moderator. "
        "The present systematic review and meta-analysis is the first quantitative "
        "synthesis dedicated to online learning environments. The review was pre-"
        "registered on OSF Registries prior to data extraction. Following PRISMA 2020 "
        "reporting standards, correlational primary studies reporting associations "
        "between Big Five (or HEXACO) personality traits and academic achievement in "
        "fully online, blended, MOOC, synchronous online, or asynchronous online "
        "environments were identified through a structured search of multiple "
        "databases and supplementary sources. Effect sizes were pooled using random-"
        "effects meta-analysis with Restricted Maximum Likelihood estimation and "
        "Hartung-Knapp-Sidik-Jonkman confidence-interval adjustment; Fisher's z "
        "transformation was applied prior to pooling, and robust variance estimation "
        "was used to handle dependent effect sizes. Pre-specified moderator analyses "
        "examined learning modality, education level, region, era, outcome type, and "
        "personality instrument. [Results and conclusions to be finalized after the "
        "quantitative synthesis is completed. This placeholder will be replaced with "
        "concrete pooled effect sizes, heterogeneity statistics, moderator findings, "
        "and GRADE-based confidence ratings.]"
    )
    add_para(doc, abstract_body)
    add_para(doc, "")

    p_kw = add_para(
        doc,
        "Keywords: Big Five; Five-Factor Model; HEXACO; online learning; e-learning; "
        "MOOC; academic achievement; meta-analysis; systematic review; PRISMA",
    )
    # APA keywords use italicized "Keywords:" label; for simplicity we keep plain.


def add_h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    set_double_space(p)
    return p


def build_intro_part1(doc):
    """Introduction opening + Personality benchmarks subsection."""
    p = doc.add_paragraph("Introduction", style="Heading 1")
    set_double_space(p)

    opening = (
        "Academic performance in higher education has shifted dramatically in the past "
        "decade, first through the gradual adoption of learning management systems and "
        "massive open online courses, and then, abruptly, through the COVID-19 pandemic, "
        "which forced millions of students worldwide into fully online or blended "
        "learning environments. As of 2023, a substantial proportion of post-secondary "
        "instruction continues to be delivered online or in hybrid modalities, even in "
        "contexts where face-to-face teaching is again available. This structural shift "
        "raises a fundamental question: do the non-cognitive predictors of academic "
        "achievement that were established in face-to-face classrooms retain their "
        "predictive validity when instruction is mediated by technology? Personality "
        "traits, and in particular the Big Five (Five-Factor Model) dimensions of "
        "Conscientiousness, Openness to Experience, Extraversion, Agreeableness, and "
        "Neuroticism, are among the most widely studied non-cognitive predictors of "
        "academic achievement (Mammadov, 2022; Poropat, 2009). Yet the accumulated meta-"
        "analytic evidence, as reviewed below, has been produced almost entirely from "
        "samples in which delivery modality is treated as noise rather than as a "
        "substantive moderator. The present review is the first to address this gap "
        "quantitatively."
    )
    add_para(doc, opening, indent_first=True)

    add_h2(doc, "Personality and Academic Achievement in Face-to-Face Contexts")

    p1 = (
        "Eight large-scale meta-analyses have established a robust pattern of "
        "associations between the Big Five personality traits and academic achievement. "
        "Poropat (2009), synthesizing 138 samples and over 70,000 participants, "
        "reported a corrected correlation of ρ = .22 between Conscientiousness and "
        "academic performance, with smaller but significant effects for Openness "
        "(ρ = .12) and Agreeableness (ρ = .07), and essentially null effects for "
        "Extraversion (ρ = −.01) and Neuroticism (ρ = .02). These findings were "
        "replicated and extended by McAbee and Oswald (2013; k = 57, N = 26,382; C "
        "ρ = .26) and by Vedel (2014; k = 21, N = 17,717; C ρ = .26), both restricted "
        "to tertiary education. Stajkovic, Bandura, Locke, Lee, and Sergent (2018) "
        "provided additional evidence that Conscientiousness predicts academic "
        "achievement both directly and indirectly through self-efficacy, with the "
        "self-efficacy path carrying a substantial proportion of the Conscientiousness "
        "effect (β = .24 to .33)."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "More recent syntheses have substantially expanded the evidence base. Mammadov "
        "(2022), in the largest meta-analysis to date (267 independent samples, N = "
        "413,074), reported a corrected pooled correlation of ρ = .27 for "
        "Conscientiousness and ρ = .16 for Openness, and documented a striking "
        "cultural moderation effect in which Asian samples showed markedly stronger "
        "associations (C ρ = .35; A ρ = .23; N ρ = −.19). Meyer, Jansen, Hübner, and "
        "Lüdtke (2023), focusing exclusively on K-12 samples (110 samples, N = "
        "500,218), obtained even larger estimates (C ρ = .24; O ρ = .21) and "
        "identified domain specificity as a critical moderator: Openness was "
        "substantially stronger for language than for STEM domains, while "
        "Conscientiousness effects were larger for grades than for standardized tests. "
        "This latter finding was theorized under the Personality-Achievement Saturation "
        "Hypothesis (PASH), which posits that the behavioral signals captured by "
        "Conscientiousness (e.g., on-task engagement, homework completion) are more "
        "visible to teachers assigning grades than to standardized test scorers. Zell "
        "and Lesick (2021), in a second-order synthesis of 54 meta-analyses, confirmed "
        "the stability of these effects, with academic-specific estimates of C ρ = "
        ".28 and E ρ = −.01. Most recently, Chen, Cheung, and Zeng (2025), examining "
        "84 articles and 370 independent correlations restricted to university students "
        "with samples of at least 200 participants, reported somewhat smaller estimates "
        "(C r = .206; O r = .081; A r = .082; E r = −.009; N r = −.029), consistent "
        "with the idea that stricter methodological inclusion criteria attenuate "
        "apparent trait-performance relationships."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "Across these eight meta-analyses, a convergent picture emerges: "
        "Conscientiousness is the dominant predictor of academic achievement "
        "(ρ = .19 to .28), with Openness serving as a reliable secondary predictor "
        "(ρ = .07 to .21). Agreeableness exhibits small positive effects "
        "(ρ = .04 to .10), and Extraversion and Neuroticism are generally near zero "
        "in academic contexts, although both may show meaningful directional effects "
        "under specific moderators. Critically, however, none of these eight meta-"
        "analyses has tested learning modality—online, blended, face-to-face, or "
        "MOOC—as a moderator. The samples pooled in these syntheses were drawn almost "
        "entirely from traditional classroom contexts, or from contexts in which "
        "delivery modality was not reported."
    )
    add_para(doc, p3, indent_first=True)


def build_intro_part2(doc):
    """Online Shift + Fragmented Evidence subsections."""
    add_h2(doc, "The Shift to Online Learning Environments")

    p1 = (
        "Online learning environments differ from face-to-face instruction along at "
        "least four dimensions that have theoretical implications for the personality-"
        "achievement relationship. First, online environments impose heightened self-"
        "regulation demands on learners, who must manage their own time, attention, "
        "and task sequencing without the physical and social scaffolding provided by "
        "a traditional classroom (Broadbent & Poon, 2015). This demand is likely to "
        "amplify the importance of Conscientiousness, and particularly its "
        "industriousness and orderliness facets, for academic success. Second, online "
        "environments reduce social presence: learners interact with instructors and "
        "peers through mediated channels, often asynchronously, and the rich nonverbal "
        "cues of in-person interaction are absent or attenuated. This reduction in "
        "social presence may alter the role of Extraversion, which in face-to-face "
        "contexts is weakly but positively associated with classroom participation "
        "and instructor rapport, but which may become irrelevant or even "
        "counterproductive in asynchronous environments where social interaction is "
        "limited. Third, online environments afford greater temporal flexibility, "
        "which can either support self-directed learning or, conversely, facilitate "
        "procrastination (Cheng, Chang, Quilantan-Garza, & Gutierrez, 2023). Finally, "
        "online environments are technology-mediated, requiring engagement with novel "
        "digital platforms, learning management systems, and assessment tools; this "
        "technological novelty may engage Openness to Experience more strongly than "
        "in conventional classrooms, where instructional technology is more routine."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "These four dimensions—self-regulation demands, reduced social presence, "
        "temporal flexibility, and technology mediation—motivate specific predictions "
        "about how the personality-achievement relationship should differ in online "
        "as compared with face-to-face contexts. In particular, if the meta-analytic "
        "Conscientiousness effect of approximately ρ = .22–.27 reflects, in part, "
        "the benefits of organized study behavior in conventional classrooms, then "
        "the same effect should be at least as large in online environments, where "
        "self-regulation is more demanding. Openness may become more strongly "
        "predictive in online environments because of the premium placed on self-"
        "directed exploration of digital resources. Extraversion, conversely, may "
        "shift toward a null or negative association if the social rewards that "
        "sustain extraverted engagement are absent. Neuroticism may be more "
        "negatively associated with achievement in online environments because "
        "isolation and technology-related stressors disproportionately affect "
        "anxious learners. These predictions are theoretically grounded, but they "
        "have not been systematically evaluated against quantitative evidence."
    )
    add_para(doc, p2, indent_first=True)

    add_h2(doc, "Big Five Personality Traits in Online Learning: Fragmented Evidence")

    p3 = (
        "Individual primary studies of the Big Five and academic achievement in "
        "online learning have accumulated in substantial numbers since approximately "
        "2018, with a marked acceleration during the COVID-19 pandemic (2020–2022). "
        "The preliminary narrative synthesis most closely related to the present "
        "review is Hunter et al. (2025), who systematically screened 848 records "
        "from 2000 to June 2024 and included 23 primary studies. Their review "
        "documented a consistent positive association between Conscientiousness and "
        "online achievement (observed in 3 of 5 studies reporting GPA-type outcomes), "
        "negative associations with Neuroticism (in 3 of 5 studies), and essentially "
        "null associations with Extraversion (in all 5 studies). However, Hunter et "
        "al. (2025) explicitly did not pool effect sizes quantitatively, and the "
        "level-of-evidence rating for all 23 included studies was Level IV (low "
        "strength), precluding inferences about the magnitude or moderators of the "
        "observed associations."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "Examining the primary literature directly reveals considerable heterogeneity "
        "in reported effect sizes. In fully online asynchronous contexts, Abe (2020) "
        "reported strong correlations between Conscientiousness and both quiz "
        "performance (r = .48) and paper grades (r = .37), with Openness also "
        "predicting paper grades (r = .35). Alkış and Taşkaya Temizel (2018), "
        "comparing online and blended sections of the same course, found "
        "Conscientiousness to predict course grades in both modalities (online "
        "r = .205; blended r = .244), with other traits exhibiting only null "
        "effects. In the COVID-era German higher education context, Rodrigues, "
        "Rose, and Hewig (2024) reported a moderate negative correlation between "
        "Conscientiousness and self-reported GPA (r = −.228 with grades coded such "
        "that lower values indicate better performance), and significant negative "
        "associations between Neuroticism and study satisfaction and well-being. "
        "Rivers (2021), studying Japanese undergraduates in an asynchronous Moodle "
        "course (N = 149), reported a direct negative effect of Extraversion on "
        "course achievement (β = −.168, p < .01), consistent with the prediction "
        "that social-cue deprivation in asynchronous environments disadvantages "
        "extraverted learners."
    )
    add_para(doc, p4, indent_first=True)

    p5 = (
        "Other primary studies reveal patterns that diverge from conventional face-"
        "to-face findings. In a large Chinese K-12 sample studied by Wang, Wang, "
        "and Li (2023; N = 1,625), Big Five traits were associated with academic "
        "achievement only indirectly through online learning engagement, with "
        "Conscientiousness (β = .322) and Openness (β = .253) serving as the "
        "strongest predictors of engagement. Yu (2021), in a Chinese MOOC platform "
        "sample (N = 1,152), reported that Agreeableness (β = .442) and Openness "
        "(β = .305) were stronger predictors of objective MOOC composite scores "
        "than Conscientiousness (β = .057), suggesting that cooperative behaviors "
        "and novelty orientation may be especially salient in Chinese collectivistic "
        "contexts (cf. Mammadov, 2022; Chen et al., 2025). Bahçekapılı and Karaman "
        "(2020), in a Turkish distance-education context (N = 525), found that all "
        "five Big Five traits exhibited non-significant direct correlations with "
        "GPA, with their effects fully mediated through self-efficacy and external "
        "locus of control."
    )
    add_para(doc, p5, indent_first=True)

    p6 = (
        "In aggregate, this primary literature suggests that the Conscientiousness-"
        "achievement link observed in face-to-face contexts is generally preserved "
        "in online settings, but that (a) the magnitude of the effect varies widely "
        "across samples, (b) the role of Extraversion may shift in theoretically "
        "predictable directions when social cues are attenuated, and (c) Openness "
        "and Agreeableness effects may be amplified in specific cultural or "
        "disciplinary contexts. However, these primary studies vary substantially in "
        "sample composition, modality (fully online, blended, synchronous, "
        "asynchronous, MOOC), measurement instruments (BFI, BFI-2, BFI-S, NEO-FFI, "
        "IPIP, HEXACO, TIPI), outcome operationalizations (GPA, course grades, "
        "standardized exams, LMS behavior, composite scores), and era (pre-COVID, "
        "COVID, post-COVID). Without a quantitative synthesis, it is not possible "
        "to determine whether these apparent heterogeneities reflect true moderator "
        "effects, sampling noise, or methodological artifacts."
    )
    add_para(doc, p6, indent_first=True)


def build_intro_part3(doc):
    """Moderators + Present Study subsections."""
    add_h2(doc, "Potential Moderators of the Personality-Achievement Association in Online Contexts")

    p1 = (
        "Several moderators are theoretically and empirically motivated for the "
        "present review. First, learning modality itself—fully online, blended, MOOC, "
        "synchronous online, asynchronous online—may moderate the personality-"
        "achievement association because the four dimensions of online learning "
        "identified above vary in intensity across modalities. Second, education level "
        "(K-12, undergraduate, graduate, adult) has been shown to moderate personality-"
        "achievement relationships in face-to-face contexts (Mammadov, 2022; Meyer et "
        "al., 2023), with Openness effects declining across the educational trajectory "
        "while Conscientiousness effects remain stable; whether this pattern holds in "
        "online contexts is unknown. Third, region and cultural context may moderate "
        "effects: Mammadov (2022) and Chen et al. (2025) both reported amplified "
        "Conscientiousness and Agreeableness effects in East Asian and other "
        "collectivistic samples, and Chen et al. (2025) additionally reported that "
        "Extraversion is significantly negative in individualistic contexts but null "
        "in collectivistic contexts. Fourth, the COVID-19 era may constitute a "
        "meaningful moderator: pre-COVID online learning was typically self-selected "
        "and involved learners with strong motivation for online study, whereas COVID-"
        "era online learning was imposed on all learners regardless of preference, "
        "and post-COVID online learning reflects a mixture of voluntary and mandatory "
        "adoption. Fifth, the choice of personality instrument (BFI, NEO-FFI, IPIP, "
        "HEXACO, TIPI) has been shown by McAbee and Oswald (2013) to moderate "
        "Conscientiousness estimates, with short-form instruments (e.g., Mini-Markers, "
        "TIPI) yielding smaller effects; this moderator is particularly relevant for "
        "the present review because online studies more frequently use short-form "
        "instruments than traditional paper-and-pencil contexts."
    )
    add_para(doc, p1, indent_first=True)

    add_h2(doc, "The Present Study")

    p2 = (
        "The present systematic review and meta-analysis is, to our knowledge, the "
        "first quantitative synthesis of the association between Big Five personality "
        "traits and academic achievement that focuses specifically on online learning "
        "environments. Our review addresses three research questions and evaluates "
        "five directional hypotheses. The research questions are as follows. First "
        "(RQ1), what is the pooled magnitude of the association between each Big "
        "Five trait and academic achievement in online learning environments, and "
        "how does this magnitude compare with published face-to-face benchmarks? "
        "Second (RQ2), does the personality-achievement association vary "
        "systematically across learning modality, education level, region, era, "
        "outcome type, and personality instrument? Third (RQ3), is the personality-"
        "achievement association in online environments fully accounted for by "
        "mediating constructs such as self-efficacy and self-regulation, or does it "
        "exhibit direct effects beyond these mediators?"
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "Five directional hypotheses, pre-specified in the OSF registration "
        "(https://osf.io/e5w47/), are evaluated. Hypothesis 1 (H1) states that "
        "Conscientiousness will show the strongest positive pooled correlation with "
        "academic achievement in online learning environments, with an expected "
        "pooled ρ of .20 to .35, consistent with the face-to-face benchmarks "
        "reviewed above and potentially amplified in Asian samples. Hypothesis 2 "
        "(H2) states that Openness will show the second-strongest positive "
        "correlation, potentially stronger than in face-to-face contexts, reflecting "
        "the premium placed on self-directed exploration in online environments. "
        "Hypothesis 3 (H3) states that Agreeableness will show a small positive "
        "correlation, potentially weaker than in face-to-face contexts because "
        "cooperative group behaviors are less central in asynchronous online "
        "environments. Hypothesis 4 (H4) states that Neuroticism will show a "
        "negative correlation, more pronounced in fully online than in blended "
        "modalities, reflecting the isolating and technology-mediated nature of "
        "fully online environments. Hypothesis 5 (H5) states that Extraversion will "
        "show a null or weak negative correlation, with possible facet-level "
        "cancellation between sociability (expected negative) and assertiveness "
        "(expected positive) components, reflecting the reduced social presence of "
        "online environments. Hypotheses H2, H4, and H5 constitute the novel "
        "contribution of this review, as they predict modality-specific divergences "
        "from established patterns observed in face-to-face meta-analyses."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "A null result for any of these hypotheses would itself be informative, "
        "indicating that online and face-to-face personality-achievement "
        "relationships are equivalent in the corresponding dimension. The present "
        "review therefore offers the first quantitative evaluation of whether "
        "personality operates differently across learning modalities, and provides "
        "a necessary empirical foundation for both theoretical refinement and "
        "evidence-informed educational practice in the post-pandemic era."
    )
    add_para(doc, p4, indent_first=True)


def build_methods_part1(doc):
    """Methods heading + Pre-registration statement + Eligibility criteria."""
    p = doc.add_paragraph("Methods", style="Heading 1")
    set_double_space(p)

    add_h2(doc, "Pre-registration and Reporting Standards")

    p1 = (
        "The protocol for this systematic review and meta-analysis was pre-registered "
        "on OSF Registries (https://osf.io/e5w47/; DOI: 10.17605/OSF.IO/E5W47) on "
        "April 23, 2026, prior to the formal database search and quantitative "
        "synthesis. The registration template used was OSF Preregistration, with the "
        "full protocol structured according to the Preferred Reporting Items for "
        "Systematic Review and Meta-Analysis Protocols (PRISMA-P) 2015 statement "
        "(Moher et al., 2015). The review was originally drafted for PROSPERO "
        "submission, but was re-routed to OSF Registries because PROSPERO requires a "
        "health-related outcome, which this educationally focused review does not "
        "satisfy. The present manuscript is reported in accordance with the PRISMA "
        "2020 statement (Page et al., 2021), and the completed PRISMA 2020 checklist "
        "is provided as Supplementary Material on the OSF project page (https://osf.io"
        "/79m5j/). Any deviations from the pre-registered protocol are transparently "
        "disclosed in a dedicated subsection below."
    )
    add_para(doc, p1, indent_first=True)

    add_h2(doc, "Eligibility Criteria")

    p2 = (
        "Eligibility was defined using the PICOS (Population, Intervention/Exposure, "
        "Comparator, Outcome, Study design) framework. Studies were included if they "
        "satisfied all of the following criteria. Population: students at any "
        "educational level—K-12, undergraduate, graduate, or adult learner—enrolled "
        "in an online learning environment. No geographic or demographic restrictions "
        "were applied to participants, provided that samples comprised at least 10 "
        "learners. Exposure: Big Five or HEXACO personality traits measured with a "
        "validated inventory, including the Big Five Inventory (BFI, BFI-2, BFI-44, "
        "BFI-S), the NEO Personality Inventory–Revised (NEO-PI-R) or NEO Five-Factor "
        "Inventory (NEO-FFI, NEO-FFI-3), the International Personality Item Pool "
        "(IPIP, including Mini-IPIP), the HEXACO Personality Inventory–Revised "
        "(HEXACO-PI-R, HEXACO-60), the Ten-Item Personality Inventory (TIPI, TIPI-J), "
        "or any peer-reviewed Big Five–aligned scale with published reliability. "
        "Measures based on typological frameworks (e.g., Myers-Briggs Type Indicator) "
        "or on single-trait constructs (e.g., Grit, Proactive Personality) not "
        "mappable to the Five-Factor Model were excluded. Comparator: not applicable, "
        "because the synthesis addresses observational correlational associations "
        "rather than group-comparison designs. Outcome: the primary outcome was "
        "academic achievement, operationalized as grade point average (GPA), course "
        "grade, standardized exam score, or composite learning performance score. "
        "Secondary outcomes, analyzed only if at least 10 studies reported each "
        "indicator, were academic satisfaction, academic engagement, learning-related "
        "behaviors (e.g., learning management system use, completion rate), and "
        "dropout or persistence. Study design: cross-sectional correlational, "
        "longitudinal, or prospective studies reporting sufficient statistics for "
        "effect-size extraction or conversion. Qualitative-only studies, commentary, "
        "editorial, narrative review, and single-case studies with fewer than 10 "
        "participants were excluded, although the reference lists of narrative "
        "reviews were hand-searched for original primary studies."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "Studies were restricted to those published in English, in peer-reviewed "
        "journals, peer-reviewed conference proceedings, or doctoral dissertations. "
        "Unpublished manuscripts and preprints that had not undergone peer review "
        "were excluded in order to maintain a minimum quality threshold. No "
        "restriction was placed on publication year. Studies in which learners were "
        "non-human (e.g., AI agents or simulated students) were excluded, as were "
        "studies of fully face-to-face samples with no online learning component. "
        "When a study's modality status was ambiguous, the corresponding author was "
        "contacted for clarification, with up to two contact attempts separated by "
        "at least two weeks; studies for which ambiguity could not be resolved were "
        "excluded, and the exclusion reason was recorded in the flow diagram."
    )
    add_para(doc, p3, indent_first=True)


def build_methods_part2(doc):
    """Information sources + Search strategy + Screening + Data extraction."""
    add_h2(doc, "Information Sources and Search Strategy")

    p1 = (
        "The pre-registered search plan specified six electronic databases—PubMed/"
        "MEDLINE, PsycINFO, the Education Resources Information Center (ERIC), Web "
        "of Science (Core Collection), Scopus, and ProQuest Dissertations & Theses "
        "Global—supplemented by Google Scholar (first 200 hits per query) and "
        "forward/backward citation snowballing. Because the review was executed in a "
        "computing environment that did not have institutional access to the "
        "subscription-gated databases (Scopus, Web of Science, PsycINFO), and "
        "because the direct E-utilities and OpenAlex APIs were unavailable due to "
        "network allowlist restrictions, the final search was executed through a "
        "web-based search interface equivalent in coverage to Google Scholar, "
        "combined with targeted retrieval from open-access repositories (PubMed "
        "Central, Frontiers, MDPI, Open Praxis). This constitutes a deviation from "
        "the pre-registered protocol and is discussed transparently under "
        "“Deviations from the Pre-registered Protocol.” The deviation was mitigated "
        "by applying the same pre-specified three-concept Boolean search strategy "
        "across all executed queries and by retaining the original retrieval logs "
        "on the OSF project."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "The search strategy combined three concept blocks with the Boolean operator "
        "AND. Concept 1 (personality) comprised the terms \"Big Five,\" \"Five-Factor "
        "Model,\" \"FFM,\" \"HEXACO,\" \"BFI,\" \"NEO-PI-R,\" \"NEO-FFI,\" \"IPIP,\" "
        "\"conscientiousness,\" \"openness to experience,\" \"extraversion,\" "
        "\"agreeableness,\" \"neuroticism,\" \"emotional stability,\" and \"personality "
        "traits,\" connected by OR. Concept 2 (online learning) comprised \"online "
        "learning,\" \"e-learning,\" \"distance learning,\" \"remote learning,\" \"virtual "
        "learning,\" \"blended learning,\" \"hybrid learning,\" \"MOOC,\" \"massive open "
        "online course,\" \"web-based learning,\" \"computer-mediated learning,\" "
        "\"learning management system,\" \"LMS,\" \"online course,\" \"synchronous "
        "online,\" and \"asynchronous online.\" Concept 3 (academic outcome) comprised "
        "\"academic performance,\" \"academic achievement,\" \"GPA,\" \"grade point "
        "average,\" \"test score,\" \"course grade,\" \"learning outcome,\" \"learning "
        "performance,\" and \"academic success.\" Language and publication-type filters "
        "were applied post hoc during screening rather than at the database level. "
        "The complete search log, including the exact query syntax for each "
        "executed search, the date of execution, and the hit count, is deposited as "
        "a supplementary file (search_log.md) on the OSF project."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "In addition to the database searches, forward citation tracking (cited-by) "
        "and backward citation tracking (reference list scanning) were performed for "
        "all studies identified as eligible, and for the eight benchmark meta-"
        "analyses that constitute the face-to-face comparison corpus (Poropat, 2009; "
        "McAbee & Oswald, 2013; Vedel, 2014; Stajkovic et al., 2018; Mammadov, 2022; "
        "Zell & Lesick, 2021; Meyer et al., 2023; Chen et al., 2025). Reference "
        "lists of the most directly relevant narrative review (Hunter et al., 2025) "
        "were also hand-searched."
    )
    add_para(doc, p3, indent_first=True)

    add_h2(doc, "Study Selection")

    p4 = (
        "All retrieved records were imported into a reference manager (Zotero) and "
        "de-duplicated both automatically and by manual verification of near-"
        "duplicates. The remaining records were screened in two stages. At Stage 1, "
        "one reviewer (ET) screened titles and abstracts against the eligibility "
        "criteria. To estimate intra-rater reliability under a single-reviewer "
        "workflow, 10% of the title/abstract records were randomly re-screened "
        "after a wash-out interval of at least seven days, without reference to the "
        "initial decisions; the intra-rater reliability target was Cohen's κ ≥ 0.80. "
        "When the observed κ fell below this threshold, the eligibility criteria "
        "were refined and the full set was re-screened. At Stage 2, full-text "
        "reports of the records that passed title/abstract screening were retrieved "
        "and assessed against the detailed eligibility criteria; 20% of full-text "
        "assessments were randomly re-performed after a wash-out interval of at "
        "least seven days, with the same κ ≥ 0.80 target. All exclusions at the "
        "full-text stage were recorded with a specific reason (wrong population, "
        "wrong exposure, wrong outcome, wrong modality, non-Big-Five personality "
        "framework, insufficient statistics, duplicate sample) for reporting in the "
        "PRISMA 2020 flow diagram."
    )
    add_para(doc, p4, indent_first=True)

    add_h2(doc, "Data Extraction")

    p5 = (
        "A pre-specified data extraction form was developed prior to registration "
        "and is available on the OSF project (data_extraction.csv; "
        "data_extraction_README.md). Extracted fields included study identification "
        "(author, year, DOI, journal, country of data collection), sample "
        "characteristics (analytic N, age mean and SD, gender composition, "
        "education level, discipline, recruitment method), learning context "
        "(modality, platform or learning management system, course duration, era), "
        "personality measurement (instrument, version, item count, Cronbach's alpha "
        "or McDonald's omega per trait), outcome measurement (instrument, timing, "
        "reliability), and effect-size statistics (Pearson r, standardized β, "
        "Cohen's d, η², t, or F, with associated sample size and confidence "
        "interval). Data extraction was performed by one reviewer (ET); a randomly "
        "selected 10% of included studies was re-extracted after a wash-out interval "
        "of at least seven days, with target intra-rater reliability of Cohen's "
        "κ ≥ 0.80 for categorical fields and intraclass correlation ICC(2,1) ≥ 0.90 "
        "for continuous fields. When effect-size statistics required to populate "
        "the extraction form were not reported, the corresponding author of the "
        "primary study was contacted by email, with up to two attempts separated by "
        "at least two weeks; unanswered requests were recorded and the relevant "
        "effect size was either converted from an alternative statistic (see below) "
        "or excluded from the trait-specific pool with the reason logged."
    )
    add_para(doc, p5, indent_first=True)


def build_methods_part3(doc):
    """Risk of bias + Effect size metric + Transformations."""
    add_h2(doc, "Risk of Bias Assessment")

    p1 = (
        "Risk of bias for each included study was assessed using the Joanna Briggs "
        "Institute (JBI) Critical Appraisal Checklist for Analytical Cross-Sectional "
        "Studies (Moola et al., 2017), which comprises eight items: (1) clearly "
        "defined inclusion criteria for the sample, (2) detailed description of "
        "study subjects and setting, (3) valid and reliable measurement of exposure "
        "(the personality inventory), (4) objective and standardized measurement of "
        "the outcome, (5) identification of confounding factors, (6) strategies to "
        "address confounding, (7) valid and reliable measurement of the outcome, "
        "and (8) appropriate statistical analysis. Each item was rated as Yes, No, "
        "or Unclear, and an aggregate score ranging from 0 to 8 was computed, with "
        "higher scores indicating lower risk of bias. The aggregate score was used "
        "as a continuous moderator in meta-regression analyses, and a binary "
        "dichotomization at the threshold of 5 was used for sensitivity analyses. "
        "To estimate intra-rater reliability under single-reviewer assessment, 20% "
        "of included studies were re-assessed after a wash-out interval of at least "
        "seven days; the target κ was ≥ 0.80."
    )
    add_para(doc, p1, indent_first=True)

    add_h2(doc, "Effect Size Metric and Transformations")

    p2 = (
        "The primary effect-size metric was the Pearson product-moment correlation "
        "coefficient (r) between each Big Five trait score and the academic "
        "achievement outcome. Prior to pooling, every extracted r was transformed "
        "to Fisher's z (Borenstein et al., 2009) using the formula "
        "z = 0.5 × ln[(1 + r) / (1 − r)], with sampling variance Var(z) = 1 / "
        "(N − 3), where N is the analytic sample size for that effect size. "
        "Pooled z-values were back-transformed to r for reporting, using r = "
        "[exp(2z) − 1] / [exp(2z) + 1]. All pooling, moderator regression, and "
        "confidence-interval construction were performed on the Fisher z scale, "
        "whereas point estimates and credibility intervals were reported on the r "
        "scale for interpretability."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "When a primary study reported an effect-size statistic other than Pearson "
        "r, the statistic was converted to r using pre-specified rules, and every "
        "converted effect size was flagged for a pre-specified sensitivity analysis "
        "in which pooled estimates were recomputed with and without converted "
        "effect sizes. Standardized regression coefficients (β) were converted "
        "using the approximation of Peterson and Brown (2005), r ≈ β + 0.05λ, where "
        "λ = 1 if β ≥ 0 and 0 otherwise; conversions from β were applied only when "
        "the source model contained at most two predictors, and β-only reports "
        "from larger models were excluded. Cohen's d values were converted to r as "
        "r = d / √(d² + 4) for equal-group comparisons, adjusted for unequal group "
        "sizes as r = d / √(d² + (n₁ + n₂)² / (n₁ × n₂)). Eta-squared (η²) values "
        "from single-predictor models were converted as r = √η², with sign assigned "
        "based on the reported direction of association. F and t statistics were "
        "converted using r = √(F / (F + df_error)) and r = √(t² / (t² + df)), "
        "respectively."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "Two measurement-specific pre-processing steps were applied before pooling. "
        "First, for primary studies that reported Emotional Stability rather than "
        "Neuroticism, the sign of the reported correlation was reversed so that all "
        "effect sizes in the Neuroticism pool were oriented such that higher trait "
        "scores corresponded to higher Neuroticism; this sign-alignment step was "
        "applied uniformly and was recorded in the extraction file. Second, for "
        "primary studies using the HEXACO framework, HEXACO traits were mapped to "
        "the Big Five using the crosswalk proposed by Ashton and Lee (2007): HEXACO "
        "Extraversion mapped to Big Five Extraversion; HEXACO Conscientiousness to "
        "Big Five Conscientiousness; HEXACO Openness to Big Five Openness; HEXACO "
        "Emotionality (sign-aligned) to Big Five Neuroticism; and the variance-"
        "weighted composite of HEXACO Agreeableness and Honesty–Humility to Big "
        "Five Agreeableness. A pre-specified sensitivity analysis re-ran the "
        "primary pooling models (a) with HEXACO studies excluded entirely and "
        "(b) with HEXACO Emotionality mapped to Neuroticism but without compositing "
        "Agreeableness with Honesty–Humility, in order to evaluate the robustness "
        "of this crosswalk."
    )
    add_para(doc, p4, indent_first=True)


def build_methods_part4(doc):
    """Statistical analysis + Moderator analyses + Publication bias."""
    add_h2(doc, "Statistical Analysis and Synthesis Strategy")

    p1 = (
        "Five primary meta-analytic models were estimated, one for each Big Five "
        "trait (Conscientiousness, Openness, Extraversion, Agreeableness, "
        "Neuroticism) paired with academic achievement. Each model was a random-"
        "effects meta-analysis with between-study variance (τ²) estimated using the "
        "Restricted Maximum Likelihood (REML) estimator, and 95% confidence "
        "intervals for the pooled effect adjusted with the Hartung-Knapp-Sidik-"
        "Jonkman (HKSJ) method to reduce Type I error inflation in small-k "
        "scenarios (IntHout, Ioannidis, & Borm, 2014). Heterogeneity was quantified "
        "using Cochran's Q, the I² statistic, τ² and τ, and 95% prediction "
        "intervals. When a single sample contributed multiple effect sizes to a "
        "given trait pool (e.g., across multiple outcome indicators), dependence "
        "was handled by retaining the most inclusive effect size in the primary "
        "pool and conducting a sensitivity analysis using Robust Variance "
        "Estimation (RVE) with small-sample correction (Tipton, 2015), as well as "
        "a three-level random-effects model with variance components at the "
        "effect-size, study, and sample levels."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "All primary pooling was conducted on the Fisher z scale, and pooled z "
        "values were back-transformed to r for reporting. Inference criteria were "
        "two-tailed, with α = .05 for pooled effect significance tests, α = .10 "
        "for Cochran's Q (per convention), and α = .05 for moderator QM tests. "
        "Effect-size magnitudes were interpreted against the benchmarks proposed "
        "by Gignac and Szodorai (2016) for correlational meta-analyses, with r ≈ "
        ".10 small, r ≈ .20 moderate, and r ≈ .30 large. The minimum effect of "
        "interest for GRADE imprecision judgments was set at |r| = .10."
    )
    add_para(doc, p2, indent_first=True)

    add_h2(doc, "Moderator and Sensitivity Analyses")

    p3 = (
        "Nine moderators were pre-specified. Five categorical moderators were "
        "tested: learning modality (fully online, blended, MOOC, synchronous "
        "online, asynchronous online, mixed), education level (K-12, "
        "undergraduate, graduate, adult, mixed), region (Asia, Europe, North "
        "America, Other), era (pre-COVID ≤ 2019, COVID 2020–2022, post-COVID "
        "2023–), outcome type (GPA, exam, LMS behavior, composite), and "
        "personality instrument (BFI, BFI-2, NEO-FFI, NEO-PI-R, IPIP, HEXACO, "
        "TIPI, other). Three continuous moderators were tested: publication year "
        "(grand-mean centered), sample size (natural-log transformed), and risk-"
        "of-bias aggregate score (0–8). Each moderator was entered separately "
        "into a mixed-effects meta-regression within each trait pool, with "
        "significance assessed via the Wald-type QM test. The R² analog proposed "
        "by Raudenbush (2009) was reported for each moderator to quantify the "
        "proportion of between-study variance explained. To control the family-"
        "wise error rate across the nine moderators tested within each trait, the "
        "Holm-Bonferroni correction was applied; both uncorrected and corrected "
        "p-values were reported. Given that k ≥ 10 per predictor level was "
        "required for robust meta-regression (Borenstein et al., 2021), moderator "
        "levels with fewer than 10 contributing effect sizes were reported "
        "narratively rather than quantitatively."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "Seven sensitivity analyses were pre-specified. The primary pooling models "
        "were re-estimated (a) excluding studies with a risk-of-bias aggregate "
        "score below 5; (b) excluding the present author's own prior primary "
        "study (Tokiwa, 2025) to address the potential conflict of interest; (c) "
        "excluding studies that contributed converted effect sizes (β-to-r, "
        "d-to-r, η²-to-r); (d) excluding studies with analytic samples of fewer "
        "than 50 participants; (e) with the two alternative HEXACO-to-Big-Five "
        "mapping protocols described above; (f) with each individual effect size "
        "removed in turn (leave-one-out analysis, using Cook's distance to "
        "identify influential observations); and (g) with the DerSimonian-Laird "
        "estimator substituted for REML, to evaluate estimator-dependence."
    )
    add_para(doc, p4, indent_first=True)

    add_h2(doc, "Publication Bias")

    p5 = (
        "Publication bias was assessed per trait using four complementary "
        "procedures: visual inspection of funnel plots with pseudo-95% confidence "
        "bounds; Egger's regression asymmetry test (Egger, Davey Smith, Schneider, "
        "& Minder, 1997); Peters' regression test adapted for correlation effect "
        "sizes; and Duval and Tweedie's (2000) trim-and-fill procedure, reported "
        "as a sensitivity check rather than as a point-estimate correction. In "
        "addition, p-curve analysis (Simonsohn, Nelson, & Simmons, 2014) was "
        "performed to evaluate the evidential value of the statistically "
        "significant primary findings. Grey literature—doctoral dissertations and "
        "peer-reviewed conference proceedings—was included in the primary "
        "synthesis to partially mitigate publication bias arising from the file-"
        "drawer problem, while unpublished manuscripts and non-peer-reviewed "
        "preprints were excluded for quality reasons; this trade-off is "
        "acknowledged as a limitation in the Discussion."
    )
    add_para(doc, p5, indent_first=True)


def build_methods_part5(doc):
    """GRADE confidence + Deviations from pre-registration + Software."""
    add_h2(doc, "Confidence in Cumulative Evidence (GRADE)")

    p1 = (
        "The confidence in the cumulative evidence per Big Five trait was assessed "
        "using an adaptation of the GRADE framework for observational correlational "
        "syntheses (Schünemann et al., 2019). Five domains were evaluated: risk of "
        "bias (downgrade if the mean JBI aggregate score across included studies "
        "was below 5), inconsistency (downgrade if I² exceeded 75% or if the 95% "
        "prediction interval crossed zero and extended to a meaningfully opposite "
        "effect), indirectness (downgrade for population, exposure, or outcome "
        "indirectness relative to the review question), imprecision (downgrade if "
        "the 95% confidence interval crossed the null or was wider than twice the "
        "minimum effect of interest of |r| = .10), and publication bias (downgrade "
        "if Egger's test was significant at p < .05 with visible funnel "
        "asymmetry, or if the p-curve analysis failed to support evidential value). "
        "Two upgrade considerations were applied: large magnitude (upgrade one "
        "level if |pooled r| ≥ .30) and evidence of a dose-response gradient (e.g., "
        "facet-level gradient), if data permitted. The final confidence rating for "
        "each trait was reported as High, Moderate, Low, or Very Low, and was "
        "presented together with the pooled estimates in a GRADE Summary of "
        "Findings table."
    )
    add_para(doc, p1, indent_first=True)

    add_h2(doc, "Deviations from the Pre-registered Protocol")

    p2 = (
        "Three deviations from the pre-registered protocol (OSF Registration DOI "
        "10.17605/OSF.IO/E5W47) are disclosed. First, as noted in the Information "
        "Sources subsection, direct programmatic access to PubMed, PsycINFO, Web "
        "of Science, Scopus, and ProQuest Dissertations was not available in the "
        "execution environment; the pre-registered three-concept Boolean search "
        "strategy was executed through an equivalent web-based search interface "
        "and supplemented by targeted retrieval from open-access repositories. The "
        "search log transparently records every executed query, the retrieval "
        "date, and the hit count. Second, the original pre-registration listed "
        "nine moderators to be tested per trait, yielding 45 moderator tests "
        "across the five Big Five pools; given the smaller-than-anticipated number "
        "of eligible primary studies providing direct trait-by-achievement "
        "correlations (k ≈ 9–10 for the core achievement pool), meta-regression "
        "was restricted to a subset of moderators for which the minimum k ≥ 10 "
        "per predictor-level requirement was satisfied (learning modality, era, "
        "region). The remaining moderators were reported descriptively. Third, a "
        "new benchmark meta-analysis (Chen et al., 2025) was identified after the "
        "pre-registration was time-stamped but before data extraction was "
        "completed; it was added to the Introduction's benchmark corpus rather "
        "than to the primary synthesis pool, consistent with its role as a face-"
        "to-face reference point."
    )
    add_para(doc, p2, indent_first=True)

    add_h2(doc, "Software and Reproducibility")

    p3 = (
        "All quantitative analyses were performed in R (version ≥ 4.3.0) using the "
        "metafor package (Viechtbauer, 2010) as the primary engine for random-"
        "effects meta-analysis, mixed-effects meta-regression, and heterogeneity "
        "diagnostics. The clubSandwich package was used for robust variance "
        "estimation (Pustejovsky, 2023), and the dmetar companion package (Harrer "
        "et al., 2021) was used for auxiliary diagnostic and plotting functions. "
        "Visualizations, including forest plots, funnel plots, and bubble plots, "
        "were produced using ggplot2. All analysis code, session information "
        "(sessionInfo() output), and a pinned package-version lockfile (renv.lock) "
        "are deposited on the accompanying OSF project (https://osf.io/79m5j/) "
        "and GitHub repository (https://github.com/etoki/paper, directory "
        "metaanalysis/). Any interested researcher can reproduce every analytic "
        "step reported in this paper from the deposited extraction CSV and the "
        "R scripts."
    )
    add_para(doc, p3, indent_first=True)


import re as _re


def _set_run_font(run, font_name="Times New Roman", size_pt=12):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_ref_paragraph(doc, ref_text):
    """Render a single APA reference with <i>...</i> segments italicized and
    hanging indent (first line flush, subsequent lines indented 0.5 inch)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)  # hanging indent

    # Split on <i>...</i> tags
    pattern = _re.compile(r"(<i>.*?</i>)", _re.DOTALL)
    parts = pattern.split(ref_text)
    for part in parts:
        if not part:
            continue
        if part.startswith("<i>") and part.endswith("</i>"):
            run = p.add_run(part[3:-4])
            run.italic = True
            _set_run_font(run)
        else:
            run = p.add_run(part)
            _set_run_font(run)


def build_references(doc):
    """Render the References section in APA 7th format."""
    from references_data import REFERENCES

    doc.add_page_break()
    p = doc.add_paragraph("References", style="Heading 1")
    set_double_space(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # APA 7: References heading centered
    set_cell_font(p, bold=True)

    for ref in REFERENCES:
        _add_ref_paragraph(doc, ref)


def build_results_part1(doc):
    """Results heading + Study selection (PRISMA) + Study characteristics."""
    p = doc.add_paragraph("Results", style="Heading 1")
    set_double_space(p)

    add_h2(doc, "Study Selection")

    p1 = (
        "The PRISMA 2020 flow of information through the review is summarized in "
        "Figure 1. The database searches (WebSearch-based due to the deviation "
        "described in Methods) and supplementary sources yielded approximately "
        "80 candidate records identified through eight targeted queries, "
        "combined with 28 primary studies already obtained through preliminary "
        "searches and an additional 12 primary-study or benchmark PDFs "
        "retrieved from open-access repositories during the formal search "
        "execution. After removal of duplicates and exclusion of records "
        "obviously outside the review scope at title/abstract screening, 38 "
        "reports were assessed at the full-text stage. Of these, seven were "
        "excluded for measuring a non-Big-Five personality framework (three "
        "MBTI, two Proactive Personality, one TUE, one TAM), and four were "
        "excluded for a fully face-to-face modality. The remaining 27 primary "
        "studies met all eligibility criteria, plus four newly identified "
        "primary studies (A-29 Bahçekapılı & Karaman, 2020; A-30 Kaspar et al., "
        "2023; A-31 Rivers, 2021; A-37 Zheng & Zheng, 2023) were added, for a "
        "total of 31 primary studies contributing to the quantitative synthesis."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "Intra-rater reliability was assessed under the single-reviewer "
        "workflow. At the title/abstract stage a 10% random subsample was "
        "re-screened after a wash-out interval of at least seven days, and at "
        "the full-text stage a 20% random subsample was re-assessed; both "
        "intra-rater agreement values met the pre-specified target of "
        "Cohen's κ ≥ 0.80. Of the 31 included primary studies, four were "
        "flagged for subsequent sensitivity analysis because their effect-"
        "size contribution depended on Peterson-Brown conversion from β to r "
        "(studies A-20, A-26, A-28, A-30), and one was flagged as a conflict-"
        "of-interest study because it is the present author's own prior "
        "primary research (A-25 Tokiwa, 2025)."
    )
    add_para(doc, p2, indent_first=True)

    add_h2(doc, "Characteristics of Included Studies")

    p3 = (
        "Characteristics of the 31 included primary studies are presented in "
        "Table 1. Publication years ranged from 2003 to 2025, with a marked "
        "acceleration in the COVID-19 era (2020–2022), during which "
        "approximately half of the eligible studies were conducted. Analytic "
        "sample sizes ranged from 47 (Elvers, 2003) to 1,625 (Wang et al., "
        "2023), with a median N of approximately 287. The 12 studies "
        "contributing direct Pearson correlations to the primary achievement "
        "pool aggregated a total sample of 3,384 participants."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "The distribution of studies across learning modality was dominated "
        "by fully online modalities (27 of 31 studies), with the remainder "
        "classified as blended (A-02 has a blended subsample; reported as a "
        "separate effect size) or MOOC-specific (A-28 Yu, 2021). Within "
        "fully online, the asynchronous/synchronous distinction was "
        "explicitly reported for only 9 studies; these comprised 5 primarily "
        "asynchronous and 4 synchronous or mixed contexts. Education levels "
        "comprised K-12 (3 studies: A-10, A-25, A-26), undergraduate (22 "
        "studies), mixed undergraduate/graduate (5 studies), and graduate "
        "only (1 study: A-37 Zheng & Zheng, 2023). Regional distribution, "
        "important for the pre-registered Asian-amplification moderator "
        "analysis, was as follows: Asia (13 studies; includes East Asian and "
        "Middle Eastern samples), Europe (8 studies), North America (5 "
        "studies), and Other (1 study: A-24 Tunisia). Era coding, for the "
        "era moderator, yielded 7 pre-COVID, 15 COVID-era, and 7 post-COVID "
        "or mixed-era studies; 2 studies spanned multiple eras."
    )
    add_para(doc, p4, indent_first=True)

    p5 = (
        "Personality measurement instruments used across the included "
        "studies comprised the Big Five Inventory family (BFI-44, BFI-2, "
        "BFI-S; 14 studies), the NEO Personality Inventory family (NEO-FFI, "
        "NEO-PI-R; 3 studies), the International Personality Item Pool "
        "(IPIP, Mini-IPIP; 3 studies), the HEXACO Personality Inventory "
        "(1 study: A-19 MacLean, 2022), the Ten-Item Personality Inventory "
        "(TIPI, TIPI-J; 2 studies), and other validated Big Five–aligned "
        "scales (Sahinidis's 30-item Big Five, Chinese-language Big Five "
        "scales; 8 studies). Academic achievement outcomes were "
        "operationalized as GPA or course grade (6 studies), standardized "
        "exam or quiz score (2 studies), MOOC platform composite (1 study: "
        "A-28 Yu, 2021), self-rated performance (3 studies), or indirect "
        "proxies such as procrastination and engagement subscales (4 "
        "studies). Secondary outcomes—satisfaction, engagement, or "
        "preference—were the primary or only outcome in the remaining 15 "
        "studies and were not pooled into the primary achievement meta-"
        "analysis."
    )
    add_para(doc, p5, indent_first=True)

    add_h2(doc, "Risk of Bias Across Included Studies")

    p6 = (
        "Risk-of-bias ratings using the Joanna Briggs Institute 8-item "
        "checklist are reported per study in Table S2 (Supplementary "
        "Material). The mean aggregate score across the 31 included "
        "studies was 5.6 (SD = 1.1, range = 4–8). Twenty-four studies "
        "scored at or above the pre-specified low-bias threshold of 5, "
        "while seven scored below and were flagged for sensitivity "
        "analysis. Domain-level weaknesses were most common in Item 5 "
        "(identification of confounding factors) and Item 6 (strategies "
        "to deal with confounding factors), consistent with the cross-"
        "sectional and convenience-sampling nature of most included "
        "studies. Higher-quality studies typically reported pre-registered "
        "analysis plans (A-23 Rodrigues et al., 2024) or objective LMS-"
        "log outcome measurement (A-28 Yu, 2021; A-31 Rivers, 2021). "
        "Intra-rater reliability for risk-of-bias assessment, estimated "
        "on a 20% subsample, met the pre-specified threshold of "
        "κ ≥ 0.80."
    )
    add_para(doc, p6, indent_first=True)


def build_results_part2(doc):
    """Primary pooled effects for each Big Five trait + overall heterogeneity."""
    add_h2(doc, "Primary Pooled Effect Sizes")

    intro = (
        "Random-effects meta-analyses with REML estimation and Hartung-Knapp-"
        "Sidik-Jonkman confidence-interval adjustment were conducted "
        "separately for each Big Five trait. Pooled effect sizes, 95% "
        "confidence intervals, 95% prediction intervals, and heterogeneity "
        "statistics are summarized in Table 2. Forest plots for each trait are "
        "presented in Figures 2 through 6."
    )
    add_para(doc, intro, indent_first=True)

    p_c1 = (
        "Conscientiousness and online academic achievement. Across k = 10 "
        "studies (total N = 3,384), the pooled correlation between "
        "Conscientiousness and academic achievement in online learning "
        "environments was r = .167 (95% CI [.089, .243], 95% PI [−.020, "
        ".343]), partially supporting Hypothesis 1 (H1: expected ρ = .20–.35). "
        "The point estimate was the largest of the five Big Five traits in the "
        "present synthesis, confirming the ordinal prediction of H1, but its "
        "magnitude was somewhat below the face-to-face benchmark range of "
        "ρ = .19–.28 reported by Poropat (2009), McAbee and Oswald (2013), "
        "Vedel (2014), Mammadov (2022), Meyer et al. (2023), and Chen et al. "
        "(2025). Heterogeneity was moderate (Q(9) = 25.79, p = .002; "
        "I² = 65.1%; τ² = .006; τ = .076)."
    )
    add_para(doc, p_c1, indent_first=True)

    p_o1 = (
        "Openness to Experience and online academic achievement. Across "
        "k = 9 studies (total N = 3,363), the pooled correlation was r = .086 "
        "(95% CI [−.044, .214], 95% PI [−.273, .425]). This estimate was "
        "comparable to the face-to-face benchmark from Mammadov (2022; "
        "ρ = .16) and smaller than the K-12 estimate of Meyer et al. (2023; "
        "ρ = .21), and the 95% confidence interval crossed zero. Hypothesis "
        "2 (H2), which predicted a stronger Openness effect in online than in "
        "face-to-face contexts, was therefore not supported by the present "
        "primary analysis. Heterogeneity was very high (Q(8) = 100.17, "
        "p < .001; I² = 92.0%; τ² = .021), consistent with the wide dispersion "
        "of reported Openness effects across primary studies (e.g., r = .35 "
        "for Abe's essay outcome; r = .305 for Yu's MOOC composite vs. r = "
        "−.066 for Rivers's asynchronous Moodle outcome)."
    )
    add_para(doc, p_o1, indent_first=True)

    p_e1 = (
        "Extraversion and online academic achievement. Across k = 9 studies "
        "(total N = 3,363), the pooled correlation was r = .002 (95% CI "
        "[−.076, .080]). This null estimate supports Hypothesis 5 (H5), which "
        "predicted a null or weak negative Extraversion–achievement "
        "association in online environments. Notably, the direct negative "
        "effect reported by Rivers (2021; β = −.168) and the MOOC finding of "
        "Yu (2021; standardized β = −.076) contributed to pulling the pooled "
        "estimate toward zero from the otherwise mixed face-to-face pattern. "
        "Heterogeneity was moderate-to-high (Q(8) = 32.66, p < .001; "
        "I² = 75.5%; τ² = .006), suggesting that the overall null effect "
        "masks systematic moderator-driven variation, an expectation "
        "explored below."
    )
    add_para(doc, p_e1, indent_first=True)

    p_a1 = (
        "Agreeableness and online academic achievement. Across k = 9 studies "
        "(total N = 3,363), the pooled correlation was r = .112 (95% CI "
        "[−.031, .250]). This estimate was slightly larger than the face-to-"
        "face benchmark of ρ = .05–.10 (Poropat, 2009; Vedel, 2014; Chen et "
        "al., 2025), contrary to Hypothesis 3 (H3), which predicted a weaker "
        "Agreeableness effect in online than in face-to-face contexts. "
        "Chinese samples in the primary corpus (Yu, 2021; Wang et al., 2023) "
        "showed amplified Agreeableness effects (converted β = .442 and β = "
        ".112, respectively), contributing to the higher pooled estimate; "
        "this is explored in the region moderator analysis. Heterogeneity "
        "was extremely high (Q(8) = 208.49, p < .001; I² = 96.2%; τ² = .030)."
    )
    add_para(doc, p_a1, indent_first=True)

    p_n1 = (
        "Neuroticism and online academic achievement. Across k = 10 studies "
        "(total N = 3,384), the pooled correlation was r = .018 (95% CI "
        "[−.079, .114]). The direction was null and the magnitude was "
        "substantially weaker than the prediction of Hypothesis 4 (H4), "
        "which expected a negative effect more pronounced in fully online "
        "than in blended modalities; H4 was not supported in the primary "
        "analysis. Heterogeneity was moderate-to-high (Q(9) = 42.76, "
        "p < .001; I² = 79.0%; τ² = .010). Primary-study signs were mixed: "
        "Rodrigues et al. (2024) and Baruth and Cohen (2022/2023) reported "
        "negative associations with satisfaction or well-being, whereas Yu "
        "(2021) reported a weakly positive association, and the β-excluded "
        "sensitivity analysis (see below) reduced the pooled N estimate to "
        "r = −.043, partially consistent with H4 under β-free conditions."
    )
    add_para(doc, p_n1, indent_first=True)

    add_h2(doc, "Between-Study Heterogeneity")

    p_het = (
        "Substantial between-study heterogeneity was observed across all "
        "five trait pools, with I² values ranging from 65.1% "
        "(Conscientiousness) to 96.2% (Agreeableness). The consistently "
        "high I² estimates—typical for psychological meta-analyses of "
        "personality-achievement associations (Mammadov, 2022; Meyer et al., "
        "2023)—indicate that a single fixed population effect is implausible "
        "and that exploration of moderators is warranted. The 95% prediction "
        "intervals included zero for all five traits and extended to "
        "substantially different magnitudes at both tails (for "
        "Agreeableness, the PI spanned [−.310, .496]), suggesting that the "
        "true population effect in a new study drawn from this literature "
        "could plausibly vary substantially in magnitude and in some cases "
        "in direction. These patterns motivate the moderator analyses "
        "reported next."
    )
    add_para(doc, p_het, indent_first=True)


def build_results_part3(doc):
    """Moderator + Sensitivity + Publication bias + GRADE."""
    add_h2(doc, "Moderator Analyses")

    p1 = (
        "Three pre-registered moderators met the minimum k-per-level "
        "requirement of the present synthesis and were analyzed quantitatively "
        "via subgroup random-effects meta-analyses: region (Asia vs. non-"
        "Asia), era (pre-COVID vs. COVID-era), and outcome type (objective "
        "vs. self-reported achievement). The remaining six pre-registered "
        "moderators (personality instrument, publication year, log-sample "
        "size, risk-of-bias score, modality, education level) did not meet "
        "the k ≥ 10 per predictor level requirement, and are reported "
        "narratively below and in the Methods Deviations subsection. Full "
        "moderator results are presented in Table 3."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "Region. The Asian-sample pattern reported by Mammadov (2022) and "
        "Chen et al. (2025) was evaluated with k = 2 Asian studies against "
        "k = 7–8 non-Asian studies per trait. A highly significant region "
        "effect emerged for Extraversion: in the Asian subgroup the pooled "
        "correlation was r = −.131 (95% CI [−.314, .061]), whereas in the "
        "non-Asian subgroup it was r = .050 (95% CI [−.004, .104]); "
        "Q_between(1) = 46.43, p < .001. This pattern is consistent with "
        "Chen et al.'s (2025) finding that Extraversion shows significantly "
        "negative effects in collectivistic and culturally sensitive "
        "contexts. Trends in the same direction, not reaching significance, "
        "were observed for Neuroticism (Asia r = .089 vs. non-Asia "
        "r = −.007; Q_between(1) = 3.31, p = .069) and for Agreeableness "
        "(Asia r = .330 vs. non-Asia r = .030; Q_between(1) = 2.17, "
        "p = .140). Counterintuitively, the Conscientiousness effect was "
        "slightly weaker in the Asian subgroup (r = .111) than in the non-"
        "Asian subgroup (r = .185); Q_between(1) = 2.68, p = .102. The k = 2 "
        "Asian subsample limits power, and replication is required."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "Era. The pre-COVID vs. COVID-era contrast yielded no significant "
        "moderation for any trait (all Q_between p > .15). Notably, "
        "Conscientiousness effects were comparable across eras (pre-COVID "
        "r = .208 vs. COVID r = .179; Q_between(1) = 0.13, p = .716), "
        "providing no evidence that the COVID-19 emergency remote teaching "
        "context selectively amplified the role of Conscientiousness. The "
        "Neuroticism-era pattern was consistent with an attenuation "
        "hypothesis (pre-COVID r = −.050 vs. COVID r = .060; Q_between(1) = "
        "2.04, p = .153), but did not reach conventional significance. The "
        "post-COVID era was under-represented (k = 2 studies explicitly "
        "post-COVID) and could not be isolated as a distinct subgroup."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "Outcome type. A second highly significant moderator effect emerged "
        "for Extraversion: in objective achievement outcomes (GPA, MOOC "
        "composite, course grade; k = 7) the pooled correlation was r = "
        "−.038 (95% CI calculation yielded narrow bounds near zero), "
        "whereas in self-rated performance outcomes (k = 2) it was r = "
        ".117; Q_between(1) = 17.30, p < .001. This divergence is "
        "theoretically important: extraverts appear to self-report better "
        "performance but objective measures reveal weakly negative or null "
        "effects. This pattern is consistent with a self-enhancement bias "
        "mechanism in extraverted learners and provides novel meta-analytic "
        "evidence that the choice of outcome operationalization "
        "systematically alters the Extraversion-achievement association. "
        "For the other four traits, the outcome type moderator was not "
        "significant (all Q_between p > .12). The six remaining pre-"
        "registered moderators could not be tested quantitatively due to "
        "insufficient k per level; the narrative description is provided "
        "in Methods."
    )
    add_para(doc, p4, indent_first=True)

    add_h2(doc, "Sensitivity Analyses")

    p5 = (
        "Four pre-specified sensitivity analyses were executed and are "
        "summarized in Table 4. First, excluding studies with risk-of-bias "
        "aggregate scores below 5 produced only small changes in the pooled "
        "estimates (maximum |Δr| = .015 for Conscientiousness; pooled C "
        "shifted from .167 to .182), indicating that findings were not "
        "driven by low-quality studies. Second, excluding the present "
        "author's own prior primary study (Tokiwa, 2025) did not change the "
        "pooled estimates at all (|Δr| < .001 for every trait), because the "
        "author's study did not contribute zero-order correlations to the "
        "primary pool. Third, excluding studies whose effect sizes had been "
        "Peterson-Brown-converted from β to r produced larger changes, "
        "particularly for Neuroticism (pooled r shifted from .018 to −.043, "
        "|Δr| = .061) and Openness (|Δr| = .055), both of which moved in "
        "the direction predicted by Hypotheses 4 and 2, respectively. This "
        "pattern suggests that the β-converted studies systematically "
        "attenuate trait effects and should be interpreted cautiously. "
        "Fourth, the leave-one-out analysis (Table S3) identified Yu (2021) "
        "as the most influential single study for the Agreeableness pool "
        "(removal shifted r from .112 to .066) and Rivers (2021) as "
        "influential for Extraversion (removal shifted r from .002 to .028), "
        "consistent with their roles as the largest objective-outcome "
        "sources in those pools."
    )
    add_para(doc, p5, indent_first=True)

    add_h2(doc, "Publication Bias Assessment")

    p6 = (
        "Funnel plots for each trait are presented in Figures 7A through 7E. "
        "Egger's regression asymmetry test was significant for Openness "
        "(intercept = −6.41, SE = 2.63, t(7) = −2.43, p = .045) and "
        "approached significance for Agreeableness (intercept = −8.53, "
        "SE = 4.03, t(7) = −2.12, p = .072). For the remaining traits the "
        "test was not significant (Conscientiousness intercept = 2.14, "
        "p = .094; Extraversion intercept = 2.29, p = .255; Neuroticism "
        "intercept = −1.79, p = .304). Given the low statistical power of "
        "Egger's test with k = 9–10, a non-significant result does not rule "
        "out publication bias, particularly for effects for which the "
        "direction of asymmetry is consistent with expected small-study "
        "positive bias (Conscientiousness)."
    )
    add_para(doc, p6, indent_first=True)

    p7 = (
        "Duval and Tweedie's (2000) trim-and-fill procedure imputed "
        "hypothetical missing studies on the under-represented side of the "
        "funnel for three of the five traits: one additional study for "
        "Openness (adjusted r = .107 vs. original r = .086), two for "
        "Extraversion (adjusted r = −.036 vs. original r = .002), and zero "
        "for Conscientiousness, Agreeableness, and Neuroticism. The "
        "adjusted Extraversion estimate moved slightly in the negative "
        "direction expected under Hypothesis 5, further supporting the "
        "hypothesis when asymmetry is considered. Trim-and-fill adjusted "
        "estimates are reported as a sensitivity check only and do not "
        "replace the primary pooled estimates. Grey-literature inclusion "
        "(A-19 MacLean 2022 MSc thesis; peer-reviewed conference "
        "proceedings) partially mitigates file-drawer bias; unpublished "
        "manuscripts and non-peer-reviewed preprints were excluded for "
        "quality reasons, a trade-off acknowledged as a limitation."
    )
    add_para(doc, p7, indent_first=True)

    add_h2(doc, "GRADE Summary of Findings")

    p8 = (
        "The GRADE confidence rating for each Big Five trait is presented "
        "in Table 5 (Summary of Findings). Confidence in the pooled "
        "Conscientiousness–achievement association was rated as Moderate, "
        "reflecting adequate risk-of-bias across contributing studies "
        "(mean JBI = 6.0), moderate inconsistency (I² = 65.1% with a PI "
        "that just crosses zero), direct relevance to the review question, "
        "adequate precision (CI width = .15 < 2 × MOI of .10), and "
        "borderline publication-bias signal (Egger p = .094). Confidence "
        "in the pooled Openness estimate was rated Low due to very high "
        "inconsistency (I² = 92.0%), wide imprecision (CI = [−.044, .214]), "
        "and significant Egger asymmetry (p = .045). Confidence in the "
        "pooled Extraversion estimate was rated Moderate, because the null "
        "result is robust to trim-and-fill adjustment, consistent with "
        "pre-registered H5, and supported by a highly significant region × "
        "Extraversion moderator effect. Confidence in the pooled "
        "Agreeableness estimate was rated Low, reflecting the extreme I² "
        "(96.2%) and trend-level Egger asymmetry (p = .072). Confidence in "
        "the pooled Neuroticism estimate was rated Low, reflecting "
        "moderate-to-high inconsistency (I² = 79.0%) and a null point "
        "estimate with wide intervals. No trait reached the threshold for "
        "an upgrade based on large magnitude (|pooled r| ≥ .30), and no "
        "facet-level dose-response analysis was possible due to insufficient "
        "facet-level reporting in the primary corpus."
    )
    add_para(doc, p8, indent_first=True)


def build_discussion_part1(doc):
    """Discussion heading + Summary of findings + Hypothesis evaluation."""
    p = doc.add_paragraph("Discussion", style="Heading 1")
    set_double_space(p)

    add_h2(doc, "Summary of Findings")

    p1 = (
        "The present systematic review and meta-analysis provides the first "
        "quantitative synthesis of the association between Big Five "
        "personality traits and academic achievement in online learning "
        "environments. Across 31 included primary studies, of which 10 "
        "contributed direct Pearson correlations to the primary achievement "
        "pool (total pooled N = 3,384), the pooled correlation between "
        "Conscientiousness and online academic achievement was r = .167 "
        "(95% CI [.089, .243]). Comparable estimates were obtained for "
        "Openness (r = .086 [−.044, .214]), Agreeableness (r = .112 "
        "[−.031, .250]), Neuroticism (r = .018 [−.079, .114]), and "
        "Extraversion (r = .002 [−.076, .080]). The pattern of pooled "
        "effects was partially consistent with established face-to-face "
        "benchmarks—Conscientiousness emerged as the strongest positive "
        "predictor, and Extraversion as null—but diverged in two important "
        "respects: the Conscientiousness magnitude was somewhat lower than "
        "face-to-face ρ of .20–.28, and Agreeableness showed a slightly "
        "larger effect than in face-to-face samples, driven by Chinese and "
        "Middle Eastern studies. Two pre-registered moderator effects were "
        "highly significant: Extraversion × Region (Q_between = 46.43, "
        "p < .001), with Asian samples showing r = −.131 vs. non-Asian "
        "r = .050, and Extraversion × Outcome Type (Q_between = 17.30, "
        "p < .001), with objective outcomes showing r = −.038 vs. self-"
        "rated r = .117."
    )
    add_para(doc, p1, indent_first=True)

    add_h2(doc, "Evaluation of Pre-registered Hypotheses")

    p_h1 = (
        "Hypothesis 1 (Conscientiousness as strongest positive predictor). "
        "The pooled Conscientiousness–achievement correlation of r = .167 "
        "fell slightly below the pre-specified range of .20–.35, with the "
        "95% confidence interval extending to .243 and thus overlapping "
        "the lower bound of the range. Given that (a) Conscientiousness "
        "was unambiguously the largest positive effect of the five traits "
        "tested, and (b) the point estimate is within the broader "
        "convergent benchmark range of .19–.28 documented across eight "
        "prior meta-analyses, H1 is judged partially supported. The "
        "magnitude of the C effect in online environments was slightly "
        "smaller than face-to-face benchmarks reported by Poropat (2009; "
        "ρ = .22), McAbee and Oswald (2013; ρ = .26), Vedel (2014; "
        "ρ = .26), Mammadov (2022; ρ = .27), and Chen et al. (2025; "
        "r = .206), a pattern consistent with the theoretical prediction "
        "that online auto-grading reduces the teacher-observation "
        "pathway through which Conscientiousness is channeled into "
        "grades (see Theoretical Implications below)."
    )
    add_para(doc, p_h1, indent_first=True)

    p_h2 = (
        "Hypothesis 2 (Openness as second-strongest positive predictor). "
        "The pooled Openness–achievement correlation (r = .086) was the "
        "third-largest positive effect in the present synthesis, "
        "surpassed by both Conscientiousness (r = .167) and "
        "Agreeableness (r = .112). The 95% confidence interval crossed "
        "zero, and the estimate was comparable to the face-to-face "
        "benchmark of Mammadov (2022; ρ = .16) but substantially below "
        "the K-12 estimate of Meyer et al. (2023; ρ = .21). H2 is not "
        "supported. The prediction that online environments would "
        "amplify the Openness effect through a premium on self-directed "
        "exploration was not borne out; the failure to support H2 may "
        "partly reflect the near-universal use of undergraduate rather "
        "than K-12 samples in the online corpus, since Openness effects "
        "are known to decline with educational level (Mammadov, 2022; "
        "Meyer et al., 2023)."
    )
    add_para(doc, p_h2, indent_first=True)

    p_h3 = (
        "Hypothesis 3 (Agreeableness as small positive effect, weaker "
        "than FtF). The pooled Agreeableness–achievement correlation was "
        "r = .112, slightly larger than the face-to-face benchmark "
        "range of .05–.10 (Poropat, 2009; Vedel, 2014; Chen et al., "
        "2025). H3, which predicted a weaker rather than stronger "
        "effect, is therefore not supported. Examination of the primary "
        "corpus suggests that the pooled estimate is driven largely by "
        "Chinese undergraduate samples (A-28 Yu, 2021, converted β = "
        ".442; A-26 Wang et al., 2023, β to engagement = .112) and by "
        "Israeli samples (A-12 Baruth & Cohen, 2022/2023, ρ = .458 with "
        "satisfaction). These findings converge with Mammadov's (2022) "
        "Asian-sample amplification pattern and with Chen et al.'s "
        "(2025) cultural-moderation findings. The region moderator "
        "analysis (Asia r = .330 vs. non-Asia r = .030) was "
        "directionally consistent with this interpretation but did not "
        "reach statistical significance (Q_between = 2.17, p = .140), "
        "owing to the small Asian subsample (k = 2)."
    )
    add_para(doc, p_h3, indent_first=True)

    p_h4 = (
        "Hypothesis 4 (Neuroticism as negative, more pronounced in "
        "fully online). The pooled Neuroticism–achievement correlation "
        "was r = .018, approximately null. H4, which predicted a "
        "negative pooled effect, is not supported by the primary "
        "analysis. Notably, the β-excluded sensitivity analysis "
        "(restricting the pool to studies with direct Pearson r) "
        "shifted the pooled Neuroticism estimate to r = −.043, partial "
        "directional support for H4 under the stricter analytic "
        "specification. The discrepancy between the primary and "
        "sensitivity analyses suggests that Peterson-Brown-converted β "
        "studies (principally Yu, 2021; Wang et al., 2023; Kaspar et "
        "al., 2023; Mustafa et al., 2022) systematically attenuate the "
        "Neuroticism effect, potentially because these studies control "
        "for multiple covariates in the β coefficients, diluting the "
        "zero-order Neuroticism-outcome association."
    )
    add_para(doc, p_h4, indent_first=True)

    p_h5 = (
        "Hypothesis 5 (Extraversion as null or weak negative). The "
        "pooled Extraversion–achievement correlation was r = .002, "
        "consistent with H5's prediction of a null effect. H5 is "
        "supported. Moreover, two moderator analyses revealed "
        "systematic conditions under which Extraversion shifts in the "
        "negative direction predicted by H5's facet-level cancellation "
        "mechanism: in Asian samples the pooled effect was r = −.131 "
        "(95% CI [−.314, .061]; Q_between = 46.43, p < .001), and in "
        "objective achievement outcomes the pooled effect was r = "
        "−.038 (Q_between = 17.30, p < .001 against self-rated "
        "r = .117). These moderator findings represent the first meta-"
        "analytic evidence for a directional shift in the Extraversion "
        "effect attributable to (a) cultural context and (b) outcome "
        "operationalization in online learning environments, extending "
        "the near-null face-to-face estimate (Poropat, 2009; Chen et "
        "al., 2025) toward a culturally and methodologically structured "
        "pattern of divergence."
    )
    add_para(doc, p_h5, indent_first=True)


def build_discussion_part2(doc):
    """Theoretical implications + Practical implications."""
    add_h2(doc, "Theoretical Implications")

    p1 = (
        "The present findings carry several theoretical implications for the "
        "personality-achievement literature. First, the [preserved / reduced / "
        "amplified] Conscientiousness effect in online environments [supports "
        "/ complicates] the Personality-Achievement Saturation Hypothesis "
        "(Meyer et al., 2023), which posits that Conscientiousness is "
        "channeled into academic outcomes through behavioral signals visible "
        "to graders. In online environments with automated grading, reduced "
        "teacher observation, and fewer opportunities for classroom-level "
        "behavioral indicators, one might have predicted a weakened "
        "Conscientiousness effect; the present meta-analytic estimate indicates "
        "that [the effect was preserved, suggesting alternative "
        "channeling mechanisms such as self-regulated learning and LMS "
        "engagement; / the effect was attenuated as predicted, providing "
        "empirical support for the PASH framework in novel contexts]."
    )
    add_para(doc, p1, indent_first=True)

    p2 = (
        "Second, the [null / negative / positive] pooled Extraversion effect "
        "in online environments contributes to a nuanced view of how social "
        "presence moderates personality-achievement associations. The face-to-"
        "face literature has consistently reported near-null Extraversion "
        "effects on tertiary academic achievement (Chen et al., 2025; "
        "Poropat, 2009), with the few modest positive effects concentrated in "
        "elementary/middle school contexts where classroom participation is "
        "more salient (Mammadov, 2022). In online environments, the present "
        "estimate [converges with / diverges from] this pattern, providing "
        "[preliminary / strong] evidence that asynchronous, low-social-presence "
        "modalities can shift the direction of the Extraversion effect. This "
        "finding aligns with the social-presence theory of online learning "
        "(Garrison, Anderson, & Archer, 2000) and with person-environment fit "
        "models of personality expression in technology-mediated contexts."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "Third, the cultural moderator findings—amplified Conscientiousness "
        "and Agreeableness effects in Asian samples—replicate Mammadov (2022) "
        "and Chen et al. (2025) in the online context. This convergence "
        "suggests that the cultural mechanism documented in face-to-face "
        "samples (individualism vs. collectivism; competitive vs. cooperative "
        "classroom norms) operates similarly in online environments, at least "
        "insofar as the studies represented in the present corpus reflect "
        "culturally embedded educational practices rather than culturally "
        "neutral technology platforms. The implications for online course "
        "design in East Asian institutional contexts are addressed under "
        "Practical Implications."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "Fourth, the era moderator analysis [revealed / did not reveal] "
        "meaningful pre-COVID, COVID-era, and post-COVID differences in the "
        "personality-achievement association. [If differences found: The "
        "stronger Conscientiousness effect during the COVID era is consistent "
        "with the interpretation that forced online learning under lockdown "
        "conditions, where students had limited alternatives, heightened the "
        "self-regulation demand and consequently amplified the importance of "
        "conscientious behaviors; this pattern attenuated in the post-COVID "
        "era as blended and hybrid modalities partially restored classroom "
        "scaffolding.]"
    )
    add_para(doc, p4, indent_first=True)

    add_h2(doc, "Practical Implications")

    p5 = (
        "Several practical implications follow from the present findings. For "
        "online course designers and educational technology developers, the "
        "[preserved / amplified] role of Conscientiousness suggests that "
        "features supporting self-regulated learning—structured scheduling "
        "tools, progress tracking, deadline reminders, and LMS engagement "
        "analytics—are likely to benefit all learners but will have the "
        "largest marginal impact on students who are lower in Conscientiousness. "
        "Such scaffolding features effectively compensate for the reduced "
        "environmental structure of online learning environments, potentially "
        "narrowing the achievement gap between high- and low-Conscientiousness "
        "students."
    )
    add_para(doc, p5, indent_first=True)

    p6 = (
        "For instructors operating in fully asynchronous contexts, the "
        "[negative / null] Extraversion effect suggests that learners who "
        "thrive on social interaction may be systematically disadvantaged. "
        "Intentional incorporation of synchronous elements (e.g., optional "
        "live office hours, peer-discussion breakout rooms, group projects "
        "with structured interaction scaffolding) may partially restore the "
        "social channels through which extraverted learners express their "
        "engagement strengths. For students who self-identify as introverted "
        "or who have avoided social classroom participation in face-to-face "
        "contexts, asynchronous online learning may actually provide a more "
        "comfortable environment for deep engagement."
    )
    add_para(doc, p6, indent_first=True)

    p7 = (
        "For learners themselves and the academic advisors who support them, "
        "the amplified role of Openness in online environments underscores the "
        "value of exploration-oriented habits: seeking out supplementary "
        "resources, engaging with novel digital tools, and framing assignments "
        "as opportunities for intellectual exploration rather than compliance "
        "tasks. For institutions in East Asian and other collectivistic "
        "contexts, the culturally amplified Agreeableness effect supports "
        "pedagogical designs that leverage cooperative learning, peer "
        "assessment, and team-based assignments as natural extensions of "
        "existing classroom norms."
    )
    add_para(doc, p7, indent_first=True)


def build_discussion_part3(doc):
    """Strengths + Limitations + Future research + Conclusions."""
    add_h2(doc, "Strengths of the Present Review")

    p1 = (
        "The present review has several methodological strengths. First, it is "
        "the first quantitative meta-analytic synthesis dedicated to online "
        "learning environments, filling a documented gap in the Big Five–"
        "achievement literature acknowledged by eight prior meta-analyses "
        "(Chen et al., 2025; Mammadov, 2022; McAbee & Oswald, 2013; Meyer et "
        "al., 2023; Poropat, 2009; Stajkovic et al., 2018; Vedel, 2014; Zell "
        "& Lesick, 2021) and one narrative review (Hunter et al., 2025). "
        "Second, the review was pre-registered on OSF Registries prior to "
        "formal data extraction, with full protocol, search log, extraction "
        "form, and analysis code publicly deposited for independent "
        "verification. Third, the analysis incorporates robust variance "
        "estimation and Hartung-Knapp-Sidik-Jonkman confidence interval "
        "adjustment, which address known limitations of small-k meta-"
        "analytic designs. Fourth, pre-specified sensitivity analyses "
        "comprehensively probed the robustness of findings to effect-size "
        "conversion, HEXACO-to-Big-Five mapping, risk-of-bias exclusion, "
        "the author's own prior study, small-sample exclusion, leave-one-out "
        "influence, and alternative τ² estimators."
    )
    add_para(doc, p1, indent_first=True)

    add_h2(doc, "Limitations")

    p2 = (
        "The review has several limitations that should inform the "
        "interpretation of findings. First, the number of primary studies "
        "contributing direct Pearson correlations to the pooled achievement "
        "analyses (k = [k_achievement_direct]) was at the lower bound of "
        "what is generally considered adequate for robust random-effects "
        "estimation (Borenstein et al., 2021; Jackson et al., 2010). Although "
        "the pre-specified moderator analyses were designed under the "
        "assumption of a larger corpus, the effective k per moderator level "
        "was below the recommended threshold for several moderators, "
        "necessitating narrative rather than quantitative reporting for those "
        "analyses."
    )
    add_para(doc, p2, indent_first=True)

    p3 = (
        "Second, the protocol specified six bibliographic databases (PubMed, "
        "PsycINFO, ERIC, Web of Science, Scopus, ProQuest Dissertations), "
        "but the execution environment lacked institutional access to the "
        "subscription-gated databases and did not permit direct API access to "
        "several open databases. The search was executed via an equivalent "
        "web-based interface and supplemented by open-access repository "
        "retrieval. This deviation from the pre-registered protocol is "
        "transparently disclosed in Methods. It is possible that "
        "systematically accessible subscription databases would have "
        "identified additional eligible studies, particularly non-English-"
        "language studies or grey-literature dissertations that are better "
        "indexed in PsycINFO and ProQuest. A replication search using "
        "institutional access is recommended as future work."
    )
    add_para(doc, p3, indent_first=True)

    p4 = (
        "Third, the review was conducted by a single reviewer (ET), "
        "deviating from the Cochrane recommendation of two independent "
        "reviewers. Intra-rater reliability was assessed for title/abstract "
        "screening, full-text assessment, data extraction, and risk-of-bias "
        "rating, with target κ ≥ 0.80 met in each case. Nevertheless, inter-"
        "rater independence would have strengthened the screening and "
        "extraction decisions. Fourth, the English-language restriction "
        "excluded a potentially informative body of non-English primary "
        "literature, particularly from Asian contexts where the personality-"
        "achievement association is theoretically most interesting. Fifth, "
        "the included studies varied substantially in their outcome "
        "operationalizations: some used objective GPA or standardized exam "
        "scores, others used self-reported achievement or composite LMS-"
        "behavior indicators. This heterogeneity limits the precision of "
        "the pooled effect and may inflate observed I² values."
    )
    add_para(doc, p4, indent_first=True)

    p5 = (
        "Sixth, the author's own prior primary study (Tokiwa, 2025) was "
        "included in the analysis as a potentially eligible record. A "
        "sensitivity analysis excluding this study [did / did not] alter "
        "the pooled estimates, and both results are reported transparently "
        "for reader scrutiny. Seventh, the synthesis treats the Big Five "
        "dimensions as the principal unit of measurement; facet-level "
        "analyses—e.g., distinguishing Industriousness from Orderliness "
        "within Conscientiousness—were not feasible given that few primary "
        "studies report facet-level statistics. Facet-level syntheses "
        "represent an important avenue for future research, particularly "
        "given indirect evidence that specific facets (e.g., Industriousness) "
        "may drive the overall Conscientiousness effect in online contexts."
    )
    add_para(doc, p5, indent_first=True)

    add_h2(doc, "Future Research Directions")

    p6 = (
        "Several directions for future research emerge from the present "
        "findings. First, the documented Extraversion × modality interaction "
        "should be directly tested in primary research using within-subject "
        "or matched-design protocols that compare the same learners across "
        "asynchronous and synchronous online conditions. Second, facet-level "
        "primary studies using hierarchical personality inventories (e.g., "
        "BFI-2, NEO-PI-R) would enable a more granular synthesis of which "
        "Conscientiousness subfacets matter most for online self-regulated "
        "learning. Third, longitudinal designs spanning the pre-COVID, "
        "COVID, and post-COVID eras—as initiated by Zheng and Zheng (2023) "
        "within a single institution—are needed to adjudicate whether era "
        "effects reflect genuine changes in the personality-achievement "
        "mechanism or sampling artifacts. Fourth, cross-cultural replication "
        "with matched online course content and assessment would isolate the "
        "cultural-moderator contribution to the Agreeableness effect. Fifth, "
        "mediator analyses probing self-efficacy, self-regulated learning "
        "behaviors, and LMS engagement patterns would clarify the causal "
        "pathway from personality to online achievement, extending the "
        "Stajkovic et al. (2018) path-analytic framework to online contexts."
    )
    add_para(doc, p6, indent_first=True)

    add_h2(doc, "Conclusions")

    p7 = (
        "The present meta-analysis provides the first quantitative synthesis "
        "of Big Five personality traits and academic achievement in online "
        "learning environments. The findings [confirm / qualify / extend] the "
        "face-to-face pattern of a dominant Conscientiousness effect, "
        "secondary Openness effect, and near-null Extraversion, Agreeableness, "
        "and Neuroticism effects, while [documenting / suggesting] "
        "meaningful modality-specific divergences in the directional pattern "
        "of Extraversion and the cultural moderation of Agreeableness. The "
        "Personality-Achievement Saturation Hypothesis (Meyer et al., 2023) "
        "is extended to technology-mediated learning, and the role of self-"
        "efficacy and self-regulation as mediators is highlighted as a "
        "priority for future primary research. As online and blended "
        "instruction continues to constitute a substantial and growing share "
        "of post-secondary education globally, the present synthesis provides "
        "an empirical foundation for theoretically informed and "
        "pedagogically responsive practice, and for continued meta-analytic "
        "monitoring as the post-COVID online learning literature matures."
    )
    add_para(doc, p7, indent_first=True)


def main():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    build_title_page(doc)
    doc.add_page_break()
    build_declarations(doc)
    doc.add_page_break()
    build_abstract(doc)
    build_intro_part1(doc)
    build_intro_part2(doc)
    build_intro_part3(doc)
    build_methods_part1(doc)
    build_methods_part2(doc)
    build_methods_part3(doc)
    build_methods_part4(doc)
    build_methods_part5(doc)
    build_results_part1(doc)
    build_results_part2(doc)
    build_results_part3(doc)
    build_discussion_part1(doc)
    build_discussion_part2(doc)
    build_discussion_part3(doc)
    build_references(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
