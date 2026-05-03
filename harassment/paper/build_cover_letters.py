"""Generate journal-tailored cover letters for the harassment manuscript.

Targets (paid OA, predatory-free, within ¥200k-¥500k @ ¥160/USD;
already-submitted journals BMC Psychology / Frontiers in Psychology /
Current Psychology / Behavioral Sciences / SAGE Open are excluded):

- Heliyon (Elsevier)                       ~$2,300  ≈ ¥368k
- Acta Psychologica (Elsevier hybrid)      ~$2,820  ≈ ¥451k
- PeerJ Life & Environment                 ~$1,395  ≈ ¥223k
- Royal Society Open Science               ~$1,800  ≈ ¥288k
- Discover Psychology (Springer Nature)    ~$1,290  ≈ ¥206k
- Healthcare (MDPI)                        CHF 2,700 ≈ ¥432k

Improvements over the uploaded template
(`Cover Letter (need edit).docx`, `Cover Letter_Behavioral Sciences.docx`):

1. Sole-author voice: "our manuscript / We analyze / we believe"
   → "my manuscript / I analyze / I believe" (the manuscript Title page
   declares Eisuke Tokiwa as the only author).
2. Preprint disclosure added (Research Square v1,
   https://doi.org/10.21203/rs.3.rs-7756124/v1) — required by most
   target journals' submission policies.
3. Applied-implication wording softened:
   "underscoring the applied value of integrating HEXACO into risk
   assessment and prevention" → "informing tentative implications for
   harassment-prevention frameworks". Cross-sectional self-report design
   does not warrant deterministic claims about selection or screening.
4. Journal-fit paragraphs tailored per venue (scope, audience, indexing,
   relevant prior publications in the journal).

Run:
    python3 harassment/paper/build_cover_letters.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


HERE = Path(__file__).resolve().parent

# ---------------------------------------------------------------------- #
# Letterhead / signature                                                 #
# ---------------------------------------------------------------------- #
HEADER_LINES = [
    "Eisuke Tokiwa",
    "Founder, SUNBLAZE Co., Ltd.",
    "6th floor, Nishi-Shinjuku Mizuma Building",
    "3-3-13 Nishi-Shinjuku, Shinjuku-ku, Tokyo 160-0023, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Tel: +81-80-7279-4835",
    "Email: eisuke.tokiwa@sunblaze.jp",
]

SIGNATURE_LINES = [
    "Sincerely,",
    "",
    "Eisuke Tokiwa",
    "Founder, SUNBLAZE Co., Ltd.",
    "Tokyo, Japan",
    "ORCID: 0009-0009-7124-6669",
    "Email: eisuke.tokiwa@sunblaze.jp",
]


# ---------------------------------------------------------------------- #
# Shared paragraphs (identical across letters)                           #
# ---------------------------------------------------------------------- #
def submission_para(journal: str) -> str:
    return (
        f"I am pleased to submit my manuscript entitled "
        f"“Associations of HEXACO and Dark Traits with Power and "
        f"Gender Harassment” for consideration in {journal}."
    )


CONTRIBUTION = (
    "__Contribution.__ "
    "This study examines whether the HEXACO model—especially "
    "Honesty–Humility—adds predictive value beyond the Dark "
    "Triad (Machiavellianism, narcissism, psychopathy) for self-reported "
    "workplace harassment perpetration tendencies in a Japanese (non-"
    "WEIRD) employed adult sample (N = 354). I evaluate both power "
    "harassment and gender harassment, addressing two questions of "
    "theoretical and applied interest: (a) do broad moral-personality "
    "dimensions provide incremental prediction over antagonistic dark "
    "traits, and (b) do these associations generalize outside Western "
    "WEIRD contexts?"
)

METHODS = (
    "__Methods and key findings.__ "
    "Using Spearman rank correlations and HC3-robust hierarchical "
    "multiple regressions, I specify three nested models (Model A: "
    "controls + Dark Triad; Model B: + HEXACO; Model C: Honesty–"
    "Humility × Dark Triad interactions), with sensitivity analyses "
    "excluding high-influence cases (Cook’s D > 4/n) and "
    "multicollinearity diagnostics (VIF ≤ 2.00). Across models, "
    "psychopathy is a consistent positive predictor of power harassment "
    "(β = .32–.40, p < .001), whereas Honesty–Humility is "
    "negatively associated (β = −.14, p = .049). For gender "
    "harassment, narcissism and psychopathy are positive predictors, "
    "while Honesty–Humility (β = −.23, p < .001) and "
    "Openness (β = −.24, p < .001) are negative predictors; "
    "Machiavellianism shows a negative coefficient consistent with a "
    "statistical suppression pattern once variance shared with narcissism "
    "and psychopathy is partialed out. Adding HEXACO yields meaningful "
    "incremental variance (ΔR² = .032, p = .036 for power "
    "harassment; ΔR² = .096, p < .001 for gender harassment), "
    "informing tentative implications for harassment-prevention "
    "frameworks rather than deterministic personnel-selection rules."
)

COMPLIANCE = (
    "__Compliance and transparency.__ "
    "Ethics approval: Institutional Review Board of the Public Health "
    "Research Foundation (Reference no. 25G0001). Informed consent was "
    "obtained from all participants; data were anonymised and stored "
    "securely. The author declares no conflicts of interest and reports "
    "no external funding."
)

OPEN_SCIENCE = (
    "__Open science and preprint.__ "
    "A preprint of this manuscript is publicly available on Research "
    "Square (v1; https://doi.org/10.21203/rs.3.rs-7756124/v1). The "
    "de-identified dataset and analysis script are openly available at "
    "https://github.com/etoki/paper (harassment/raw.csv, harassment/"
    "analysis.py)."
)

NONDUP = (
    "This manuscript is original and is not currently under consideration "
    "by any other journal. The Research Square preprint (v1) is "
    "explicitly disclosed above; any substantive revisions during peer "
    "review will be reflected in a versioned preprint update under the "
    "same DOI stem. I am the sole author and accept full accountability "
    "for all aspects of the work."
)

CLOSING = (
    "Thank you for considering my submission. I look forward to your "
    "decision."
)


# ---------------------------------------------------------------------- #
# Per-journal fit paragraphs                                             #
# ---------------------------------------------------------------------- #
JOURNALS: dict[str, dict[str, object]] = {
    "Heliyon": {
        "addressee": ["To the Editors,", "Heliyon (Elsevier)"],
        "section_suggestion": (
            "I suggest the manuscript be handled within the "
            "“Psychology” section."
        ),
        "fit": (
            "__Fit with Heliyon.__ "
            "Heliyon’s broad multidisciplinary scope and emphasis on "
            "methodologically sound research — regardless of perceived "
            "novelty premium — align directly with the present study’s "
            "transparent reporting of pre-specified hierarchical models, "
            "sensitivity checks, and openly deposited data and code. The "
            "manuscript intersects personality science, organizational "
            "behavior, and cross-cultural psychology, all areas in which "
            "Heliyon publishes regularly."
        ),
        "extras": (
            "All Elsevier submission requirements are addressed: ORCID "
            "in the header, structured declarations, data availability "
            "statement, and CRediT statement (sole author)."
        ),
    },
    "Acta Psychologica": {
        "addressee": ["To the Editors,", "Acta Psychologica (Elsevier)"],
        "section_suggestion": "",
        "fit": (
            "__Fit with Acta Psychologica.__ "
            "Acta Psychologica has a long-standing record of publishing "
            "individual-differences research that bridges personality "
            "structure and applied behavioural outcomes. The present "
            "manuscript addresses the journal’s core interest in "
            "rigorous correlational psychology by quantifying incremental "
            "validity of the HEXACO model relative to the Dark Triad in a "
            "non-Western sample — an underrepresented context in the "
            "individual-differences literature."
        ),
        "extras": (
            "All Elsevier submission requirements are addressed: ORCID, "
            "structured declarations, data availability, and CRediT "
            "statement (sole author)."
        ),
    },
    "PeerJ": {
        "addressee": ["To the Editors,", "PeerJ Life & Environment"],
        "section_suggestion": (
            "Given PeerJ’s subject-area structure, I suggest "
            "handling under “Psychology / Industrial-Organisational "
            "Psychology”."
        ),
        "fit": (
            "__Fit with PeerJ.__ "
            "PeerJ’s editorial policy of evaluating submissions on "
            "scientific and methodological soundness rather than perceived "
            "impact is well matched to a confirmatory cross-sectional "
            "investigation with pre-specified hierarchical models, robust "
            "standard errors, sensitivity analyses, and fully open data "
            "and code. The integration of HEXACO and Dark Triad in a "
            "Japanese employee sample contributes incremental cross-"
            "cultural evidence relevant to PeerJ’s readership in "
            "behavioural and organisational research."
        ),
        "extras": (
            "All PeerJ submission requirements are addressed: ORCID, "
            "data availability (linked to public GitHub repository), "
            "and a sole-author declaration."
        ),
    },
    "Royal Society Open Science": {
        "addressee": [
            "To the Editors,",
            "Royal Society Open Science (The Royal Society)",
        ],
        "section_suggestion": (
            "The manuscript would best fit the “Psychology and "
            "cognitive neuroscience” subject area."
        ),
        "fit": (
            "__Fit with Royal Society Open Science.__ "
            "The journal’s emphasis on transparent, reproducible, "
            "open research aligns with the present submission, which "
            "deposits both the de-identified raw dataset and the full "
            "analytic pipeline in a public repository, pre-specifies "
            "hierarchical models, and reports sensitivity analyses. The "
            "manuscript contributes cross-cultural evidence on "
            "individual differences and workplace behaviour that should "
            "interest readers across personality, social, and "
            "organisational psychology."
        ),
        "extras": (
            "All Royal Society submission requirements are addressed: "
            "ORCID, ethics statement, data accessibility statement linked "
            "to the public repository, and authorship statement (sole "
            "author)."
        ),
    },
    "Discover Psychology": {
        "addressee": [
            "To the Editors,",
            "Discover Psychology (Springer Nature)",
        ],
        "section_suggestion": "",
        "fit": (
            "__Fit with Discover Psychology.__ "
            "Discover Psychology’s broad and inclusive scope across "
            "subfields of psychology — including personality, "
            "organisational, and cross-cultural research — is an "
            "appropriate venue for an empirical study that integrates the "
            "HEXACO and Dark Triad frameworks within a non-Western "
            "employed sample. The journal’s commitment to open and "
            "reproducible research is mirrored by the publicly available "
            "dataset and analysis script accompanying this submission."
        ),
        "extras": (
            "All Springer Nature submission requirements are addressed: "
            "ORCID, structured declarations (ethics, consent, conflicts, "
            "funding, data availability), and a sole-author CRediT "
            "statement."
        ),
    },
    "Healthcare": {
        "addressee": ["To the Editors,", "Healthcare (MDPI)"],
        "section_suggestion": (
            "The manuscript fits the “Mental Health, Workplace "
            "Health, and Well-being” topical area."
        ),
        "fit": (
            "__Fit with Healthcare.__ "
            "Workplace harassment is a major occupational-health risk "
            "associated with deteriorated mental health, increased "
            "turnover, and reduced organisational well-being. Healthcare "
            "regularly publishes empirical work on psychosocial "
            "occupational health, and the present manuscript contributes "
            "individual-difference evidence on perpetration tendencies "
            "from a Japanese employee sample — a population context "
            "still under-represented in the workplace-health literature. "
            "The findings are directly relevant to readers concerned with "
            "evidence-based prevention strategies and organisational "
            "well-being."
        ),
        "extras": (
            "All MDPI submission requirements are addressed: ORCID, "
            "structured declarations conforming to the Healthcare "
            "template, and data availability linked to a public "
            "repository (sole author)."
        ),
    },
}


# ---------------------------------------------------------------------- #
# docx writer                                                            #
# ---------------------------------------------------------------------- #
def add_paragraph(doc: Document, text: str, *, space_after_pt: int = 6) -> None:
    """Add a paragraph; supports inline bold via __segment__ markers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.line_spacing = 1.15

    parts = text.split("__")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True


def build_letter(journal: str, info: dict, out_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Letterhead
    for line in HEADER_LINES:
        add_paragraph(doc, line, space_after_pt=0)
    add_paragraph(doc, "", space_after_pt=0)

    # Date
    add_paragraph(doc, date.today().strftime("%B %d, %Y"))
    add_paragraph(doc, "", space_after_pt=0)

    # Addressee
    for line in info["addressee"]:
        add_paragraph(doc, line, space_after_pt=0)
    add_paragraph(doc, "", space_after_pt=0)

    # Body
    add_paragraph(doc, "Dear Editors,")
    sub = submission_para(journal)
    if info.get("section_suggestion"):
        sub = sub + " " + info["section_suggestion"]
    add_paragraph(doc, sub)
    add_paragraph(doc, CONTRIBUTION)
    add_paragraph(doc, METHODS)
    add_paragraph(doc, info["fit"])
    add_paragraph(doc, COMPLIANCE)
    add_paragraph(doc, OPEN_SCIENCE)
    add_paragraph(doc, NONDUP)
    if info.get("extras"):
        add_paragraph(doc, info["extras"])
    add_paragraph(doc, CLOSING)

    # Signature
    for line in SIGNATURE_LINES:
        add_paragraph(doc, line, space_after_pt=0)

    doc.save(str(out_path))


def main() -> int:
    for journal, info in JOURNALS.items():
        out = HERE / f"Cover Letter_{journal}.docx"
        build_letter(journal, info, out)
        print(f"Wrote {out.relative_to(HERE.parent.parent)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
